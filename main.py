import os
import customtkinter as ctk
from customtkinter import CTkImage  # Add this import at the top
import tkinter as tk
from tkinter import messagebox
from PIL import Image, ImageDraw
from datetime import datetime
import comtypes.client
import webbrowser
import pandas as pd
import openpyxl
from docx import Document
from docx.shared import Pt
from tkinter import ttk
import tempfile
from copy import deepcopy
import shutil
import os
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx import Document
from tkinter import filedialog
import sys
from docxtpl import DocxTemplate # <<< NEW IMPORT
import json # New import for saving configurations
import re   # New import for logic parsing
from tkinter import simpledialog # For the logic input


TEMPLATE_FILE = "Invoice_Template_No_Borders.docx"
BILL_FOLDER = "bill"
BILL_COUNTER_FILE = "bill_counter.txt"
LOGO_PATH = os.path.abspath("logo.png")
EXCEL_FILE = "bills.xlsx"
DOCTORS_FILE = "doctors.txt"
AGENTS_FILE = "agents.txt"

REPORT_TEMPLATES_FOLDER = "ReportTemplates" 
OUTPUT_FOLDER = "GeneratedReports"

CUSTOM_REPORTS_FILE = "custom_reports.json"



ctk.set_appearance_mode("System")
ctk.set_default_color_theme("green")


class ManageTestsFrame(ctk.CTkFrame):
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller
        
        # --- 1. Top Header ---
        header_frame = ctk.CTkFrame(self, fg_color="transparent")
        header_frame.pack(fill="x", padx=20, pady=10)
        
        # Back Button
        ctk.CTkButton(header_frame, text="⬅ Back to Dashboard", width=150, height=40,
                      fg_color="gray", hover_color="#555",
                      command=self.go_back).pack(side="left")
        
        # Title
        ctk.CTkLabel(header_frame, text="Manage Tests Database", 
                     font=ctk.CTkFont(size=24, weight="bold")).pack(side="left", padx=20)

        # --- 2. Add New Test Section ---
        add_frame = ctk.CTkFrame(self)
        add_frame.pack(fill="x", padx=20, pady=10)
        
        ctk.CTkLabel(add_frame, text="Add New Test:", font=ctk.CTkFont(weight="bold")).pack(side="left", padx=15, pady=15)
        
        self.new_test_name = ctk.CTkEntry(add_frame, placeholder_text="Test Name", width=300)
        self.new_test_name.pack(side="left", padx=10)
        
        self.new_test_rate = ctk.CTkEntry(add_frame, placeholder_text="Amount (₹)", width=150)
        self.new_test_rate.pack(side="left", padx=10)
        
        ctk.CTkButton(add_frame, text="➕ Add Test", fg_color="#2E8B57", hover_color="#245c3d",
                      command=self.add_test).pack(side="left", padx=10)

        # --- 3. Tests List Section ---
        list_container = ctk.CTkFrame(self)
        list_container.pack(fill="both", expand=True, padx=20, pady=(0, 20))
        
        # Header for the list
        labels_frame = ctk.CTkFrame(list_container, fg_color="transparent", height=30)
        labels_frame.pack(fill="x", padx=10, pady=5)
        ctk.CTkLabel(labels_frame, text="Test Name", anchor="w").pack(side="left", fill="x", expand=True, padx=10)
        ctk.CTkLabel(labels_frame, text="Amount", width=100).pack(side="left", padx=10)
        ctk.CTkLabel(labels_frame, text="Actions", width=150).pack(side="left", padx=10)

        # Scrollable area
        self.scroll_frame = ctk.CTkScrollableFrame(list_container)
        self.scroll_frame.pack(fill="both", expand=True, padx=10, pady=5)

        # Load data immediately
        self.refresh_list()

    def go_back(self):
        """Return to dashboard and reload data"""
        self.controller.load_test_data() # Refresh main app data
        self.controller.show_dashboard()

    def refresh_list(self):
        # Clear list
        for widget in self.scroll_frame.winfo_children():
            widget.destroy()
            
        # Load data from file directly to ensure sync
        test_data = {}
        if os.path.exists("test_amount.txt"):
            with open("test_amount.txt", "r", encoding="utf-8") as f:
                for line in f:
                    try:
                        name, amt = line.strip().split(" - ")
                        test_data[name] = amt
                    except: continue

        # Populate rows
        for name, amount in test_data.items():
            row = ctk.CTkFrame(self.scroll_frame, fg_color="transparent")
            row.pack(fill="x", pady=2)
            
            # Name
            ctk.CTkLabel(row, text=name, anchor="w").pack(side="left", fill="x", expand=True, padx=10)
            
            # Amount
            ctk.CTkLabel(row, text=f"₹{amount}", width=100).pack(side="left", padx=10)
            
            # Buttons
            btn_frame = ctk.CTkFrame(row, fg_color="transparent", width=150)
            btn_frame.pack(side="left")
            
            # Edit Button (Pencil)
            ctk.CTkButton(btn_frame, text="✎", width=40, fg_color="#F39C12", hover_color="#D68910",
                          command=lambda n=name, a=amount: self.edit_test_popup(n, a)).pack(side="left", padx=2)
            
            # Delete Button (Trash)
            ctk.CTkButton(btn_frame, text="🗑", width=40, fg_color="#C0392B", hover_color="#922B21",
                          command=lambda n=name: self.delete_test(n)).pack(side="left", padx=2)

    def add_test(self):
        name = self.new_test_name.get().strip()
        rate = self.new_test_rate.get().strip()
        
        if not name or not rate:
            return
            
        try:
            float(rate) # Validation
            with open("test_amount.txt", "a", encoding="utf-8") as f:
                f.write(f"{name} - {rate}\n")
            
            self.new_test_name.delete(0, "end")
            self.new_test_rate.delete(0, "end")
            self.refresh_list()
        except ValueError:
            messagebox.showerror("Error", "Invalid Amount")

    def delete_test(self, name_to_delete):
        if not messagebox.askyesno("Confirm", f"Delete '{name_to_delete}'?"):
            return
            
        # Read all, filter out deleted, write back
        lines = []
        if os.path.exists("test_amount.txt"):
            with open("test_amount.txt", "r", encoding="utf-8") as f:
                lines = f.readlines()
        
        with open("test_amount.txt", "w", encoding="utf-8") as f:
            for line in lines:
                if not line.startswith(f"{name_to_delete} - "):
                    f.write(line)
        
        self.refresh_list()

    def edit_test_popup(self, old_name, old_rate):
        # Simple popup for editing
        dialog = ctk.CTkToplevel(self)
        dialog.title("Edit Test")
        dialog.geometry("400x200")
        dialog.attributes('-topmost', True)
        
        ctk.CTkLabel(dialog, text="Test Name:").pack(pady=5)
        name_var = ctk.StringVar(value=old_name)
        ctk.CTkEntry(dialog, textvariable=name_var, width=250).pack(pady=5)
        
        ctk.CTkLabel(dialog, text="Amount:").pack(pady=5)
        rate_var = ctk.StringVar(value=str(old_rate))
        ctk.CTkEntry(dialog, textvariable=rate_var, width=250).pack(pady=5)
        
        def save_edit():
            new_n = name_var.get().strip()
            new_r = rate_var.get().strip()
            if not new_n or not new_r: return
            
            # Delete old, add new logic
            self.delete_test_silently(old_name)
            with open("test_amount.txt", "a", encoding="utf-8") as f:
                f.write(f"{new_n} - {new_r}\n")
            
            dialog.destroy()
            self.refresh_list()

        ctk.CTkButton(dialog, text="Save Changes", command=save_edit).pack(pady=20)

    def delete_test_silently(self, name_to_delete):
        """Helper to delete without confirmation dialog (used during edit)"""
        lines = []
        if os.path.exists("test_amount.txt"):
            with open("test_amount.txt", "r", encoding="utf-8") as f:
                lines = f.readlines()
        with open("test_amount.txt", "w", encoding="utf-8") as f:
            for line in lines:
                if not line.startswith(f"{name_to_delete} - "):
                    f.write(line)



class ManageDoctorsFrame(ctk.CTkFrame):
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller
        
        # --- 1. Top Header ---
        header_frame = ctk.CTkFrame(self, fg_color="transparent")
        header_frame.pack(fill="x", padx=20, pady=10)
        
        # Back Button
        ctk.CTkButton(header_frame, text="⬅ Back to Dashboard", width=150, height=40,
                      fg_color="gray", hover_color="#555",
                      command=self.go_back).pack(side="left")
        
        # Title
        ctk.CTkLabel(header_frame, text="Manage Doctors Database", 
                     font=ctk.CTkFont(size=24, weight="bold")).pack(side="left", padx=20)

        # --- 2. Add New Doctor Section ---
        add_frame = ctk.CTkFrame(self)
        add_frame.pack(fill="x", padx=20, pady=10)
        
        ctk.CTkLabel(add_frame, text="Add New Doctor:", font=ctk.CTkFont(weight="bold")).pack(side="left", padx=15, pady=15)
        
        self.new_doctor_name = ctk.CTkEntry(add_frame, placeholder_text="Doctor Name", width=400)
        self.new_doctor_name.pack(side="left", padx=10)
        
        ctk.CTkButton(add_frame, text="➕ Add Doctor", fg_color="#2E8B57", hover_color="#245c3d",
                      command=self.add_doctor).pack(side="left", padx=10)

        # --- 3. Doctors List Section ---
        list_container = ctk.CTkFrame(self)
        list_container.pack(fill="both", expand=True, padx=20, pady=(0, 20))
        
        # Header for the list
        labels_frame = ctk.CTkFrame(list_container, fg_color="transparent", height=30)
        labels_frame.pack(fill="x", padx=10, pady=5)
        ctk.CTkLabel(labels_frame, text="Doctor Name", anchor="w").pack(side="left", fill="x", expand=True, padx=10)
        ctk.CTkLabel(labels_frame, text="Actions", width=150).pack(side="left", padx=10)

        # Scrollable area
        self.scroll_frame = ctk.CTkScrollableFrame(list_container)
        self.scroll_frame.pack(fill="both", expand=True, padx=10, pady=5)

        # Load data immediately
        self.refresh_list()

    def go_back(self):
        """Return to dashboard and reload data"""
        self.controller.doctors_data = self.controller.load_doctors_data() 
        self.controller.show_dashboard()

    def refresh_list(self):
        # Clear list
        for widget in self.scroll_frame.winfo_children():
            widget.destroy()
            
        # Load data from file directly
        doctors = []
        if os.path.exists(DOCTORS_FILE):
            with open(DOCTORS_FILE, "r", encoding="utf-8") as f:
                for line in f:
                    doc = line.strip()
                    if doc: doctors.append(doc)

        # Populate rows
        for doctor in doctors:
            row = ctk.CTkFrame(self.scroll_frame, fg_color="transparent")
            row.pack(fill="x", pady=2)
            
            # Name
            ctk.CTkLabel(row, text=doctor, anchor="w").pack(side="left", fill="x", expand=True, padx=10)
            
            # Buttons
            btn_frame = ctk.CTkFrame(row, fg_color="transparent", width=150)
            btn_frame.pack(side="left")
            
            # Edit Button
            ctk.CTkButton(btn_frame, text="✎", width=40, fg_color="#F39C12", hover_color="#D68910",
                          command=lambda d=doctor: self.edit_doctor_popup(d)).pack(side="left", padx=2)
            
            # Delete Button
            ctk.CTkButton(btn_frame, text="🗑", width=40, fg_color="#C0392B", hover_color="#922B21",
                          command=lambda d=doctor: self.delete_doctor(d)).pack(side="left", padx=2)

    def add_doctor(self):
        name = self.new_doctor_name.get().strip()
        if not name: return
            
        # Check for duplicate
        existing = []
        if os.path.exists(DOCTORS_FILE):
            with open(DOCTORS_FILE, "r", encoding="utf-8") as f:
                existing = [line.strip() for line in f]
        
        if name in existing:
            messagebox.showerror("Error", "Doctor already exists")
            return

        with open(DOCTORS_FILE, "a", encoding="utf-8") as f:
            f.write(f"{name}\n")
        
        self.new_doctor_name.delete(0, "end")
        self.refresh_list()

    def delete_doctor(self, name_to_delete):
        if not messagebox.askyesno("Confirm", f"Delete '{name_to_delete}'?"):
            return
            
        lines = []
        if os.path.exists(DOCTORS_FILE):
            with open(DOCTORS_FILE, "r", encoding="utf-8") as f:
                lines = f.readlines()
        
        with open(DOCTORS_FILE, "w", encoding="utf-8") as f:
            for line in lines:
                if line.strip() != name_to_delete:
                    f.write(line)
        
        self.refresh_list()

    def edit_doctor_popup(self, old_name):
        dialog = ctk.CTkToplevel(self)
        dialog.title("Edit Doctor")
        dialog.geometry("400x150")
        dialog.attributes('-topmost', True)
        
        ctk.CTkLabel(dialog, text="Doctor Name:").pack(pady=5)
        name_var = ctk.StringVar(value=old_name)
        ctk.CTkEntry(dialog, textvariable=name_var, width=300).pack(pady=5)
        
        def save_edit():
            new_n = name_var.get().strip()
            if not new_n: return
            
            # Read all lines
            lines = []
            if os.path.exists(DOCTORS_FILE):
                with open(DOCTORS_FILE, "r", encoding="utf-8") as f:
                    lines = [line.strip() for line in f.readlines()]
            
            # Replace
            with open(DOCTORS_FILE, "w", encoding="utf-8") as f:
                for line in lines:
                    if line == old_name:
                        f.write(f"{new_n}\n")
                    else:
                        f.write(f"{line}\n")
            
            dialog.destroy()
            self.refresh_list()

        ctk.CTkButton(dialog, text="Save Changes", command=save_edit).pack(pady=20)


class ManageAgentsFrame(ctk.CTkFrame):
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller
        
        # --- 1. Top Header ---
        header_frame = ctk.CTkFrame(self, fg_color="transparent")
        header_frame.pack(fill="x", padx=20, pady=10)
        
        # Back Button
        ctk.CTkButton(header_frame, text="⬅ Back to Dashboard", width=150, height=40,
                      fg_color="gray", hover_color="#555",
                      command=self.go_back).pack(side="left")
        
        # Title
        ctk.CTkLabel(header_frame, text="Manage Agents Database", 
                     font=ctk.CTkFont(size=24, weight="bold")).pack(side="left", padx=20)

        # --- 2. Add New Agent Section ---
        add_frame = ctk.CTkFrame(self)
        add_frame.pack(fill="x", padx=20, pady=10)
        
        ctk.CTkLabel(add_frame, text="Add New Agent:", font=ctk.CTkFont(weight="bold")).pack(side="left", padx=15, pady=15)
        
        self.new_agent_name = ctk.CTkEntry(add_frame, placeholder_text="Agent Name", width=400)
        self.new_agent_name.pack(side="left", padx=10)
        
        ctk.CTkButton(add_frame, text="➕ Add Agent", fg_color="#2E8B57", hover_color="#245c3d",
                      command=self.add_agent).pack(side="left", padx=10)

        # --- 3. Agents List Section ---
        list_container = ctk.CTkFrame(self)
        list_container.pack(fill="both", expand=True, padx=20, pady=(0, 20))
        
        # Header for the list
        labels_frame = ctk.CTkFrame(list_container, fg_color="transparent", height=30)
        labels_frame.pack(fill="x", padx=10, pady=5)
        ctk.CTkLabel(labels_frame, text="Agent Name", anchor="w").pack(side="left", fill="x", expand=True, padx=10)
        ctk.CTkLabel(labels_frame, text="Actions", width=150).pack(side="left", padx=10)

        # Scrollable area
        self.scroll_frame = ctk.CTkScrollableFrame(list_container)
        self.scroll_frame.pack(fill="both", expand=True, padx=10, pady=5)

        # Load data immediately
        self.refresh_list()

    def go_back(self):
        """Return to dashboard and reload data"""
        self.controller.agents_data = self.controller.load_agents_data() 
        self.controller.show_dashboard()

    def refresh_list(self):
        # Clear list
        for widget in self.scroll_frame.winfo_children():
            widget.destroy()
            
        # Load data from file directly
        agents = []
        if os.path.exists(AGENTS_FILE):
            with open(AGENTS_FILE, "r", encoding="utf-8") as f:
                for line in f:
                    agt = line.strip()
                    if agt: agents.append(agt)

        # Populate rows
        for agent in agents:
            row = ctk.CTkFrame(self.scroll_frame, fg_color="transparent")
            row.pack(fill="x", pady=2)
            
            # Name
            ctk.CTkLabel(row, text=agent, anchor="w").pack(side="left", fill="x", expand=True, padx=10)
            
            # Buttons
            btn_frame = ctk.CTkFrame(row, fg_color="transparent", width=150)
            btn_frame.pack(side="left")
            
            # Edit Button
            ctk.CTkButton(btn_frame, text="✎", width=40, fg_color="#F39C12", hover_color="#D68910",
                          command=lambda a=agent: self.edit_agent_popup(a)).pack(side="left", padx=2)
            
            # Delete Button
            ctk.CTkButton(btn_frame, text="🗑", width=40, fg_color="#C0392B", hover_color="#922B21",
                          command=lambda a=agent: self.delete_agent(a)).pack(side="left", padx=2)

    def add_agent(self):
        name = self.new_agent_name.get().strip()
        if not name: return
            
        # Check for duplicate
        existing = []
        if os.path.exists(AGENTS_FILE):
            with open(AGENTS_FILE, "r", encoding="utf-8") as f:
                existing = [line.strip() for line in f]
        
        if name in existing:
            messagebox.showerror("Error", "Agent already exists")
            return

        with open(AGENTS_FILE, "a", encoding="utf-8") as f:
            f.write(f"{name}\n")
        
        self.new_agent_name.delete(0, "end")
        self.refresh_list()

    def delete_agent(self, name_to_delete):
        if not messagebox.askyesno("Confirm", f"Delete '{name_to_delete}'?"):
            return
            
        lines = []
        if os.path.exists(AGENTS_FILE):
            with open(AGENTS_FILE, "r", encoding="utf-8") as f:
                lines = f.readlines()
        
        with open(AGENTS_FILE, "w", encoding="utf-8") as f:
            for line in lines:
                if line.strip() != name_to_delete:
                    f.write(line)
        
        self.refresh_list()

    def edit_agent_popup(self, old_name):
        dialog = ctk.CTkToplevel(self)
        dialog.title("Edit Agent")
        dialog.geometry("400x150")
        dialog.attributes('-topmost', True)
        
        ctk.CTkLabel(dialog, text="Agent Name:").pack(pady=5)
        name_var = ctk.StringVar(value=old_name)
        ctk.CTkEntry(dialog, textvariable=name_var, width=300).pack(pady=5)
        
        def save_edit():
            new_n = name_var.get().strip()
            if not new_n: return
            
            lines = []
            if os.path.exists(AGENTS_FILE):
                with open(AGENTS_FILE, "r", encoding="utf-8") as f:
                    lines = [line.strip() for line in f.readlines()]
            
            with open(AGENTS_FILE, "w", encoding="utf-8") as f:
                for line in lines:
                    if line == old_name:
                        f.write(f"{new_n}\n")
                    else:
                        f.write(f"{line}\n")
            
            dialog.destroy()
            self.refresh_list()

        ctk.CTkButton(dialog, text="Save Changes", command=save_edit).pack(pady=20)



class HistoryFrame(ctk.CTkFrame):
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller
        
        # --- 1. Top Header ---
        header_frame = ctk.CTkFrame(self, fg_color="transparent")
        header_frame.pack(fill="x", padx=20, pady=10)
        
        ctk.CTkButton(header_frame, text="⬅ Back to Dashboard", width=150, height=40,
                      fg_color="gray", hover_color="#555",
                      command=self.go_back).pack(side="left")
        
        ctk.CTkLabel(header_frame, text="Patient History & Records", 
                     font=ctk.CTkFont(size=24, weight="bold")).pack(side="left", padx=20)

        # --- 2. Controls Section (Search & Actions) ---
        controls_frame = ctk.CTkFrame(self)
        controls_frame.pack(fill="x", padx=20, pady=10)
        
        # Search
        ctk.CTkLabel(controls_frame, text="Search:").pack(side="left", padx=10)
        self.search_var = ctk.StringVar()
        search_entry = ctk.CTkEntry(controls_frame, textvariable=self.search_var, width=300,
                                    placeholder_text="Search by Name, Bill No, Date...")
        search_entry.pack(side="left", padx=10)
        search_entry.bind("<KeyRelease>", self.filter_history_data)
        
        # Action Buttons (Right Aligned)
        ctk.CTkButton(controls_frame, text="Download Excel 📥", command=self.download_excel_copy,
                      fg_color="#2E8B57", width=120).pack(side="right", padx=10)
                      
        ctk.CTkButton(controls_frame, text="Refresh 🔄", command=self.load_history_data,
                      width=100).pack(side="right", padx=10)

        # --- 3. Table Section ---
        table_frame = ctk.CTkFrame(self)
        table_frame.pack(fill="both", expand=True, padx=20, pady=(0, 20))
        
        # Create Scrollbars
        vsb = ttk.Scrollbar(table_frame, orient="vertical")
        hsb = ttk.Scrollbar(table_frame, orient="horizontal")

        # Create Treeview
        self.history_tree = ttk.Treeview(table_frame, yscrollcommand=vsb.set, xscrollcommand=hsb.set)
        
        # Grid layout for table and scrollbars
        self.history_tree.grid(row=0, column=0, sticky="nsew")
        vsb.grid(row=0, column=1, sticky="ns")
        hsb.grid(row=1, column=0, sticky="ew")
        
        # Configure table frame weights
        table_frame.grid_rowconfigure(0, weight=1)
        table_frame.grid_columnconfigure(0, weight=1)
        
        # Link scrollbars
        vsb.config(command=self.history_tree.yview)
        hsb.config(command=self.history_tree.xview)

        # Bind Context Menu
        self.history_tree.bind("<Button-3>", self.show_context_menu)

        # Load Data
        self.load_history_data()

    def go_back(self):
        self.controller.show_dashboard()

    def load_history_data(self):
        """Loads data from Excel with column sizing fixes for scrolling"""
        if not os.path.exists(EXCEL_FILE): return

        # Clear table
        for item in self.history_tree.get_children():
            self.history_tree.delete(item)

        try:
            df = pd.read_excel(EXCEL_FILE)
            if 'Due Payment Date' not in df.columns: df['Due Payment Date'] = ''
            
            # Set Columns
            self.history_tree["columns"] = list(df.columns)
            self.history_tree["show"] = "headings"
            
            for col in df.columns:
                self.history_tree.heading(col, text=col)
                
                # --- FIX: Custom Widths to enable Scrolling ---
                if col == "Tests":
                    # Make 'Tests' very wide and prevent it from shrinking
                    # stretch=False forces the scrollbar to appear if it goes off screen
                    self.history_tree.column(col, width=1000, minwidth=400, anchor="w", stretch=False)
                elif col == "Bill Number":
                    self.history_tree.column(col, width=100, minwidth=100, anchor="w", stretch=False)
                elif col in ["Age", "Gender"]:
                    self.history_tree.column(col, width=60, minwidth=60, anchor="w", stretch=True)
                else:
                    # Default for other columns
                    self.history_tree.column(col, width=120, minwidth=100, anchor="w", stretch=True)

            # Configure Tags for Highlights
            self.history_tree.tag_configure('due', background='#FFF3CD') # Light Yellow
            self.history_tree.tag_configure('match', background='#D1E7DD') # Light Green
            self.history_tree.tag_configure('match_due', background='#FFD700') # Gold

            # Insert Data
            for _, row in df.iterrows():
                values = []
                for col, val in row.items():
                    if pd.isna(val):
                        values.append("")
                    elif col == "Contact":
                        try:
                            values.append(str(int(float(val))))
                        except (ValueError, TypeError):
                            values.append(str(val))
                    else:
                        values.append(val)
                
                due_amt = float(row.get("Due Amount", 0)) if pd.notna(row.get("Due Amount")) else 0
                tag = ('due',) if due_amt > 0 else ()
                
                self.history_tree.insert("", "end", values=values, tags=tag)

        except Exception as e:
            messagebox.showerror("Error", f"Failed to load history: {e}")

            
    def filter_history_data(self, event=None):
        search_term = self.search_var.get().lower()
        
        # First reset all tags to base state
        for item in self.history_tree.get_children():
            current_tags = list(self.history_tree.item(item, 'tags'))
            is_due = 'due' in current_tags or 'match_due' in current_tags
            
            if not search_term:
                self.history_tree.item(item, tags=('due',) if is_due else ())
                self.history_tree.detach(item) 
                self.history_tree.move(item, '', 'end') 
                continue

            # Check for match
            values = self.history_tree.item(item, 'values')
            row_text = " ".join(str(v).lower() for v in values)
            
            if search_term in row_text:
                new_tag = 'match_due' if is_due else 'match'
                self.history_tree.item(item, tags=(new_tag,))
                self.history_tree.move(item, '', 'end') 
            else:
                self.history_tree.detach(item) 

    def show_context_menu(self, event):
        item = self.history_tree.identify_row(event.y)
        if not item: return
        
        self.history_tree.selection_set(item)
        values = self.history_tree.item(item, 'values')
        if not values: return
        
        bill_number = values[0]
        try:
            cols = self.history_tree['columns']
            if isinstance(cols, tuple): cols = list(cols)
            due_idx = cols.index('Due Amount')
            due_amount = float(values[due_idx]) if values[due_idx] else 0.0
        except: due_amount = 0.0

        menu = tk.Menu(self, tearoff=0)
        menu.add_command(label="Edit Bill", command=lambda: self.controller.edit_bill(bill_number))
        
        if due_amount > 0:
            menu.add_command(label="Pay Due", command=lambda: self.controller.pay_due(bill_number, due_amount))
            
        try: menu.tk_popup(event.x_root, event.y_root)
        finally: menu.grab_release()

    def download_excel_copy(self):
        self.controller.download_excel_copy()


class CreateReportFrame(ctk.CTkFrame):
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller
        self.reset_state() # Initialize variables

        # --- Top Bar ---
        top_bar = ctk.CTkFrame(self, fg_color="transparent")
        top_bar.pack(fill="x", pady=(6, 12), padx=6)
        
        ctk.CTkButton(top_bar, text=" ⬅️ Back", width=100,
                      command=lambda: controller.show_frame(ReportSelectionFrame)).pack(side="left")
        
        self.title_label = ctk.CTkLabel(top_bar, text="Create New Report Module", font=ctk.CTkFont(size=20, weight="bold"))
        self.title_label.pack(side="left", padx=20)

        # Upload Template Button
        self.upload_btn = ctk.CTkButton(top_bar, text="Upload Template (DOCX)", command=self.upload_template, fg_color="#E67E22", hover_color="#D35400")
        self.upload_btn.pack(side="right")

        # --- Report Name Input ---
        name_frame = ctk.CTkFrame(self, fg_color="transparent")
        name_frame.pack(fill="x", padx=20, pady=5)
        ctk.CTkLabel(name_frame, text="Report Name:").pack(side="left", padx=5)
        self.report_name_entry = ctk.CTkEntry(name_frame, width=300, placeholder_text="e.g., Thyroid Profile")
        self.report_name_entry.pack(side="left", padx=5)

        # --- Middle Section: Field Definition ---
        input_frame = ctk.CTkFrame(self)
        input_frame.pack(fill="x", padx=20, pady=10)

        self.field_name_entry = ctk.CTkEntry(input_frame, width=200, placeholder_text="Enter Field Name")
        self.field_name_entry.pack(side="left", padx=10, pady=10)

        self.field_code_entry = ctk.CTkEntry(input_frame, width=150, placeholder_text="Enter Field Code (e.g. lbt)")
        self.field_code_entry.pack(side="left", padx=10, pady=10)

        ctk.CTkButton(input_frame, text="Add Field", width=100, command=self.add_field).pack(side="left", padx=5)
        ctk.CTkButton(input_frame, text="Add Logic", width=100, fg_color="#5D6D7E", command=self.add_logic_field).pack(side="left", padx=5)

        # --- Bottom Section: Field List ---
        ctk.CTkLabel(self, text="Added Fields:", font=ctk.CTkFont(weight="bold")).pack(anchor="w", padx=20)
        
        self.field_list_frame = ctk.CTkScrollableFrame(self, height=250)
        self.field_list_frame.pack(fill="x", padx=20, pady=5)

        # --- Save Button ---
        self.save_btn = ctk.CTkButton(self, text="Create Report Module", height=50, font=ctk.CTkFont(size=16, weight="bold"),
                      command=self.save_report_configuration)
        self.save_btn.pack(pady=20, padx=20, fill="x")

    def reset_state(self):
        """Resets the frame to 'Create' mode"""
        self.fields = []
        self.template_path = None
        self.edit_mode = False
        self.original_name = None
        if hasattr(self, 'report_name_entry'):
            self.report_name_entry.delete(0, "end")
            self.field_name_entry.delete(0, "end")
            self.field_code_entry.delete(0, "end")
            self.upload_btn.configure(text="Upload Template (DOCX)", fg_color="#E67E22")
            self.title_label.configure(text="Create New Report Module")
            self.save_btn.configure(text="Create Report Module")
            for widget in self.field_list_frame.winfo_children():
                widget.destroy()

    def load_for_editing(self, report_name, config):
        """Loads an existing report into the frame for editing"""
        self.reset_state()
        self.edit_mode = True
        self.original_name = report_name
        self.fields = config['fields']
        self.template_path = config['template']

        # Update UI
        self.title_label.configure(text=f"Edit Report: {report_name}")
        self.save_btn.configure(text="Save Changes")
        self.report_name_entry.insert(0, report_name)
        self.upload_btn.configure(text=f"✅ {self.template_path}", fg_color="green")
        
        self.refresh_field_list()

    def upload_template(self):
        filename = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")])
        if filename:
            if not os.path.exists(REPORT_TEMPLATES_FOLDER):
                os.makedirs(REPORT_TEMPLATES_FOLDER)
            
            basename = os.path.basename(filename)
            destination = os.path.join(REPORT_TEMPLATES_FOLDER, basename)
            try:
                shutil.copy(filename, destination)
                self.template_path = basename
                self.upload_btn.configure(text=f"✅ {basename}", fg_color="green")
            except Exception as e:
                messagebox.showerror("Error", f"Could not save template: {e}")

    def add_field(self):
        name = self.field_name_entry.get().strip()
        code = self.field_code_entry.get().strip()
        if not name or not code: return
        
        self.fields.append({"type": "input", "name": name, "code": code, "logic": None})
        self._clear_inputs_and_refresh()

    def add_logic_field(self):
        name = self.field_name_entry.get().strip()
        code = self.field_code_entry.get().strip()
        if not name or not code: return

        current_count = len(self.fields)
        logic = simpledialog.askstring("Input Logic", f"Enter logic (e.g., (2-1)*20).\nAvailable IDs: 1 to {current_count}")
        
        if logic:
            self.fields.append({"type": "logic", "name": name, "code": code, "logic": logic})
            self._clear_inputs_and_refresh()

    def _clear_inputs_and_refresh(self):
        self.field_name_entry.delete(0, "end")
        self.field_code_entry.delete(0, "end")
        self.field_name_entry.focus()
        self.refresh_field_list()

    def edit_logic(self, index):
        """Opens dialog to edit the logic of a specific field"""
        field = self.fields[index]
        new_logic = simpledialog.askstring("Edit Logic", f"Update logic for '{field['name']}':", initialvalue=field['logic'])
        if new_logic:
            field['logic'] = new_logic
            self.refresh_field_list()

    def delete_field(self, index):
        if 0 <= index < len(self.fields):
            del self.fields[index]
            self.refresh_field_list()

    def refresh_field_list(self):
        for widget in self.field_list_frame.winfo_children():
            widget.destroy()

        for index, field_data in enumerate(self.fields):
            serial_id = index + 1
            field_data['id'] = serial_id 
            
            row_frame = ctk.CTkFrame(self.field_list_frame, fg_color="transparent")
            row_frame.pack(fill="x", pady=2, padx=5)

            text = f"{serial_id}. {field_data['name']} ({{{field_data['code']}}})"
            if field_data['type'] == "logic":
                text += f" - Logic: [ {field_data['logic']} ]"
            else:
                text += " - No logic"
                
            lbl = ctk.CTkLabel(row_frame, text=text, anchor="w")
            lbl.pack(side="left", fill="x", expand=True)

            # --- Delete Button (❌) ---
            del_btn = ctk.CTkButton(row_frame, text="❌", width=30, height=25, 
                                    fg_color="#C0392B", hover_color="#922B21",
                                    command=lambda i=index: self.delete_field(i))
            del_btn.pack(side="right", padx=5)

            # --- Edit Logic Button (✏️) - Only for Logic fields ---
            if field_data['type'] == "logic":
                edit_btn = ctk.CTkButton(row_frame, text="✏️", width=30, height=25,
                                         fg_color="#F39C12", hover_color="#D68910",
                                         command=lambda i=index: self.edit_logic(i))
                edit_btn.pack(side="right", padx=5)

    def save_report_configuration(self):
        report_name = self.report_name_entry.get().strip()
        if not report_name or not self.template_path or not self.fields:
            messagebox.showerror("Error", "Missing information (Name, Template, or Fields)")
            return

        new_report = {"name": report_name, "template": self.template_path, "fields": self.fields}

        data = {}
        if os.path.exists(CUSTOM_REPORTS_FILE):
            with open(CUSTOM_REPORTS_FILE, 'r') as f:
                try: data = json.load(f)
                except: data = {}
        
        # If Editing and name changed, remove the old key
        if self.edit_mode and self.original_name != report_name:
            if self.original_name in data:
                del data[self.original_name]
        
        data[report_name] = new_report
        
        with open(CUSTOM_REPORTS_FILE, 'w') as f:
            json.dump(data, f, indent=4)

        messagebox.showinfo("Success", "Report Saved Successfully!")
        self.reset_state()
        self.controller.frames[ReportSelectionFrame].refresh_reports()
        self.controller.show_frame(ReportSelectionFrame)


class DynamicReportFrame(ctk.CTkFrame):
    def __init__(self, parent, controller, report_config):
        super().__init__(parent)
        self.controller = controller
        self.config = report_config
        self.entries = {} 
        
        # --- Top Bar ---
        top_bar = ctk.CTkFrame(self, fg_color="transparent")
        top_bar.pack(fill="x", pady=(6, 12), padx=6)
        
        # Back Button
        ctk.CTkButton(top_bar, text=" ⬅️ Back", width=100,
                      command=lambda: controller.show_frame(ReportSelectionFrame)).pack(side="left", padx=6, pady=6)
        
        # Title
        ctk.CTkLabel(top_bar, text=f"{self.config['name']} Entry", font=ctk.CTkFont(size=18, weight="bold")).pack(side="left", padx=8)

        # --- Right Side Action Buttons (Edit / Delete) ---
        # Delete Button
        ctk.CTkButton(top_bar, text="Delete Report 🗑️", width=120, fg_color="#C0392B", hover_color="#922B21",
                      command=self.delete_report).pack(side="right", padx=5)
        
        # Edit Button
        ctk.CTkButton(top_bar, text="Edit Report ✏️", width=120, fg_color="#2980B9", hover_color="#2471A3",
                      command=self.edit_report).pack(side="right", padx=5)

        # --- Content Area (Same as before) ---
        content = ctk.CTkFrame(self)
        content.pack(expand=True, fill="both", padx=6, pady=6)
        
        left = ctk.CTkFrame(content)
        left.pack(side="left", expand=True, fill="both", padx=(0,8), pady=8)
        self.patient_entries = {}
        self.add_patient_details(left)

        right = ctk.CTkScrollableFrame(content)
        right.pack(side="right", expand=True, fill="both", padx=(8,0), pady=8)
        
        for field in self.config['fields']:
            row_frame = ctk.CTkFrame(right, fg_color="transparent")
            row_frame.pack(fill="x", pady=2)
            ctk.CTkLabel(row_frame, text=field['name'], width=150, anchor="w").pack(side="left", padx=5)
            entry = ctk.CTkEntry(row_frame, width=200)
            entry.pack(side="left", padx=5)
            if field['type'] == 'logic':
                entry.configure(placeholder_text="Auto-calculated")
            self.entries[field['id']] = entry

        ctk.CTkButton(self, text="Save Report", height=40, command=self.generate_report).pack(pady=12, padx=6)

    def delete_report(self):
        """Deletes the report config, template, and generated files."""
        if not messagebox.askyesno("Confirm Delete", 
                                   f"Are you sure you want to delete '{self.config['name']}'?\n\n"
                                   "This will delete:\n1. The Report Configuration\n2. The Template File\n3. The Generated Output Folder"):
            return

        report_name = self.config['name']
        template_name = self.config['template']
        folder_name = report_name.replace(" ", "")

        # 1. Delete Template File
        template_path = os.path.join(REPORT_TEMPLATES_FOLDER, template_name)
        if os.path.exists(template_path):
            try: os.remove(template_path)
            except Exception as e: print(f"Error deleting template: {e}")

        # 2. Delete Generated Output Folder
        output_path = os.path.join(OUTPUT_FOLDER, folder_name)
        if os.path.exists(output_path):
            try: shutil.rmtree(output_path)
            except Exception as e: print(f"Error deleting output folder: {e}")

        # 3. Delete from JSON
        if os.path.exists(CUSTOM_REPORTS_FILE):
            try:
                with open(CUSTOM_REPORTS_FILE, 'r') as f:
                    data = json.load(f)
                
                if report_name in data:
                    del data[report_name]
                
                with open(CUSTOM_REPORTS_FILE, 'w') as f:
                    json.dump(data, f, indent=4)
            except Exception as e:
                messagebox.showerror("Error", f"Failed to update database: {e}")
                return

        messagebox.showinfo("Deleted", f"Report '{report_name}' has been deleted.")
        
        # Refresh and Navigate Back
        self.controller.frames[ReportSelectionFrame].refresh_reports()
        self.controller.show_frame(ReportSelectionFrame)

    def edit_report(self):
        """Opens the Create Page in Edit Mode with current config"""
        create_frame = self.controller.frames[CreateReportFrame]
        create_frame.load_for_editing(self.config['name'], self.config)
        self.controller.show_frame(CreateReportFrame)

    # ... (Keep existing methods: add_patient_details, populate_patient_data, etc. exactly as they were) ...
    def add_patient_details(self, parent):
        labels = ["Patient Name", "Age", "Sex", "Ref. by Dr.", "Patient ID", "Collection Date", "Reporting Date"]
        keys = ["name", "age", "sex", "doctor", "patient_id", "collection_date", "report_date"]
        for i, (lbl, key) in enumerate(zip(labels, keys)):
            ctk.CTkLabel(parent, text=lbl, anchor="w").grid(row=i, column=0, padx=5, pady=4, sticky="w")
            e = ctk.CTkEntry(parent, width=220)
            e.grid(row=i, column=1, padx=5, pady=4)
            self.patient_entries[key] = e
            if "date" in key: e.bind("<FocusOut>", self._autocomplete_date)

    def populate_patient_data(self, data):
        mapping = {'name': 'Patient Name', 'age': 'Age', 'sex': 'Gender', 'doctor': 'Ref By', 'patient_id': 'Bill Number'}
        for form_key, data_key in mapping.items():
            if form_key in self.patient_entries:
                self.patient_entries[form_key].delete(0, "end")
                val = data.get(data_key, '')
                if form_key == 'age': val = str(val).split('.')[0]
                self.patient_entries[form_key].insert(0, val)

    def clear_patient_fields(self):
        for key, entry in self.patient_entries.items():
            if "date" not in key: entry.delete(0, "end")

    def _autocomplete_date(self, event):
        widget = event.widget
        day_str = widget.get().strip()
        if day_str.isdigit() and len(day_str) <= 2:
            try:
                day = int(day_str)
                now = datetime.now()
                widget.delete(0, "end")
                widget.insert(0, f"{day:02d}.{now.month:02d}.{now.year}")
            except: pass

    def generate_report(self):
        context = {k: v.get() for k, v in self.patient_entries.items()}
        field_values = {}
        
        # Gather inputs
        for field in self.config['fields']:
            fid = field['id']
            val = self.entries[fid].get()
            try: field_values[fid] = float(val) if val else 0.0
            except: field_values[fid] = 0.0
            context[field['code']] = val

        # Process Logic
        for field in self.config['fields']:
            if field['type'] == 'logic' and field['logic']:
                def replace_id(match):
                    return str(field_values.get(int(match.group(0)), 0.0))
                try:
                    calc_str = re.sub(r'\b\d+\b', replace_id, field['logic'])
                    result = eval(calc_str)
                    formatted = f"{result:.2f}"
                    context[field['code']] = formatted
                    self.entries[field['id']].delete(0, "end")
                    self.entries[field['id']].insert(0, formatted)
                    field_values[field['id']] = result
                except Exception as e:
                    print(f"Logic Error: {e}")

        # Generate Doc
        try:
            full_template_path = os.path.join(REPORT_TEMPLATES_FOLDER, self.config['template'])
            doc = DocxTemplate(full_template_path)
            doc.render(context)
            
            folder_name = self.config['name'].replace(" ", "")
            target_dir = os.path.join(OUTPUT_FOLDER, folder_name)
            os.makedirs(target_dir, exist_ok=True)
            
            p_id = context.get("patient_id", "NoID").strip().replace(' ', '_')
            out_path = os.path.join(target_dir, f"{folder_name}_{p_id}.docx")
            
            counter = 1
            while os.path.exists(out_path):
                out_path = os.path.join(target_dir, f"{folder_name}_{p_id}({counter}).docx")
                counter += 1
            
            doc.save(out_path)
            messagebox.showinfo("Success", f"Report saved:\n{out_path}")
            os.startfile(out_path)
        except Exception as e:
            messagebox.showerror("Error", f"Failed to generate: {e}")

class ReportModuleWindow(ctk.CTkToplevel):
    def __init__(self, master=None):
        super().__init__(master)
        self.title("Report Module")
        self.after(100, lambda: self.state("zoomed"))
        
        self.current_patient_data = None

        container = ctk.CTkFrame(self)
        container.pack(fill="both", expand=True)
        container.grid_rowconfigure(0, weight=1)
        container.grid_columnconfigure(0, weight=1)

        self.frames = {}
        # Add CreateReportFrame to the list
        for F in (ReportSelectionFrame, CreateReportFrame, CBCFormFrame, BilirubinFormFrame, ElectrolytesFormFrame, LipidProfileFormFrame, LFTFormFrame,RFTFormFrame, 
                  SugerPPBSFormFrame, SugerRBSFormFrame, UricAcidFormFrame, RenalProfileFormFrame, ComprehensiveProfileFormFrame, UreaCreatinineFormFrame, 
                  BSGTFormFrame, SugerFBSFormFrame, SugerFBS_PPBSFormFrame, SodiumPotassiumFormFrame, RBS_Urea_CreatinineFormFrame):
            frame = F(container, self)
            self.frames[F] = frame
            frame.grid(row=0, column=0, sticky="nsew")

        self.show_frame(ReportSelectionFrame)

    def show_frame(self, cont):
        frame = self.frames[cont]
        # Refresh dynamic list if going back to selection
        if cont == ReportSelectionFrame:
            frame.refresh_reports()
            
        if self.current_patient_data is not None and hasattr(frame, 'populate_patient_data'):
            frame.populate_patient_data(self.current_patient_data)
        else:
            if hasattr(frame, 'clear_patient_fields'):
                frame.clear_patient_fields()
        frame.tkraise()

    def show_dynamic_frame(self, report_config):
        """Creates and shows a DynamicReportFrame on the fly based on JSON config"""
        # We don't store this in self.frames permanently because it changes based on selection
        # We create it inside the container used by other frames
        
        # Access the container from one of the existing frames (hacky but works)
        container = list(self.frames.values())[0].master 
        
        # Destroy previous dynamic frame if exists
        if hasattr(self, 'dynamic_frame'):
            self.dynamic_frame.destroy()
            
        self.dynamic_frame = DynamicReportFrame(container, self, report_config)
        self.dynamic_frame.grid(row=0, column=0, sticky="nsew")
        
        if self.current_patient_data is not None:
            self.dynamic_frame.populate_patient_data(self.current_patient_data)
            
        self.dynamic_frame.tkraise()


class ReportSelectionFrame(ctk.CTkFrame):
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller

        # --- Main Layout Configuration ---
        self.grid_columnconfigure(0, weight=1) 
        self.grid_columnconfigure(1, weight=2) 
        self.grid_rowconfigure(0, weight=1)

        # --- 1. Left Frame (Report Buttons) ---
        left_frame = ctk.CTkFrame(self)
        left_frame.grid(row=0, column=0, padx=(10, 5), pady=10, sticky="nsew")
        left_frame.grid_rowconfigure(3, weight=1) 
        left_frame.grid_columnconfigure(0, weight=1)

        self.left_header_label = ctk.CTkLabel(left_frame, text="Select a Report", font=ctk.CTkFont(size=22, weight="bold"))
        self.left_header_label.grid(row=0, column=0, pady=20, padx=20)

        # Create Report Button
        create_btn = ctk.CTkButton(left_frame, text="➕ Create a Report", fg_color="#8E44AD", hover_color="#732D91",
                                   command=lambda: self.controller.show_frame(CreateReportFrame))
        create_btn.grid(row=1, column=0, pady=5, padx=20, sticky="ew")
        
        self.clear_search_button = ctk.CTkButton(left_frame, text="⬅️ Show All Reports", command=self.reset_search)
        
        self.report_buttons_frame = ctk.CTkScrollableFrame(left_frame)
        self.report_buttons_frame.grid(row=3, column=0, sticky="nsew", padx=10, pady=(0, 10))
        self.report_buttons_frame.grid_columnconfigure(0, weight=1)

        # --- 2. Right Frame (Patient Search) ---
        right_frame = ctk.CTkFrame(self)
        right_frame.grid(row=0, column=1, padx=(5, 10), pady=10, sticky="nsew")
        right_frame.grid_rowconfigure(2, weight=1) 
        right_frame.grid_columnconfigure(0, weight=1)

        ctk.CTkLabel(right_frame, text="Patient & Report Search", font=ctk.CTkFont(size=22, weight="bold")).grid(row=0, column=0, pady=20, padx=20)
        
        search_bar_frame = ctk.CTkFrame(right_frame, fg_color="transparent")
        search_bar_frame.grid(row=1, column=0, sticky="ew", padx=10, pady=(0, 10))
        search_bar_frame.grid_columnconfigure(0, weight=1)

        self.search_var = ctk.StringVar()
        search_entry = ctk.CTkEntry(search_bar_frame, placeholder_text="Type a test name to filter, or a Bill Number to search...", textvariable=self.search_var)
        search_entry.grid(row=0, column=0, sticky="ew", padx=(0, 10))
        
        search_entry.bind("<KeyRelease>", self.filter_reports_by_name)
        search_entry.bind("<Return>", self.search_patient_by_id)      

        button_panel = ctk.CTkFrame(search_bar_frame, fg_color="transparent")
        button_panel.grid(row=0, column=1, sticky="e")
        
        search_button = ctk.CTkButton(button_panel, text="Search 🔎", width=100, command=self.search_patient_by_id)
        search_button.pack(side="left", padx=(0, 10))
        
        reset_button = ctk.CTkButton(button_panel, text="Reset", width=80, command=self.reset_search, fg_color="#D2691E", hover_color="#8B4513")
        reset_button.pack(side="left", padx=(0, 10))

        saved_reports_button = ctk.CTkButton(button_panel, text="Saved Reports 📁", width=140, command=self.open_saved_reports_folder, fg_color="#1F6AA5", hover_color="#144569")
        saved_reports_button.pack(side="left")

        self.results_textbox = ctk.CTkTextbox(right_frame, wrap="word", state="disabled", font=("Segoe UI", 14))
        self.results_textbox.grid(row=2, column=0, sticky="nsew", padx=10, pady=(0, 10))
        self.results_textbox.tag_config("header", underline=True, spacing1=12, spacing3=5)
        
        # Initial Population
        self.refresh_reports()
    
    def open_saved_reports_folder(self):
        folder_path = os.path.abspath(OUTPUT_FOLDER)
        if not os.path.exists(folder_path):
            os.makedirs(folder_path)
            messagebox.showinfo("Folder Created", f"The folder '{folder_path}' was created as it did not exist.", parent=self)
        try:
            os.startfile(folder_path)
        except Exception as e:
            messagebox.showerror("Error", f"Could not open the reports folder.\nError: {e}", parent=self)

    def reset_search(self, event=None):
        """Clears search results, clears stored patient data, and restores the default view."""
        self.search_var.set("")
        self.results_textbox.configure(state="normal")
        self.results_textbox.delete("1.0", "end")
        self.results_textbox.configure(state="disabled")
        
        self.controller.current_patient_data = None
        
        self.refresh_reports()

    def _clear_report_buttons(self):
        for widget in self.report_buttons_frame.winfo_children():
            widget.destroy()

    def _populate_all_report_buttons(self):
        self.refresh_reports()

    def refresh_reports(self):
        """Loads ONLY Custom Reports into the UI."""
        self._clear_report_buttons()
        self.left_header_label.configure(text="Select a Report")
        self.clear_search_button.grid_forget()

        row_index = 0

        # Load Custom Reports from JSON (No Static Reports)
        if os.path.exists(CUSTOM_REPORTS_FILE):
            try:
                with open(CUSTOM_REPORTS_FILE, 'r') as f:
                    custom_data = json.load(f)
                
                if custom_data:
                    for r_name, r_config in custom_data.items():
                        # Pass the CONFIG to the dynamic frame
                        ctk.CTkButton(self.report_buttons_frame, text=f"📄 {r_name}", width=240, height=44, fg_color="#2980B9", hover_color="#2471A3",
                                      command=lambda cfg=r_config: self.controller.show_dynamic_frame(cfg)).grid(row=row_index, column=0, pady=10)
                        row_index += 1
                else:
                    ctk.CTkLabel(self.report_buttons_frame, text="No custom reports created yet.", font=ctk.CTkFont(slant="italic")).grid(row=0, column=0, pady=20)
            except Exception as e:
                print(f"Error loading custom reports: {e}")
                ctk.CTkLabel(self.report_buttons_frame, text="Error loading reports.", text_color="red").grid(row=0, column=0, pady=20)
        else:
            ctk.CTkLabel(self.report_buttons_frame, text="No custom reports found.\nClick 'Create a Report' to start.", 
                         font=ctk.CTkFont(slant="italic")).grid(row=0, column=0, pady=20)

    def filter_reports_by_name(self, event=None):
        """Filters ONLY custom reports."""
        if event and event.keysym == 'Return':
            return
        
        search_term = self.search_var.get().strip().lower()
        if not search_term:
            self.refresh_reports()
            return
            
        self._clear_report_buttons()
        self.left_header_label.configure(text="Matching Reports")
        self.clear_search_button.grid(row=2, column=0, pady=(0,10), padx=20, sticky="ew")
        
        row_index = 0
        found_any = False

        # Filter Custom Reports Only
        if os.path.exists(CUSTOM_REPORTS_FILE):
            try:
                with open(CUSTOM_REPORTS_FILE, 'r') as f:
                    custom = json.load(f)
                for name, cfg in custom.items():
                    if search_term in name.lower():
                        ctk.CTkButton(self.report_buttons_frame, text=f"📄 {name}", width=240, height=44, fg_color="#2980B9",
                                      command=lambda c=cfg: self.controller.show_dynamic_frame(c)).grid(row=row_index, column=0, pady=10)
                        row_index += 1
                        found_any = True
            except: pass

        if not found_any:
            ctk.CTkLabel(self.report_buttons_frame, text="No matching reports found.", font=ctk.CTkFont(slant="italic")).grid(row=0, column=0, pady=20)

    def search_patient_by_id(self, event=None):
        """Fetches patient data and stores it in the controller."""
        bill_number_to_find = self.search_var.get().strip()
        if not bill_number_to_find:
            self.controller.current_patient_data = None
            messagebox.showwarning("Input Error", "Please enter a Bill Number to search.", parent=self)
            return

        self.results_textbox.configure(state="normal")
        self.results_textbox.delete("1.0", "end")

        try:
            if not os.path.exists(EXCEL_FILE):
                self.controller.current_patient_data = None
                self.results_textbox.insert("end", f"❌ Error: The data file '{EXCEL_FILE}' was not found.")
                self.refresh_reports()
                return

            df = pd.read_excel(EXCEL_FILE)
            df['Bill Number'] = df['Bill Number'].astype(str)
            patient_data = df[df['Bill Number'].str.lower() == bill_number_to_find.lower()]

            if patient_data.empty:
                self.controller.current_patient_data = None
                self.results_textbox.insert("end", f"🤷 No record found for Bill Number: {bill_number_to_find}")
            else:
                record = patient_data.iloc[0]
                self.controller.current_patient_data = record.to_dict()
                
                # --- Update UI Textbox ---
                header_text = f"Name: {record.get('Patient Name', 'N/A')}  |  Age: {record.get('Age', 'N/A')}  |  Gender: {record.get('Gender', 'N/A')}\n"
                self.results_textbox.insert("end", header_text)
                self.results_textbox.insert("end", "—" * 50 + "\n")
                self.results_textbox.insert("end", "👤 PATIENT DETAILS\n", "header")
                patient_details = (f"  Bill Number: {record.get('Bill Number', 'N/A')}\n" +
                                   f"  Date: {record.get('Date', 'N/A')}\n" +
                                   f"  Referred By: {record.get('Ref By', 'N/A')}\n" +
                                   f"  Agent: {record.get('Agent', 'N/A')}\n" +
                                   f"  Address: {record.get('Address', 'N/A')}\n" +
                                   f"  Contact: {str(record.get('Contact', 'N/A')).split('.')[0]}\n\n")
                self.results_textbox.insert("end", patient_details)
                self.results_textbox.insert("end", "🧪 TEST DETAILS\n", "header")
                tests_str = record.get('Tests', '')
                formatted_tests = "  " + tests_str.replace("), ", ")\n  ")
                self.results_textbox.insert("end", f"{formatted_tests}\n\n")
                self.results_textbox.insert("end", "💰 AMOUNT DETAILS\n", "header")
                amount_details = (f"  Total Amount: ₹{record.get('Total Amount', 0):.2f}\n" +
                                  f"  Discount: ₹{record.get('Discount', 0):.2f}\n" +
                                  f"  Advanced Payment: ₹{record.get('Advanced Payment', 0):.2f}\n" +
                                  f"  Due Amount: ₹{record.get('Due Amount', 0):.2f}\n" +
                                  f"  Due Paid On: {record.get('Due Payment Date') if pd.notna(record.get('Due Payment Date')) else 'N/A'}\n")
                self.results_textbox.insert("end", amount_details)

                # --- Update Left Panel (Custom Reports Only) ---
                if tests_str:
                    patient_tests = [t.split(' (')[0].strip() for t in tests_str.split(',')]
                    self._display_filtered_report_buttons(patient_tests)
                else:
                    self._display_filtered_report_buttons([])

        except Exception as e:
            self.controller.current_patient_data = None
            self.results_textbox.insert("end", f"An error occurred: {e}")
            self.refresh_reports()
        finally:
            self.results_textbox.configure(state="disabled")

    def _display_filtered_report_buttons(self, patient_tests):
        """Updates the button list to show only CUSTOM reports relevant to the patient's tests."""
        self._clear_report_buttons()
        self.left_header_label.configure(text="Patient's Reports")
        self.clear_search_button.grid(row=2, column=0, pady=(0,10), padx=20, sticky="ew")
        
        found_reports = False
        row_index = 0
        added_reports = set()

        # Check Custom Reports Only
        if os.path.exists(CUSTOM_REPORTS_FILE):
            try:
                with open(CUSTOM_REPORTS_FILE, 'r') as f:
                    custom = json.load(f)
                
                for test_name in patient_tests:
                    for name, cfg in custom.items():
                        if name.lower() in test_name.lower() or test_name.lower() in name.lower():
                            if name not in added_reports:
                                ctk.CTkButton(self.report_buttons_frame, text=f"📄 {name}", width=240, height=44, 
                                              fg_color="#2980B9", hover_color="#2471A3",
                                              command=lambda c=cfg: self.controller.show_dynamic_frame(c)).grid(row=row_index, column=0, pady=10)
                                row_index += 1
                                found_reports = True
                                added_reports.add(name)
            except Exception as e:
                print(f"Error filtering custom reports: {e}")

        if not found_reports:
            ctk.CTkLabel(self.report_buttons_frame, text="No specific report forms\nfound for these tests.",
                         font=ctk.CTkFont(size=14, slant="italic")).grid(row=0, column=0, pady=20)
# --- 3. The CBC Report Form Page ---

class CBCFormFrame(ctk.CTkFrame):
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller
        
        # ... (The rest of the __init__ method that creates the form is unchanged) ...
        self.patient_entries = {}
        self.test_entries = {}

        top_bar = ctk.CTkFrame(self, fg_color="transparent")
        top_bar.pack(fill="x", pady=(6, 12), padx=6)
        
        ctk.CTkButton(top_bar, text=" ⬅️ Back to List", width=120,
                      command=lambda: controller.show_frame(ReportSelectionFrame)).pack(side="left", padx=6, pady=6)
        ctk.CTkLabel(top_bar, text="Complete Blood Count (CBC) Entry", font=ctk.CTkFont(size=18, weight="bold")).pack(side="left", padx=8)

        content = ctk.CTkFrame(self)
        content.pack(expand=True, fill="both", padx=6, pady=6)
        left = ctk.CTkFrame(content)
        left.pack(side="left", expand=True, fill="both", padx=(0,8), pady=8)
        right = ctk.CTkScrollableFrame(content, width=460, height=420)
        right.pack(side="right", expand=True, fill="both", padx=(8,0), pady=8)

        self.add_patient_field(left, 0, "Patient Name", "name")
        self.add_patient_field(left, 1, "Age", "age")
        self.add_patient_field(left, 2, "Sex", "sex")
        self.add_patient_field(left, 3, "Ref. by Dr.", "doctor")
        self.add_patient_field(left, 4, "Patient ID", "patient_id")
        self.add_patient_field(left, 5, "Collection Date", "collection_date")
        self.add_patient_field(left, 6, "Reporting Date", "report_date")
        
        self.add_test_field(right, 0, "HEAMOGLOBIN", "hemoglobin")
        self.add_test_field(right, 1, "TOTAL WBC COUNT", "wbc")
        self.add_test_field(right, 2, "Neutrophils", "neutrophils")
        self.add_test_field(right, 3, "Lymphocytes", "lymphocytes")
        self.add_test_field(right, 4, "Monocytes", "monocytes")
        self.add_test_field(right, 5, "Eosinophils", "eosinophils")
        self.add_test_field(right, 6, "Basophils", "basophils")
        self.add_test_field(right, 7, "PLATELET COUNT", "platelets")
        self.add_test_field(right, 8, "RBC COUNT", "rbc")
        self.add_test_field(right, 9, "HEAMATOCRIT (PCV)", "pcv")
        self.add_test_field(right, 10, "RDW - CV", "rdwcv")
        self.add_test_field(right, 11, "RDW C SD", "rdwsd")
        self.add_test_field(right, 12, "MPV", "mpv")
        self.add_test_field(right, 13, "PDW", "pdw")
        self.add_test_field(right, 14, "Erythrocyte Sedimentation Rate", "esr")

        ctk.CTkButton(self, text="Save Report", height=40, command=self.generate_report).pack(pady=12, padx=6)

    def add_patient_field(self, parent, row, label, key):
        ctk.CTkLabel(parent, text=label, anchor="w").grid(row=row, column=0, padx=5, pady=4, sticky="w")
        e = ctk.CTkEntry(parent, width=220)
        e.grid(row=row, column=1, padx=5, pady=4)
        self.patient_entries[key] = e
        # This line activates the autocomplete function for date fields
        if "date" in key.lower():
            e.bind("<FocusOut>", self._autocomplete_date)

    def add_test_field(self, parent, row, label, key):
        ctk.CTkLabel(parent, text=label, anchor="w").grid(row=row, column=0, padx=5, pady=4, sticky="w")
        e = ctk.CTkEntry(parent, width=220)
        e.grid(row=row, column=1, padx=5, pady=4)
        self.test_entries[key] = e

    def clear_all_fields(self):
        for e in list(self.patient_entries.values()) + list(self.test_entries.values()):
            e.delete(0, "end")

    def generate_report(self):
        # ... (context dictionary and calculations are the same) ...
        def get_float(key):
            try:
                return float(self.test_entries[key].get())
            except (ValueError, KeyError):
                return 0.0

        context = {key: entry.get() for key, entry in self.patient_entries.items()}
        context.update({key: entry.get() for key, entry in self.test_entries.items()})
        
        hemo, wbc, neut = get_float("hemoglobin"), get_float("wbc"), get_float("neutrophils")
        lymph, eos, rbc = get_float("lymphocytes"), get_float("eosinophils"), get_float("rbc")
        pcv, platelet, mpv = get_float("pcv"), get_float("platelets"), get_float("mpv")

        context['abs_neut'] = f"{(wbc * neut / 100) / 1000:.2f}" if wbc and neut else "0.00"
        context['abs_lymph'] = f"{(wbc * lymph / 100) / 1000:.2f}" if wbc and lymph else "0.00"
        context['abs_eos'] = f"{(wbc * eos / 100) / 1000:.2f}" if wbc and eos else "0.00"
        context['mcv'] = f"{(pcv / rbc * 10):.2f}" if pcv and rbc else "0.00"
        context['mch'] = f"{(hemo / rbc * 10):.2f}" if hemo and rbc else "0.00"
        context['mchc'] = f"{(hemo / pcv * 100):.2f}" if hemo and pcv else "0.00"
        context['pct'] = f"{(platelet * mpv / 10000) * 100:.2f}" if platelet and mpv else "0.00"
        
        try:
            # <<< THIS IS THE MAIN CHANGE >>>
            # Define the specific template file for this form
            template_name = "CBC NEW 2025.docx"
            # Construct the full path to the template inside the 'ReportTemplates' folder
            full_template_path = os.path.join(REPORT_TEMPLATES_FOLDER, template_name)

            doc = DocxTemplate(full_template_path)
            doc.render(context)
            
            # ... (The rest of the file saving logic is the same as before) ...
            target_dir = os.path.join(OUTPUT_FOLDER, "CBC")
            os.makedirs(target_dir, exist_ok=True)
            
            patient_id = context.get("patient_id", "NoID").strip().replace(' ', '_') or "NoID"
            base_fname = f"CBC_{patient_id}.docx"
            out_path = os.path.join(target_dir, base_fname)
            
            counter = 1
            while os.path.exists(out_path):
                new_fname = f"CBC_{patient_id}({counter}).docx"
                out_path = os.path.join(target_dir, new_fname)
                counter += 1
            
            doc.save(out_path)
            
            messagebox.showinfo("Success", f"Report saved:\n{out_path}", parent=self)
            self.clear_all_fields()
            
            os.startfile(out_path)
            self.controller.iconify()

        except FileNotFoundError:
            messagebox.showerror(
                "Template Not Found",
                f"The template file '{template_name}' was not found.\n\nPlease make sure it is inside the '{REPORT_TEMPLATES_FOLDER}' folder.",
                parent=self
            )
        except Exception as e:
            messagebox.showerror("Error", f"Failed to generate report:\n{e}", parent=self)

    def populate_patient_data(self, data):
        """Receives patient data and fills the form's entry fields."""
        # Clear any old data first, except for date fields
        for key, entry in self.patient_entries.items():
            if "date" not in key:
                entry.delete(0, "end")
        
        # Insert new data using keys from the Excel file
        # The .get() method prevents errors if a key is missing
        if 'name' in self.patient_entries:
            self.patient_entries['name'].insert(0, data.get('Patient Name', ''))
        if 'age' in self.patient_entries:
            # Format age to remove .0 if it exists
            age_val = str(data.get('Age', '')).split('.')[0]
            self.patient_entries['age'].insert(0, age_val)
        if 'sex' in self.patient_entries:
            self.patient_entries['sex'].insert(0, data.get('Gender', ''))
        if 'doctor' in self.patient_entries:
            self.patient_entries['doctor'].insert(0, data.get('Ref By', ''))
        if 'patient_id' in self.patient_entries:
            self.patient_entries['patient_id'].insert(0, data.get('Bill Number', ''))

    def clear_patient_fields(self):
        """Clears only the patient-related entry fields."""
        for key, entry in self.patient_entries.items():
            # Keep date fields, clear everything else
            if "date" not in key:
                entry.delete(0, "end")

    def _autocomplete_date(self, event):
        """Autocompletes a date field when focus leaves the widget."""
        widget = event.widget
        day_str = widget.get().strip()

        # Check if the input is a 1 or 2 digit number (a day)
        if day_str.isdigit() and len(day_str) <= 2:
            try:
                day = int(day_str)
                if 1 <= day <= 31:
                    # Get the current month and year
                    now = datetime.now()
                    # Format with leading zeros for day and month
                    full_date = f"{day:02d}.{now.month:02d}.{now.year}"
                    
                    # Update the entry field with the full date
                    widget.delete(0, "end")
                    widget.insert(0, full_date)
            except ValueError:
                # Ignore if it's not a valid number
                pass

# <<< NEW: The Bilirubin Report Form Page >>>
class BilirubinFormFrame(ctk.CTkFrame):
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller
        
        self.patient_entries = {}
        self.test_entries = {}

        # --- Top Bar with Back Button ---
        top_bar = ctk.CTkFrame(self, fg_color="transparent")
        top_bar.pack(fill="x", pady=(6, 12), padx=6)
        
        ctk.CTkButton(top_bar, text=" ⬅️ Back to List", width=120,
                      command=lambda: controller.show_frame(ReportSelectionFrame)).pack(side="left", padx=6, pady=6)
        ctk.CTkLabel(top_bar, text="Bilirubin Report Entry", font=ctk.CTkFont(size=18, weight="bold")).pack(side="left", padx=8)

        # --- Main Content Area ---
        content = ctk.CTkFrame(self)
        content.pack(expand=True, fill="both", padx=6, pady=6)
        left = ctk.CTkFrame(content)
        left.pack(side="left", expand=True, fill="both", padx=(0,8), pady=8)
        right = ctk.CTkFrame(content) # Not scrollable as there are few fields
        right.pack(side="right", expand=True, fill="both", padx=(8,0), pady=8)

        # --- Patient Details Fields ---
        self.add_patient_field(left, 0, "Patient Name", "name")
        self.add_patient_field(left, 1, "Age", "age")
        self.add_patient_field(left, 2, "Sex", "sex")
        self.add_patient_field(left, 3, "Ref. by Dr.", "doctor")
        self.add_patient_field(left, 4, "Patient ID", "patient_id")
        self.add_patient_field(left, 5, "Collection Date", "collection_date")
        self.add_patient_field(left, 6, "Reporting Date", "report_date")
        
        # --- Test Details Fields ---
        self.add_test_field(right, 0, "SERUM BILIRUBIN (TOTAL)", "total")
        self.add_test_field(right, 1, "SERUM BILIRUBIN (DIRECT)", "direct")
        
        ctk.CTkButton(self, text="Save Report", height=40, command=self.generate_report).pack(pady=12, padx=6)

    def add_patient_field(self, parent, row, label, key):
        ctk.CTkLabel(parent, text=label, anchor="w").grid(row=row, column=0, padx=5, pady=4, sticky="w")
        e = ctk.CTkEntry(parent, width=220)
        e.grid(row=row, column=1, padx=5, pady=4)
        self.patient_entries[key] = e
        # This line activates the autocomplete function for date fields
        if "date" in key.lower():
            e.bind("<FocusOut>", self._autocomplete_date)

    def add_test_field(self, parent, row, label, key):
        ctk.CTkLabel(parent, text=label, anchor="w").grid(row=row, column=0, padx=5, pady=4, sticky="w")
        e = ctk.CTkEntry(parent, width=220)
        e.grid(row=row, column=1, padx=5, pady=4)
        self.test_entries[key] = e

    def clear_all_fields(self):
        for e in list(self.patient_entries.values()) + list(self.test_entries.values()):
            e.delete(0, "end")

    def generate_report(self):
        def get_float(key):
            try:
                return float(self.test_entries[key].get())
            except (ValueError, KeyError):
                return 0.0

        context = {key: entry.get() for key, entry in self.patient_entries.items()}
        context.update({key: entry.get() for key, entry in self.test_entries.items()})
        
        # Perform the calculation for INDIRECT bilirubin
        bili_total = get_float("total")
        bili_direct = get_float("direct")
        bili_indirect = bili_total - bili_direct
        context['indirect'] = f"{bili_indirect:.2f}"

        try:
            template_name = "Bilirubin_Template.docx"
            full_template_path = os.path.join(REPORT_TEMPLATES_FOLDER, template_name)

            doc = DocxTemplate(full_template_path)
            doc.render(context)
            
            # Save in a "Bilirubin" subfolder
            target_dir = os.path.join(OUTPUT_FOLDER, "Bilirubin")
            os.makedirs(target_dir, exist_ok=True)
            
            patient_id = context.get("patient_id", "NoID").strip().replace(' ', '_') or "NoID"
            base_fname = f"Bilirubin_{patient_id}.docx"
            out_path = os.path.join(target_dir, base_fname)
            
            counter = 1
            while os.path.exists(out_path):
                new_fname = f"Bilirubin_{patient_id}({counter}).docx"
                out_path = os.path.join(target_dir, new_fname)
                counter += 1
            
            doc.save(out_path)
            
            messagebox.showinfo("Success", f"Report saved:\n{out_path}", parent=self)
            self.clear_all_fields()
            
            os.startfile(out_path)
            self.controller.iconify()

        except FileNotFoundError:
            messagebox.showerror(
                "Template Not Found",
                f"The template file '{template_name}' was not found.\n\nPlease make sure it is inside the '{REPORT_TEMPLATES_FOLDER}' folder.",
                parent=self
            )
        except Exception as e:
            messagebox.showerror("Error", f"Failed to generate report:\n{e}", parent=self)
    def populate_patient_data(self, data):
        """Receives patient data and fills the form's entry fields."""
        # Clear any old data first, except for date fields
        for key, entry in self.patient_entries.items():
            if "date" not in key:
                entry.delete(0, "end")
        
        # Insert new data using keys from the Excel file
        # The .get() method prevents errors if a key is missing
        if 'name' in self.patient_entries:
            self.patient_entries['name'].insert(0, data.get('Patient Name', ''))
        if 'age' in self.patient_entries:
            # Format age to remove .0 if it exists
            age_val = str(data.get('Age', '')).split('.')[0]
            self.patient_entries['age'].insert(0, age_val)
        if 'sex' in self.patient_entries:
            self.patient_entries['sex'].insert(0, data.get('Gender', ''))
        if 'doctor' in self.patient_entries:
            self.patient_entries['doctor'].insert(0, data.get('Ref By', ''))
        if 'patient_id' in self.patient_entries:
            self.patient_entries['patient_id'].insert(0, data.get('Bill Number', ''))

    def clear_patient_fields(self):
        """Clears only the patient-related entry fields."""
        for key, entry in self.patient_entries.items():
            # Keep date fields, clear everything else
            if "date" not in key:
                entry.delete(0, "end")

    def _autocomplete_date(self, event):
        """Autocompletes a date field when focus leaves the widget."""
        widget = event.widget
        day_str = widget.get().strip()

        # Check if the input is a 1 or 2 digit number (a day)
        if day_str.isdigit() and len(day_str) <= 2:
            try:
                day = int(day_str)
                if 1 <= day <= 31:
                    # Get the current month and year
                    now = datetime.now()
                    # Format with leading zeros for day and month
                    full_date = f"{day:02d}.{now.month:02d}.{now.year}"
                    
                    # Update the entry field with the full date
                    widget.delete(0, "end")
                    widget.insert(0, full_date)
            except ValueError:
                # Ignore if it's not a valid number
                pass

class ElectrolytesFormFrame(ctk.CTkFrame):
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller
        
        self.patient_entries = {}
        self.test_entries = {}

        # --- Top Bar with Back Button ---
        top_bar = ctk.CTkFrame(self, fg_color="transparent")
        top_bar.pack(fill="x", pady=(6, 12), padx=6)
        
        ctk.CTkButton(top_bar, text=" ⬅️ Back to List", width=120,
                      command=lambda: controller.show_frame(ReportSelectionFrame)).pack(side="left", padx=6, pady=6)
        ctk.CTkLabel(top_bar, text="FBS_UREA, CREAT, NA+,K+, CL- NEW", font=ctk.CTkFont(size=18, weight="bold")).pack(side="left", padx=8)

        # --- Main Content Area ---
        content = ctk.CTkFrame(self)
        content.pack(expand=True, fill="both", padx=6, pady=6)
        left = ctk.CTkFrame(content)
        left.pack(side="left", expand=True, fill="both", padx=(0,8), pady=8)
        right = ctk.CTkScrollableFrame(content)
        right.pack(side="right", expand=True, fill="both", padx=(8,0), pady=8)

        # --- Patient Details Fields ---
        self.add_patient_field(left, 0, "Patient Name", "name")
        self.add_patient_field(left, 1, "Age", "age")
        self.add_patient_field(left, 2, "Sex", "sex")
        self.add_patient_field(left, 3, "Ref. by Dr.", "doctor")
        self.add_patient_field(left, 4, "Patient ID", "patient_id")
        self.add_patient_field(left, 5, "Collection Date", "collection_date")
        self.add_patient_field(left, 6, "Reporting Date", "report_date")
        
        # --- Test Details Fields ---
        self.add_test_field(right, 0, "SUGAR (F)", "sugar_f")
        self.add_test_field(right, 1, "UREA", "urea")
        self.add_test_field(right, 2, "CREATININE", "creatinine")
        self.add_test_field(right, 3, "URIC ACID", "uric_acid")
        self.add_test_field(right, 4, "SODIUM (Na+)", "sodium")
        self.add_test_field(right, 5, "POTASSIUM (K+)", "potassium")
        self.add_test_field(right, 6, "CALCIUM (Ca++)", "calcium")
        self.add_test_field(right, 7, "CHLORIDE (Cl-)", "chloride")
        
        ctk.CTkButton(self, text="Save Report", height=40, command=self.generate_report).pack(pady=12, padx=6)

    def add_patient_field(self, parent, row, label, key):
        ctk.CTkLabel(parent, text=label, anchor="w").grid(row=row, column=0, padx=5, pady=4, sticky="w")
        e = ctk.CTkEntry(parent, width=220)
        e.grid(row=row, column=1, padx=5, pady=4)
        self.patient_entries[key] = e
        # This line activates the autocomplete function for date fields
        if "date" in key.lower():
            e.bind("<FocusOut>", self._autocomplete_date)

    def add_test_field(self, parent, row, label, key):
        ctk.CTkLabel(parent, text=label, anchor="w").grid(row=row, column=0, padx=5, pady=4, sticky="w")
        e = ctk.CTkEntry(parent, width=220)
        e.grid(row=row, column=1, padx=5, pady=4)
        self.test_entries[key] = e

    def clear_all_fields(self):
        for e in list(self.patient_entries.values()) + list(self.test_entries.values()):
            e.delete(0, "end")

    def generate_report(self):
        # Since there are no calculations, we just collect the data
        context = {key: entry.get() for key, entry in self.patient_entries.items()}
        context.update({key: entry.get() for key, entry in self.test_entries.items()})

        try:
            template_name = "FBS_UREA, CREAT, NA+,K+, CL- NEW.docx"
            full_template_path = os.path.join(REPORT_TEMPLATES_FOLDER, template_name)

            doc = DocxTemplate(full_template_path)
            doc.render(context)
            
            # Save in an "Electrolytes" subfolder
            target_dir = os.path.join(OUTPUT_FOLDER, "FBS_UREA, CREAT, NA+,K+, CL- NEW")
            os.makedirs(target_dir, exist_ok=True)
            
            patient_id = context.get("patient_id", "NoID").strip().replace(' ', '_') or "NoID"
            base_fname = f"FBS_UREA, CREAT, NA+,K+, CL- NEW_{patient_id}.docx"
            out_path = os.path.join(target_dir, base_fname)
            
            counter = 1
            while os.path.exists(out_path):
                new_fname = f"FBS_UREA, CREAT, NA+,K+, CL- NEW_{patient_id}({counter}).docx"
                out_path = os.path.join(target_dir, new_fname)
                counter += 1
            
            doc.save(out_path)
            
            messagebox.showinfo("Success", f"Report saved:\n{out_path}", parent=self)
            self.clear_all_fields()
            
            os.startfile(out_path)
            self.controller.iconify()

        except FileNotFoundError:
            messagebox.showerror(
                "Template Not Found",
                f"The template file '{template_name}' was not found.\n\nPlease make sure it is inside the '{REPORT_TEMPLATES_FOLDER}' folder.",
                parent=self
            )
        except Exception as e:
            messagebox.showerror("Error", f"Failed to generate report:\n{e}", parent=self)
    def populate_patient_data(self, data):
        """Receives patient data and fills the form's entry fields."""
        # Clear any old data first, except for date fields
        for key, entry in self.patient_entries.items():
            if "date" not in key:
                entry.delete(0, "end")
        
        # Insert new data using keys from the Excel file
        # The .get() method prevents errors if a key is missing
        if 'name' in self.patient_entries:
            self.patient_entries['name'].insert(0, data.get('Patient Name', ''))
        if 'age' in self.patient_entries:
            # Format age to remove .0 if it exists
            age_val = str(data.get('Age', '')).split('.')[0]
            self.patient_entries['age'].insert(0, age_val)
        if 'sex' in self.patient_entries:
            self.patient_entries['sex'].insert(0, data.get('Gender', ''))
        if 'doctor' in self.patient_entries:
            self.patient_entries['doctor'].insert(0, data.get('Ref By', ''))
        if 'patient_id' in self.patient_entries:
            self.patient_entries['patient_id'].insert(0, data.get('Bill Number', ''))

    def clear_patient_fields(self):
        """Clears only the patient-related entry fields."""
        for key, entry in self.patient_entries.items():
            # Keep date fields, clear everything else
            if "date" not in key:
                entry.delete(0, "end")

    def _autocomplete_date(self, event):
        """Autocompletes a date field when focus leaves the widget."""
        widget = event.widget
        day_str = widget.get().strip()

        # Check if the input is a 1 or 2 digit number (a day)
        if day_str.isdigit() and len(day_str) <= 2:
            try:
                day = int(day_str)
                if 1 <= day <= 31:
                    # Get the current month and year
                    now = datetime.now()
                    # Format with leading zeros for day and month
                    full_date = f"{day:02d}.{now.month:02d}.{now.year}"
                    
                    # Update the entry field with the full date
                    widget.delete(0, "end")
                    widget.insert(0, full_date)
            except ValueError:
                # Ignore if it's not a valid number
                pass

class LipidProfileFormFrame(ctk.CTkFrame):
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller
        
        self.patient_entries = {}
        self.test_entries = {}

        top_bar = ctk.CTkFrame(self, fg_color="transparent")
        top_bar.pack(fill="x", pady=(6, 12), padx=6)
        
        ctk.CTkButton(top_bar, text=" ⬅️ Back to List", width=120,
                      command=lambda: controller.show_frame(ReportSelectionFrame)).pack(side="left", padx=6, pady=6)
        ctk.CTkLabel(top_bar, text="Lipid Profile Report Entry", font=ctk.CTkFont(size=18, weight="bold")).pack(side="left", padx=8)

        content = ctk.CTkFrame(self)
        content.pack(expand=True, fill="both", padx=6, pady=6)
        left = ctk.CTkFrame(content)
        left.pack(side="left", expand=True, fill="both", padx=(0,8), pady=8)
        right = ctk.CTkFrame(content)
        right.pack(side="right", expand=True, fill="both", padx=(8,0), pady=8)

        # --- Patient Details Fields ---
        self.add_patient_field(left, 0, "Patient Name", "name")
        self.add_patient_field(left, 1, "Age", "age")
        self.add_patient_field(left, 2, "Sex", "sex")
        self.add_patient_field(left, 3, "Ref. by Dr.", "doctor")
        self.add_patient_field(left, 4, "Patient ID", "patient_id")
        self.add_patient_field(left, 5, "Collection Date", "collection_date")
        self.add_patient_field(left, 6, "Reporting Date", "report_date")
        
        # --- Test Details Fields ---
        self.add_test_field(right, 0, "TOTAL CHOLESTEROL", "total_chol")
        self.add_test_field(right, 1, "TRIGLYCERIDES", "trig")
        self.add_test_field(right, 2, "HDL CHOLESTEROL - DIRECT", "hdl")
        
        ctk.CTkButton(self, text="Save Report", height=40, command=self.generate_report).pack(pady=12, padx=6)

    def add_patient_field(self, parent, row, label, key):
        ctk.CTkLabel(parent, text=label, anchor="w").grid(row=row, column=0, padx=5, pady=4, sticky="w")
        e = ctk.CTkEntry(parent, width=220)
        e.grid(row=row, column=1, padx=5, pady=4)
        self.patient_entries[key] = e
        # This line activates the autocomplete function for date fields
        if "date" in key.lower():
            e.bind("<FocusOut>", self._autocomplete_date)

    def add_test_field(self, parent, row, label, key):
        ctk.CTkLabel(parent, text=label, anchor="w").grid(row=row, column=0, padx=5, pady=4, sticky="w")
        e = ctk.CTkEntry(parent, width=220)
        e.grid(row=row, column=1, padx=5, pady=4)
        self.test_entries[key] = e

    def clear_all_fields(self):
        for e in list(self.patient_entries.values()) + list(self.test_entries.values()):
            e.delete(0, "end")

    def generate_report(self):
        def get_float(key):
            try:
                return float(self.test_entries[key].get())
            except (ValueError, KeyError):
                return 0.0

        context = {key: entry.get() for key, entry in self.patient_entries.items()}
        context.update({key: entry.get() for key, entry in self.test_entries.items()})
        
        # Get manual inputs for calculation
        total_chol = get_float("total_chol")
        trig = get_float("trig")
        hdl = get_float("hdl")
        
        # --- Perform all calculations ---
        vldl = trig / 5 if trig > 0 else 0
        ldl = total_chol - hdl - vldl
        
        # Handle division by zero for ratios
        ldl_hdl_ratio = ldl / hdl if hdl > 0 else 0
        total_hdl_ratio = total_chol / hdl if hdl > 0 else 0
        
        # Add calculated values to the context dictionary
        context['vldl'] = f"{vldl:.2f}"
        context['ldl'] = f"{ldl:.2f}"
        context['ldl_hdl_ratio'] = f"{ldl_hdl_ratio:.2f}"
        context['total_hdl_ratio'] = f"{total_hdl_ratio:.2f}"

        try:
            template_name = "Lipid_Profile_Template.docx"
            full_template_path = os.path.join(REPORT_TEMPLATES_FOLDER, template_name)

            doc = DocxTemplate(full_template_path)
            doc.render(context)
            
            target_dir = os.path.join(OUTPUT_FOLDER, "LipidProfile")
            os.makedirs(target_dir, exist_ok=True)
            
            patient_id = context.get("patient_id", "NoID").strip().replace(' ', '_') or "NoID"
            base_fname = f"Lipid_Profile_{patient_id}.docx"
            out_path = os.path.join(target_dir, base_fname)
            
            counter = 1
            while os.path.exists(out_path):
                new_fname = f"Lipid_Profile_{patient_id}({counter}).docx"
                out_path = os.path.join(target_dir, new_fname)
                counter += 1
            
            doc.save(out_path)
            
            messagebox.showinfo("Success", f"Report saved:\n{out_path}", parent=self)
            self.clear_all_fields()
            
            os.startfile(out_path)
            self.controller.iconify()

        except FileNotFoundError:
            messagebox.showerror(
                "Template Not Found",
                f"The template file '{template_name}' was not found.\n\nPlease make sure it is inside the '{REPORT_TEMPLATES_FOLDER}' folder.",
                parent=self
            )
        except Exception as e:
            messagebox.showerror("Error", f"Failed to generate report:\n{e}", parent=self)

    def populate_patient_data(self, data):
        """Receives patient data and fills the form's entry fields."""
        # Clear any old data first, except for date fields
        for key, entry in self.patient_entries.items():
            if "date" not in key:
                entry.delete(0, "end")
        
        # Insert new data using keys from the Excel file
        # The .get() method prevents errors if a key is missing
        if 'name' in self.patient_entries:
            self.patient_entries['name'].insert(0, data.get('Patient Name', ''))
        if 'age' in self.patient_entries:
            # Format age to remove .0 if it exists
            age_val = str(data.get('Age', '')).split('.')[0]
            self.patient_entries['age'].insert(0, age_val)
        if 'sex' in self.patient_entries:
            self.patient_entries['sex'].insert(0, data.get('Gender', ''))
        if 'doctor' in self.patient_entries:
            self.patient_entries['doctor'].insert(0, data.get('Ref By', ''))
        if 'patient_id' in self.patient_entries:
            self.patient_entries['patient_id'].insert(0, data.get('Bill Number', ''))

    def clear_patient_fields(self):
        """Clears only the patient-related entry fields."""
        for key, entry in self.patient_entries.items():
            # Keep date fields, clear everything else
            if "date" not in key:
                entry.delete(0, "end")

    def _autocomplete_date(self, event):
        """Autocompletes a date field when focus leaves the widget."""
        widget = event.widget
        day_str = widget.get().strip()

        # Check if the input is a 1 or 2 digit number (a day)
        if day_str.isdigit() and len(day_str) <= 2:
            try:
                day = int(day_str)
                if 1 <= day <= 31:
                    # Get the current month and year
                    now = datetime.now()
                    # Format with leading zeros for day and month
                    full_date = f"{day:02d}.{now.month:02d}.{now.year}"
                    
                    # Update the entry field with the full date
                    widget.delete(0, "end")
                    widget.insert(0, full_date)
            except ValueError:
                # Ignore if it's not a valid number
                pass

class LFTFormFrame(ctk.CTkFrame):
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller
        
        self.patient_entries = {}
        self.test_entries = {}

        top_bar = ctk.CTkFrame(self, fg_color="transparent")
        top_bar.pack(fill="x", pady=(6, 12), padx=6)
        
        ctk.CTkButton(top_bar, text=" ⬅️ Back to List", width=120,
                      command=lambda: controller.show_frame(ReportSelectionFrame)).pack(side="left", padx=6, pady=6)
        ctk.CTkLabel(top_bar, text="Liver Function Test (LFT) Report Entry", font=ctk.CTkFont(size=18, weight="bold")).pack(side="left", padx=8)

        content = ctk.CTkFrame(self)
        content.pack(expand=True, fill="both", padx=6, pady=6)
        left = ctk.CTkFrame(content)
        left.pack(side="left", expand=True, fill="both", padx=(0,8), pady=8)
        right = ctk.CTkScrollableFrame(content)
        right.pack(side="right", expand=True, fill="both", padx=(8,0), pady=8)

        # --- Patient Details Fields ---
        self.add_patient_field(left, 0, "Patient Name", "name")
        self.add_patient_field(left, 1, "Age", "age")
        self.add_patient_field(left, 2, "Sex", "sex")
        self.add_patient_field(left, 3, "Ref. by Dr.", "doctor")
        self.add_patient_field(left, 4, "Patient ID", "patient_id")
        self.add_patient_field(left, 5, "Collection Date", "collection_date")
        self.add_patient_field(left, 6, "Reporting Date", "report_date")
        
        # --- Test Details Fields ---
        self.add_test_field(right, 0, "Total bilirubin", "total_bili")
        self.add_test_field(right, 1, "Direct bilirubin", "direct_bili")
        self.add_test_field(right, 2, "SGPT", "sgpt")
        self.add_test_field(right, 3, "SGOT", "sgot")
        self.add_test_field(right, 4, "ALKALINE PHOSPHATE", "alk_phos")
        self.add_test_field(right, 5, "Total protein", "total_prot")
        self.add_test_field(right, 6, "Albumin", "albumin")

        ctk.CTkButton(self, text="Save Report", height=40, command=self.generate_report).pack(pady=12, padx=6)

    def add_patient_field(self, parent, row, label, key):
        ctk.CTkLabel(parent, text=label, anchor="w").grid(row=row, column=0, padx=5, pady=4, sticky="w")
        e = ctk.CTkEntry(parent, width=220)
        e.grid(row=row, column=1, padx=5, pady=4)
        self.patient_entries[key] = e
        # This line activates the autocomplete function for date fields
        if "date" in key.lower():
            e.bind("<FocusOut>", self._autocomplete_date)

    def add_test_field(self, parent, row, label, key):
        ctk.CTkLabel(parent, text=label, anchor="w").grid(row=row, column=0, padx=5, pady=4, sticky="w")
        e = ctk.CTkEntry(parent, width=220)
        e.grid(row=row, column=1, padx=5, pady=4)
        self.test_entries[key] = e

    def clear_all_fields(self):
        for e in list(self.patient_entries.values()) + list(self.test_entries.values()):
            e.delete(0, "end")

    def generate_report(self):
        def get_float(key):
            try:
                return float(self.test_entries[key].get())
            except (ValueError, KeyError):
                return 0.0

        context = {key: entry.get() for key, entry in self.patient_entries.items()}
        context.update({key: entry.get() for key, entry in self.test_entries.items()})
        
        # --- Perform LFT calculations ---
        total_bili = get_float("total_bili")
        direct_bili = get_float("direct_bili")
        total_prot = get_float("total_prot")
        albumin = get_float("albumin")

        indirect_bili = total_bili - direct_bili
        globulin = total_prot - albumin
        ag_ratio = albumin / globulin if globulin > 0 else 0

        # Add calculated values to the context
        context['indirect_bili'] = f"{indirect_bili:.2f}"
        context['globulin'] = f"{globulin:.2f}"
        context['ag_ratio'] = f"{ag_ratio:.2f}"
        
        try:
            template_name = "LFT_Template.docx"
            full_template_path = os.path.join(REPORT_TEMPLATES_FOLDER, template_name)

            doc = DocxTemplate(full_template_path)
            doc.render(context)
            
            target_dir = os.path.join(OUTPUT_FOLDER, "LFT")
            os.makedirs(target_dir, exist_ok=True)
            
            patient_id = context.get("patient_id", "NoID").strip().replace(' ', '_') or "NoID"
            base_fname = f"LFT_{patient_id}.docx"
            out_path = os.path.join(target_dir, base_fname)
            
            counter = 1
            while os.path.exists(out_path):
                new_fname = f"LFT_{patient_id}({counter}).docx"
                out_path = os.path.join(target_dir, new_fname)
                counter += 1
            
            doc.save(out_path)
            
            messagebox.showinfo("Success", f"Report saved:\n{out_path}", parent=self)
            self.clear_all_fields()
            
            os.startfile(out_path)
            self.controller.iconify()

        except FileNotFoundError:
            messagebox.showerror(
                "Template Not Found",
                f"The template file '{template_name}' was not found.\n\nPlease make sure it is inside the '{REPORT_TEMPLATES_FOLDER}' folder.",
                parent=self
            )
        except Exception as e:
            messagebox.showerror("Error", f"Failed to generate report:\n{e}", parent=self)
    def populate_patient_data(self, data):
        """Receives patient data and fills the form's entry fields."""
        # Clear any old data first, except for date fields
        for key, entry in self.patient_entries.items():
            if "date" not in key:
                entry.delete(0, "end")
        
        # Insert new data using keys from the Excel file
        # The .get() method prevents errors if a key is missing
        if 'name' in self.patient_entries:
            self.patient_entries['name'].insert(0, data.get('Patient Name', ''))
        if 'age' in self.patient_entries:
            # Format age to remove .0 if it exists
            age_val = str(data.get('Age', '')).split('.')[0]
            self.patient_entries['age'].insert(0, age_val)
        if 'sex' in self.patient_entries:
            self.patient_entries['sex'].insert(0, data.get('Gender', ''))
        if 'doctor' in self.patient_entries:
            self.patient_entries['doctor'].insert(0, data.get('Ref By', ''))
        if 'patient_id' in self.patient_entries:
            self.patient_entries['patient_id'].insert(0, data.get('Bill Number', ''))

    def clear_patient_fields(self):
        """Clears only the patient-related entry fields."""
        for key, entry in self.patient_entries.items():
            # Keep date fields, clear everything else
            if "date" not in key:
                entry.delete(0, "end")
    
    def _autocomplete_date(self, event):
        """Autocompletes a date field when focus leaves the widget."""
        widget = event.widget
        day_str = widget.get().strip()

        # Check if the input is a 1 or 2 digit number (a day)
        if day_str.isdigit() and len(day_str) <= 2:
            try:
                day = int(day_str)
                if 1 <= day <= 31:
                    # Get the current month and year
                    now = datetime.now()
                    # Format with leading zeros for day and month
                    full_date = f"{day:02d}.{now.month:02d}.{now.year}"
                    
                    # Update the entry field with the full date
                    widget.delete(0, "end")
                    widget.insert(0, full_date)
            except ValueError:
                # Ignore if it's not a valid number
                pass

class RFTFormFrame(ctk.CTkFrame):
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller
        
        self.patient_entries = {}
        self.test_entries = {}

        # --- Top Bar with Back Button ---
        top_bar = ctk.CTkFrame(self, fg_color="transparent")
        top_bar.pack(fill="x", pady=(6, 12), padx=6)
        
        ctk.CTkButton(top_bar, text=" ⬅️ Back to List", width=120,
                      command=lambda: controller.show_frame(ReportSelectionFrame)).pack(side="left", padx=6, pady=6)
        ctk.CTkLabel(top_bar, text="PPBS, UREA, CREAT, NA+,K+, CL- NEW Entry", font=ctk.CTkFont(size=18, weight="bold")).pack(side="left", padx=8)

        # --- Main Content Area ---
        content = ctk.CTkFrame(self)
        content.pack(expand=True, fill="both", padx=6, pady=6)
        left = ctk.CTkFrame(content)
        left.pack(side="left", expand=True, fill="both", padx=(0,8), pady=8)
        right = ctk.CTkScrollableFrame(content)
        right.pack(side="right", expand=True, fill="both", padx=(8,0), pady=8)

        # --- Patient Details Fields ---
        self.add_patient_field(left, 0, "Patient Name", "name")
        self.add_patient_field(left, 1, "Age", "age")
        self.add_patient_field(left, 2, "Sex", "sex")
        self.add_patient_field(left, 3, "Ref. by Dr.", "doctor")
        self.add_patient_field(left, 4, "Patient ID", "patient_id")
        self.add_patient_field(left, 5, "Collection Date", "collection_date")
        self.add_patient_field(left, 6, "Reporting Date", "report_date")
        
        # --- Test Details Fields ---
        self.add_test_field(right, 0, "UREA", "urea")
        self.add_test_field(right, 1, "CREATININE", "creatinine")
        self.add_test_field(right, 2, "URIC ACID", "uric_acid")
        self.add_test_field(right, 3, "SODIUM (Na+)", "sodium")
        self.add_test_field(right, 4, "POTASSIUM (K+)", "potassium")
        self.add_test_field(right, 5, "CALCIUM (Ca++)", "calcium")
        self.add_test_field(right, 6, "CHLORIDE (Cl-)", "chloride")
        
        ctk.CTkButton(self, text="Save Report", height=40, command=self.generate_report).pack(pady=12, padx=6)

    def add_patient_field(self, parent, row, label, key):
        ctk.CTkLabel(parent, text=label, anchor="w").grid(row=row, column=0, padx=5, pady=4, sticky="w")
        e = ctk.CTkEntry(parent, width=220)
        e.grid(row=row, column=1, padx=5, pady=4)
        self.patient_entries[key] = e
        # This line activates the autocomplete function for date fields
        if "date" in key.lower():
            e.bind("<FocusOut>", self._autocomplete_date)

    def add_test_field(self, parent, row, label, key):
        ctk.CTkLabel(parent, text=label, anchor="w").grid(row=row, column=0, padx=5, pady=4, sticky="w")
        e = ctk.CTkEntry(parent, width=220)
        e.grid(row=row, column=1, padx=5, pady=4)
        self.test_entries[key] = e

    def clear_all_fields(self):
        for e in list(self.patient_entries.values()) + list(self.test_entries.values()):
            e.delete(0, "end")

    def generate_report(self):
        context = {key: entry.get() for key, entry in self.patient_entries.items()}
        context.update({key: entry.get() for key, entry in self.test_entries.items()})

        try:
            template_name = "PPBS, UREA, CREAT, NA+,K+, CL- NEW.docx"
            full_template_path = os.path.join(REPORT_TEMPLATES_FOLDER, template_name)

            doc = DocxTemplate(full_template_path)
            doc.render(context)
            
            target_dir = os.path.join(OUTPUT_FOLDER, "PPBS, UREA, CREAT, NA+,K+, CL- NEW")
            os.makedirs(target_dir, exist_ok=True)
            
            patient_id = context.get("patient_id", "NoID").strip().replace(' ', '_') or "NoID"
            base_fname = f"PPBS, UREA, CREAT, NA+,K+, CL- NEW_{patient_id}.docx"
            out_path = os.path.join(target_dir, base_fname)
            
            counter = 1
            while os.path.exists(out_path):
                new_fname = f"PPBS, UREA, CREAT, NA+,K+, CL- NEW_{patient_id}({counter}).docx"
                out_path = os.path.join(target_dir, new_fname)
                counter += 1
            
            doc.save(out_path)
            
            messagebox.showinfo("Success", f"Report saved:\n{out_path}", parent=self)
            self.clear_all_fields()
            
            os.startfile(out_path)
            self.controller.iconify()

        except FileNotFoundError:
            messagebox.showerror(
                "Template Not Found",
                f"The template file '{template_name}' was not found.\n\nPlease make sure it is inside the '{REPORT_TEMPLATES_FOLDER}' folder.",
                parent=self
            )
        except Exception as e:
            messagebox.showerror("Error", f"Failed to generate report:\n{e}", parent=self)
    def populate_patient_data(self, data):
        """Receives patient data and fills the form's entry fields."""
        # Clear any old data first, except for date fields
        for key, entry in self.patient_entries.items():
            if "date" not in key:
                entry.delete(0, "end")
        
        # Insert new data using keys from the Excel file
        # The .get() method prevents errors if a key is missing
        if 'name' in self.patient_entries:
            self.patient_entries['name'].insert(0, data.get('Patient Name', ''))
        if 'age' in self.patient_entries:
            # Format age to remove .0 if it exists
            age_val = str(data.get('Age', '')).split('.')[0]
            self.patient_entries['age'].insert(0, age_val)
        if 'sex' in self.patient_entries:
            self.patient_entries['sex'].insert(0, data.get('Gender', ''))
        if 'doctor' in self.patient_entries:
            self.patient_entries['doctor'].insert(0, data.get('Ref By', ''))
        if 'patient_id' in self.patient_entries:
            self.patient_entries['patient_id'].insert(0, data.get('Bill Number', ''))

    def clear_patient_fields(self):
        """Clears only the patient-related entry fields."""
        for key, entry in self.patient_entries.items():
            # Keep date fields, clear everything else
            if "date" not in key:
                entry.delete(0, "end")

    def _autocomplete_date(self, event):
        """Autocompletes a date field when focus leaves the widget."""
        widget = event.widget
        day_str = widget.get().strip()

        # Check if the input is a 1 or 2 digit number (a day)
        if day_str.isdigit() and len(day_str) <= 2:
            try:
                day = int(day_str)
                if 1 <= day <= 31:
                    # Get the current month and year
                    now = datetime.now()
                    # Format with leading zeros for day and month
                    full_date = f"{day:02d}.{now.month:02d}.{now.year}"
                    
                    # Update the entry field with the full date
                    widget.delete(0, "end")
                    widget.insert(0, full_date)
            except ValueError:
                # Ignore if it's not a valid number
                pass
class SugerPPBSFormFrame(ctk.CTkFrame):
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller
        
        self.patient_entries = {}
        self.test_entries = {}

        # --- Top Bar with Back Button ---
        top_bar = ctk.CTkFrame(self, fg_color="transparent")
        top_bar.pack(fill="x", pady=(6, 12), padx=6)
        
        ctk.CTkButton(top_bar, text=" ⬅️ Back to List", width=120,
                      command=lambda: controller.show_frame(ReportSelectionFrame)).pack(side="left", padx=6, pady=6)
        ctk.CTkLabel(top_bar, text="Suger PPBS Report Entry", font=ctk.CTkFont(size=18, weight="bold")).pack(side="left", padx=8)

        # --- Main Content Area ---
        content = ctk.CTkFrame(self)
        content.pack(expand=True, fill="both", padx=6, pady=6)
        left = ctk.CTkFrame(content)
        left.pack(side="left", expand=True, fill="both", padx=(0,8), pady=8)
        right = ctk.CTkFrame(content)
        right.pack(side="right", expand=True, fill="both", padx=(8,0), pady=8)

        # --- Patient Details Fields ---
        self.add_patient_field(left, 0, "Patient Name", "name")
        self.add_patient_field(left, 1, "Age", "age")
        self.add_patient_field(left, 2, "Sex", "sex")
        self.add_patient_field(left, 3, "Ref. by Dr.", "doctor")
        self.add_patient_field(left, 4, "Patient ID", "patient_id")
        self.add_patient_field(left, 5, "Collection Date", "collection_date")
        self.add_patient_field(left, 6, "Reporting Date", "report_date")
        
        # --- Test Details Fields ---
        self.add_test_field(right, 0, "GLUCOSE POSTPRANDIAL (PLASMA)", "glucose_pp")
        
        ctk.CTkButton(self, text="Save Report", height=40, command=self.generate_report).pack(pady=12, padx=6)

    def add_patient_field(self, parent, row, label, key):
        ctk.CTkLabel(parent, text=label, anchor="w").grid(row=row, column=0, padx=5, pady=4, sticky="w")
        e = ctk.CTkEntry(parent, width=220)
        e.grid(row=row, column=1, padx=5, pady=4)
        self.patient_entries[key] = e
        # This line activates the autocomplete function for date fields
        if "date" in key.lower():
            e.bind("<FocusOut>", self._autocomplete_date)

    def add_test_field(self, parent, row, label, key):
        ctk.CTkLabel(parent, text=label, anchor="w").grid(row=row, column=0, padx=5, pady=4, sticky="w")
        e = ctk.CTkEntry(parent, width=220)
        e.grid(row=row, column=1, padx=5, pady=4)
        self.test_entries[key] = e

    def clear_all_fields(self):
        for e in list(self.patient_entries.values()) + list(self.test_entries.values()):
            e.delete(0, "end")

    def generate_report(self):
        context = {key: entry.get() for key, entry in self.patient_entries.items()}
        context.update({key: entry.get() for key, entry in self.test_entries.items()})

        try:
            template_name = "SUGAR PPBS NEW.docx"
            full_template_path = os.path.join(REPORT_TEMPLATES_FOLDER, template_name)

            doc = DocxTemplate(full_template_path)
            doc.render(context)
            
            target_dir = os.path.join(OUTPUT_FOLDER, "SugerPPBS")
            os.makedirs(target_dir, exist_ok=True)
            
            patient_id = context.get("patient_id", "NoID").strip().replace(' ', '_') or "NoID"
            base_fname = f"Suger_PPBS_{patient_id}.docx"
            out_path = os.path.join(target_dir, base_fname)
            
            counter = 1
            while os.path.exists(out_path):
                new_fname = f"Suger_PPBS_{patient_id}({counter}).docx"
                out_path = os.path.join(target_dir, new_fname)
                counter += 1
            
            doc.save(out_path)
            
            messagebox.showinfo("Success", f"Report saved:\n{out_path}", parent=self)
            self.clear_all_fields()
            
            os.startfile(out_path)
            self.controller.iconify()

        except FileNotFoundError:
            messagebox.showerror(
                "Template Not Found",
                f"The template file '{template_name}' was not found.\n\nPlease make sure it is inside the '{REPORT_TEMPLATES_FOLDER}' folder.",
                parent=self
            )
        except Exception as e:
            messagebox.showerror("Error", f"Failed to generate report:\n{e}", parent=self)
    def populate_patient_data(self, data):
        """Receives patient data and fills the form's entry fields."""
        # Clear any old data first, except for date fields
        for key, entry in self.patient_entries.items():
            if "date" not in key:
                entry.delete(0, "end")
        
        # Insert new data using keys from the Excel file
        # The .get() method prevents errors if a key is missing
        if 'name' in self.patient_entries:
            self.patient_entries['name'].insert(0, data.get('Patient Name', ''))
        if 'age' in self.patient_entries:
            # Format age to remove .0 if it exists
            age_val = str(data.get('Age', '')).split('.')[0]
            self.patient_entries['age'].insert(0, age_val)
        if 'sex' in self.patient_entries:
            self.patient_entries['sex'].insert(0, data.get('Gender', ''))
        if 'doctor' in self.patient_entries:
            self.patient_entries['doctor'].insert(0, data.get('Ref By', ''))
        if 'patient_id' in self.patient_entries:
            self.patient_entries['patient_id'].insert(0, data.get('Bill Number', ''))

    def clear_patient_fields(self):
        """Clears only the patient-related entry fields."""
        for key, entry in self.patient_entries.items():
            # Keep date fields, clear everything else
            if "date" not in key:
                entry.delete(0, "end")
    
    def _autocomplete_date(self, event):
        """Autocompletes a date field when focus leaves the widget."""
        widget = event.widget
        day_str = widget.get().strip()

        # Check if the input is a 1 or 2 digit number (a day)
        if day_str.isdigit() and len(day_str) <= 2:
            try:
                day = int(day_str)
                if 1 <= day <= 31:
                    # Get the current month and year
                    now = datetime.now()
                    # Format with leading zeros for day and month
                    full_date = f"{day:02d}.{now.month:02d}.{now.year}"
                    
                    # Update the entry field with the full date
                    widget.delete(0, "end")
                    widget.insert(0, full_date)
            except ValueError:
                # Ignore if it's not a valid number
                pass

class SugerRBSFormFrame(ctk.CTkFrame):
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller
        
        self.patient_entries = {}
        self.test_entries = {}

        # --- Top Bar with Back Button ---
        top_bar = ctk.CTkFrame(self, fg_color="transparent")
        top_bar.pack(fill="x", pady=(6, 12), padx=6)
        
        ctk.CTkButton(top_bar, text=" ⬅️ Back to List", width=120,
                      command=lambda: controller.show_frame(ReportSelectionFrame)).pack(side="left", padx=6, pady=6)
        ctk.CTkLabel(top_bar, text="Suger RBS Report Entry", font=ctk.CTkFont(size=18, weight="bold")).pack(side="left", padx=8)

        # --- Main Content Area ---
        content = ctk.CTkFrame(self)
        content.pack(expand=True, fill="both", padx=6, pady=6)
        left = ctk.CTkFrame(content)
        left.pack(side="left", expand=True, fill="both", padx=(0,8), pady=8)
        right = ctk.CTkFrame(content)
        right.pack(side="right", expand=True, fill="both", padx=(8,0), pady=8)

        # --- Patient Details Fields ---
        self.add_patient_field(left, 0, "Patient Name", "name")
        self.add_patient_field(left, 1, "Age", "age")
        self.add_patient_field(left, 2, "Sex", "sex")
        self.add_patient_field(left, 3, "Ref. by Dr.", "doctor")
        self.add_patient_field(left, 4, "Patient ID", "patient_id")
        self.add_patient_field(left, 5, "Collection Date", "collection_date")
        self.add_patient_field(left, 6, "Reporting Date", "report_date")
        
        # --- Test Details Fields ---
        self.add_test_field(right, 0, "GLUCOSE RANDOM SUGAR", "glucose_rs")
        
        ctk.CTkButton(self, text="Save Report", height=40, command=self.generate_report).pack(pady=12, padx=6)

    def add_patient_field(self, parent, row, label, key):
        ctk.CTkLabel(parent, text=label, anchor="w").grid(row=row, column=0, padx=5, pady=4, sticky="w")
        e = ctk.CTkEntry(parent, width=220)
        e.grid(row=row, column=1, padx=5, pady=4)
        self.patient_entries[key] = e
        # This line activates the autocomplete function for date fields
        if "date" in key.lower():
            e.bind("<FocusOut>", self._autocomplete_date)

    def add_test_field(self, parent, row, label, key):
        ctk.CTkLabel(parent, text=label, anchor="w").grid(row=row, column=0, padx=5, pady=4, sticky="w")
        e = ctk.CTkEntry(parent, width=220)
        e.grid(row=row, column=1, padx=5, pady=4)
        self.test_entries[key] = e

    def clear_all_fields(self):
        for e in list(self.patient_entries.values()) + list(self.test_entries.values()):
            e.delete(0, "end")

    def generate_report(self):
        context = {key: entry.get() for key, entry in self.patient_entries.items()}
        context.update({key: entry.get() for key, entry in self.test_entries.items()})

        try:
            template_name = "SUGAR RBS  NEW.docx"
            full_template_path = os.path.join(REPORT_TEMPLATES_FOLDER, template_name)

            doc = DocxTemplate(full_template_path)
            doc.render(context)
            
            target_dir = os.path.join(OUTPUT_FOLDER, "SugerRBS")
            os.makedirs(target_dir, exist_ok=True)
            
            patient_id = context.get("patient_id", "NoID").strip().replace(' ', '_') or "NoID"
            base_fname = f"Suger_RBS_{patient_id}.docx"
            out_path = os.path.join(target_dir, base_fname)
            
            counter = 1
            while os.path.exists(out_path):
                new_fname = f"Suger_RBS_{patient_id}({counter}).docx"
                out_path = os.path.join(target_dir, new_fname)
                counter += 1
            
            doc.save(out_path)
            
            messagebox.showinfo("Success", f"Report saved:\n{out_path}", parent=self)
            self.clear_all_fields()
            
            os.startfile(out_path)
            self.controller.iconify()

        except FileNotFoundError:
            messagebox.showerror(
                "Template Not Found",
                f"The template file '{template_name}' was not found.\n\nPlease make sure it is inside the '{REPORT_TEMPLATES_FOLDER}' folder.",
                parent=self
            )
        except Exception as e:
            messagebox.showerror("Error", f"Failed to generate report:\n{e}", parent=self)
    def populate_patient_data(self, data):
        """Receives patient data and fills the form's entry fields."""
        # Clear any old data first, except for date fields
        for key, entry in self.patient_entries.items():
            if "date" not in key:
                entry.delete(0, "end")
        
        # Insert new data using keys from the Excel file
        # The .get() method prevents errors if a key is missing
        if 'name' in self.patient_entries:
            self.patient_entries['name'].insert(0, data.get('Patient Name', ''))
        if 'age' in self.patient_entries:
            # Format age to remove .0 if it exists
            age_val = str(data.get('Age', '')).split('.')[0]
            self.patient_entries['age'].insert(0, age_val)
        if 'sex' in self.patient_entries:
            self.patient_entries['sex'].insert(0, data.get('Gender', ''))
        if 'doctor' in self.patient_entries:
            self.patient_entries['doctor'].insert(0, data.get('Ref By', ''))
        if 'patient_id' in self.patient_entries:
            self.patient_entries['patient_id'].insert(0, data.get('Bill Number', ''))

    def clear_patient_fields(self):
        """Clears only the patient-related entry fields."""
        for key, entry in self.patient_entries.items():
            # Keep date fields, clear everything else
            if "date" not in key:
                entry.delete(0, "end")

    def _autocomplete_date(self, event):
        """Autocompletes a date field when focus leaves the widget."""
        widget = event.widget
        day_str = widget.get().strip()

        # Check if the input is a 1 or 2 digit number (a day)
        if day_str.isdigit() and len(day_str) <= 2:
            try:
                day = int(day_str)
                if 1 <= day <= 31:
                    # Get the current month and year
                    now = datetime.now()
                    # Format with leading zeros for day and month
                    full_date = f"{day:02d}.{now.month:02d}.{now.year}"
                    
                    # Update the entry field with the full date
                    widget.delete(0, "end")
                    widget.insert(0, full_date)
            except ValueError:
                # Ignore if it's not a valid number
                pass

class UricAcidFormFrame(ctk.CTkFrame):
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller
        
        self.patient_entries = {}
        self.test_entries = {}

        # --- Top Bar with Back Button ---
        top_bar = ctk.CTkFrame(self, fg_color="transparent")
        top_bar.pack(fill="x", pady=(6, 12), padx=6)
        
        ctk.CTkButton(top_bar, text=" ⬅️ Back to List", width=120,
                      command=lambda: controller.show_frame(ReportSelectionFrame)).pack(side="left", padx=6, pady=6)
        ctk.CTkLabel(top_bar, text="Uric Acid Report Entry", font=ctk.CTkFont(size=18, weight="bold")).pack(side="left", padx=8)

        # --- Main Content Area ---
        content = ctk.CTkFrame(self)
        content.pack(expand=True, fill="both", padx=6, pady=6)
        left = ctk.CTkFrame(content)
        left.pack(side="left", expand=True, fill="both", padx=(0,8), pady=8)
        right = ctk.CTkFrame(content)
        right.pack(side="right", expand=True, fill="both", padx=(8,0), pady=8)

        # --- Patient Details Fields ---
        self.add_patient_field(left, 0, "Patient Name", "name")
        self.add_patient_field(left, 1, "Age", "age")
        self.add_patient_field(left, 2, "Sex", "sex")
        self.add_patient_field(left, 3, "Ref. by Dr.", "doctor")
        self.add_patient_field(left, 4, "Patient ID", "patient_id")
        self.add_patient_field(left, 5, "Collection Date", "collection_date")
        self.add_patient_field(left, 6, "Reporting Date", "report_date")
        
        # --- Test Details Fields ---
        self.add_test_field(right, 0, "Serum uric acid", "uric_acid")
        
        ctk.CTkButton(self, text="Save Report", height=40, command=self.generate_report).pack(pady=12, padx=6)

    def add_patient_field(self, parent, row, label, key):
        ctk.CTkLabel(parent, text=label, anchor="w").grid(row=row, column=0, padx=5, pady=4, sticky="w")
        e = ctk.CTkEntry(parent, width=220)
        e.grid(row=row, column=1, padx=5, pady=4)
        self.patient_entries[key] = e
        # This line activates the autocomplete function for date fields
        if "date" in key.lower():
            e.bind("<FocusOut>", self._autocomplete_date)

    def add_test_field(self, parent, row, label, key):
        ctk.CTkLabel(parent, text=label, anchor="w").grid(row=row, column=0, padx=5, pady=4, sticky="w")
        e = ctk.CTkEntry(parent, width=220)
        e.grid(row=row, column=1, padx=5, pady=4)
        self.test_entries[key] = e

    def clear_all_fields(self):
        for e in list(self.patient_entries.values()) + list(self.test_entries.values()):
            e.delete(0, "end")

    def generate_report(self):
        context = {key: entry.get() for key, entry in self.patient_entries.items()}
        context.update({key: entry.get() for key, entry in self.test_entries.items()})

        try:
            template_name = "URIC ACID   NEW.docx"
            full_template_path = os.path.join(REPORT_TEMPLATES_FOLDER, template_name)

            doc = DocxTemplate(full_template_path)
            doc.render(context)
            
            target_dir = os.path.join(OUTPUT_FOLDER, "UricAcid")
            os.makedirs(target_dir, exist_ok=True)
            
            patient_id = context.get("patient_id", "NoID").strip().replace(' ', '_') or "NoID"
            base_fname = f"Uric_Acid_{patient_id}.docx"
            out_path = os.path.join(target_dir, base_fname)
            
            counter = 1
            while os.path.exists(out_path):
                new_fname = f"Uric_Acid_{patient_id}({counter}).docx"
                out_path = os.path.join(target_dir, new_fname)
                counter += 1
            
            doc.save(out_path)
            
            messagebox.showinfo("Success", f"Report saved:\n{out_path}", parent=self)
            self.clear_all_fields()
            
            os.startfile(out_path)
            self.controller.iconify()

        except FileNotFoundError:
            messagebox.showerror(
                "Template Not Found",
                f"The template file '{template_name}' was not found.\n\nPlease make sure it is inside the '{REPORT_TEMPLATES_FOLDER}' folder.",
                parent=self
            )
        except Exception as e:
            messagebox.showerror("Error", f"Failed to generate report:\n{e}", parent=self)
    def populate_patient_data(self, data):
        """Receives patient data and fills the form's entry fields."""
        # Clear any old data first, except for date fields
        for key, entry in self.patient_entries.items():
            if "date" not in key:
                entry.delete(0, "end")
        
        # Insert new data using keys from the Excel file
        # The .get() method prevents errors if a key is missing
        if 'name' in self.patient_entries:
            self.patient_entries['name'].insert(0, data.get('Patient Name', ''))
        if 'age' in self.patient_entries:
            # Format age to remove .0 if it exists
            age_val = str(data.get('Age', '')).split('.')[0]
            self.patient_entries['age'].insert(0, age_val)
        if 'sex' in self.patient_entries:
            self.patient_entries['sex'].insert(0, data.get('Gender', ''))
        if 'doctor' in self.patient_entries:
            self.patient_entries['doctor'].insert(0, data.get('Ref By', ''))
        if 'patient_id' in self.patient_entries:
            self.patient_entries['patient_id'].insert(0, data.get('Bill Number', ''))

    def clear_patient_fields(self):
        """Clears only the patient-related entry fields."""
        for key, entry in self.patient_entries.items():
            # Keep date fields, clear everything else
            if "date" not in key:
                entry.delete(0, "end")

    def _autocomplete_date(self, event):
        """Autocompletes a date field when focus leaves the widget."""
        widget = event.widget
        day_str = widget.get().strip()

        # Check if the input is a 1 or 2 digit number (a day)
        if day_str.isdigit() and len(day_str) <= 2:
            try:
                day = int(day_str)
                if 1 <= day <= 31:
                    # Get the current month and year
                    now = datetime.now()
                    # Format with leading zeros for day and month
                    full_date = f"{day:02d}.{now.month:02d}.{now.year}"
                    
                    # Update the entry field with the full date
                    widget.delete(0, "end")
                    widget.insert(0, full_date)
            except ValueError:
                # Ignore if it's not a valid number
                pass

class RenalProfileFormFrame(ctk.CTkFrame):
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller
        
        self.patient_entries = {}
        self.test_entries = {}

        # --- Top Bar with Back Button ---
        top_bar = ctk.CTkFrame(self, fg_color="transparent")
        top_bar.pack(fill="x", pady=(6, 12), padx=6)
        
        ctk.CTkButton(top_bar, text=" ⬅️ Back to List", width=120,
                      command=lambda: controller.show_frame(ReportSelectionFrame)).pack(side="left", padx=6, pady=6)
        ctk.CTkLabel(top_bar, text="Renal Profile Report Entry", font=ctk.CTkFont(size=18, weight="bold")).pack(side="left", padx=8)

        # --- Main Content Area ---
        content = ctk.CTkFrame(self)
        content.pack(expand=True, fill="both", padx=6, pady=6)
        left = ctk.CTkFrame(content)
        left.pack(side="left", expand=True, fill="both", padx=(0,8), pady=8)
        right = ctk.CTkScrollableFrame(content)
        right.pack(side="right", expand=True, fill="both", padx=(8,0), pady=8)

        # --- Patient Details Fields ---
        self.add_patient_field(left, 0, "Patient Name", "name")
        self.add_patient_field(left, 1, "Age", "age")
        self.add_patient_field(left, 2, "Sex", "sex")
        self.add_patient_field(left, 3, "Ref. by Dr.", "doctor")
        self.add_patient_field(left, 4, "Patient ID", "patient_id")
        self.add_patient_field(left, 5, "Collection Date", "collection_date")
        self.add_patient_field(left, 6, "Reporting Date", "report_date")
        
        # --- Test Details Fields ---
        self.add_test_field(right, 0, "SUGAR (PP)", "sugar_pp")
        self.add_test_field(right, 1, "UREA", "urea")
        self.add_test_field(right, 2, "CREATININE", "creatinine")
        self.add_test_field(right, 3, "URIC ACID", "uric_acid")
        self.add_test_field(right, 4, "SODIUM (Na+)", "sodium")
        self.add_test_field(right, 5, "POTASSIUM (K+)", "potassium")
        self.add_test_field(right, 6, "CALCIUM (Ca++)", "calcium")
        self.add_test_field(right, 7, "CHLORIDE (Cl-)", "chloride")
        
        ctk.CTkButton(self, text="Save Report", height=40, command=self.generate_report).pack(pady=12, padx=6)

    def add_patient_field(self, parent, row, label, key):
        ctk.CTkLabel(parent, text=label, anchor="w").grid(row=row, column=0, padx=5, pady=4, sticky="w")
        e = ctk.CTkEntry(parent, width=220)
        e.grid(row=row, column=1, padx=5, pady=4)
        self.patient_entries[key] = e
        # This line activates the autocomplete function for date fields
        if "date" in key.lower():
            e.bind("<FocusOut>", self._autocomplete_date)

    def add_test_field(self, parent, row, label, key):
        ctk.CTkLabel(parent, text=label, anchor="w").grid(row=row, column=0, padx=5, pady=4, sticky="w")
        e = ctk.CTkEntry(parent, width=220)
        e.grid(row=row, column=1, padx=5, pady=4)
        self.test_entries[key] = e

    def clear_all_fields(self):
        for e in list(self.patient_entries.values()) + list(self.test_entries.values()):
            e.delete(0, "end")

    def generate_report(self):
        context = {key: entry.get() for key, entry in self.patient_entries.items()}
        context.update({key: entry.get() for key, entry in self.test_entries.items()})

        try:
            template_name = "RENAL PROFILE NEW.docx"
            full_template_path = os.path.join(REPORT_TEMPLATES_FOLDER, template_name)

            doc = DocxTemplate(full_template_path)
            doc.render(context)
            
            target_dir = os.path.join(OUTPUT_FOLDER, "RenalProfile")
            os.makedirs(target_dir, exist_ok=True)
            
            patient_id = context.get("patient_id", "NoID").strip().replace(' ', '_') or "NoID"
            base_fname = f"Renal_Profile_{patient_id}.docx"
            out_path = os.path.join(target_dir, base_fname)
            
            counter = 1
            while os.path.exists(out_path):
                new_fname = f"Renal_Profile_{patient_id}({counter}).docx"
                out_path = os.path.join(target_dir, new_fname)
                counter += 1
            
            doc.save(out_path)
            
            messagebox.showinfo("Success", f"Report saved:\n{out_path}", parent=self)
            self.clear_all_fields()
            
            os.startfile(out_path)
            self.controller.iconify()

        except FileNotFoundError:
            messagebox.showerror(
                "Template Not Found",
                f"The template file '{template_name}' was not found.\n\nPlease make sure it is inside the '{REPORT_TEMPLATES_FOLDER}' folder.",
                parent=self
            )
        except Exception as e:
            messagebox.showerror("Error", f"Failed to generate report:\n{e}", parent=self)
    def populate_patient_data(self, data):
        """Receives patient data and fills the form's entry fields."""
        # Clear any old data first, except for date fields
        for key, entry in self.patient_entries.items():
            if "date" not in key:
                entry.delete(0, "end")
        
        # Insert new data using keys from the Excel file
        # The .get() method prevents errors if a key is missing
        if 'name' in self.patient_entries:
            self.patient_entries['name'].insert(0, data.get('Patient Name', ''))
        if 'age' in self.patient_entries:
            # Format age to remove .0 if it exists
            age_val = str(data.get('Age', '')).split('.')[0]
            self.patient_entries['age'].insert(0, age_val)
        if 'sex' in self.patient_entries:
            self.patient_entries['sex'].insert(0, data.get('Gender', ''))
        if 'doctor' in self.patient_entries:
            self.patient_entries['doctor'].insert(0, data.get('Ref By', ''))
        if 'patient_id' in self.patient_entries:
            self.patient_entries['patient_id'].insert(0, data.get('Bill Number', ''))

    def clear_patient_fields(self):
        """Clears only the patient-related entry fields."""
        for key, entry in self.patient_entries.items():
            # Keep date fields, clear everything else
            if "date" not in key:
                entry.delete(0, "end")

    def _autocomplete_date(self, event):
        """Autocompletes a date field when focus leaves the widget."""
        widget = event.widget
        day_str = widget.get().strip()

        # Check if the input is a 1 or 2 digit number (a day)
        if day_str.isdigit() and len(day_str) <= 2:
            try:
                day = int(day_str)
                if 1 <= day <= 31:
                    # Get the current month and year
                    now = datetime.now()
                    # Format with leading zeros for day and month
                    full_date = f"{day:02d}.{now.month:02d}.{now.year}"
                    
                    # Update the entry field with the full date
                    widget.delete(0, "end")
                    widget.insert(0, full_date)
            except ValueError:
                # Ignore if it's not a valid number
                pass

class ComprehensiveProfileFormFrame(ctk.CTkFrame):
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller
        
        self.patient_entries = {}
        self.test_entries = {}

        # --- Top Bar with Back Button ---
        top_bar = ctk.CTkFrame(self, fg_color="transparent")
        top_bar.pack(fill="x", pady=(6, 12), padx=6)
        
        ctk.CTkButton(top_bar, text=" ⬅️ Back to List", width=120,
                      command=lambda: controller.show_frame(ReportSelectionFrame)).pack(side="left", padx=6, pady=6)
        ctk.CTkLabel(top_bar, text="Comprehensive Profile Report Entry", font=ctk.CTkFont(size=18, weight="bold")).pack(side="left", padx=8)

        # --- Main Content Area ---
        content = ctk.CTkFrame(self)
        content.pack(expand=True, fill="both", padx=6, pady=6)
        left = ctk.CTkFrame(content)
        left.pack(side="left", expand=True, fill="both", padx=(0,8), pady=8)
        right = ctk.CTkScrollableFrame(content)
        right.pack(side="right", expand=True, fill="both", padx=(8,0), pady=8)

        # --- Patient Details Fields ---
        self.add_patient_field(left, 0, "Patient Name", "name")
        self.add_patient_field(left, 1, "Age", "age")
        self.add_patient_field(left, 2, "Sex", "sex")
        self.add_patient_field(left, 3, "Ref. by Dr.", "doctor")
        self.add_patient_field(left, 4, "Patient ID", "patient_id")
        self.add_patient_field(left, 5, "Collection Date", "collection_date")
        self.add_patient_field(left, 6, "Reporting Date", "report_date")
        
        # --- Test Details Fields ---
        self.add_test_field(right, 0, "SUGAR (F)", "sugar_f")
        self.add_test_field(right, 1, "SUGAR (PP)", "sugar_pp")
        self.add_test_field(right, 2, "UREA", "urea")
        self.add_test_field(right, 3, "CREATININE", "creatinine")
        self.add_test_field(right, 4, "URIC ACID", "uric_acid")
        self.add_test_field(right, 5, "SODIUM (Na+)", "sodium")
        self.add_test_field(right, 6, "POTASSIUM (K+)", "potassium")
        self.add_test_field(right, 7, "TOTAL BILIRUBIN", "total_bili")
        self.add_test_field(right, 8, "SGPT", "sgpt")
        self.add_test_field(right, 9, "SGOT", "sgot")
        self.add_test_field(right, 10, "TRIGLYCERIDE", "trig")
        self.add_test_field(right, 11, "T. CHOLESTEROL", "total_chol")

        ctk.CTkButton(self, text="Save Report", height=40, command=self.generate_report).pack(pady=12, padx=6)

    def add_patient_field(self, parent, row, label, key):
        ctk.CTkLabel(parent, text=label, anchor="w").grid(row=row, column=0, padx=5, pady=4, sticky="w")
        e = ctk.CTkEntry(parent, width=220)
        e.grid(row=row, column=1, padx=5, pady=4)
        self.patient_entries[key] = e
        # This line activates the autocomplete function for date fields
        if "date" in key.lower():
            e.bind("<FocusOut>", self._autocomplete_date)

    def add_test_field(self, parent, row, label, key):
        ctk.CTkLabel(parent, text=label, anchor="w").grid(row=row, column=0, padx=5, pady=4, sticky="w")
        e = ctk.CTkEntry(parent, width=220)
        e.grid(row=row, column=1, padx=5, pady=4)
        self.test_entries[key] = e

    def clear_all_fields(self):
        for e in list(self.patient_entries.values()) + list(self.test_entries.values()):
            e.delete(0, "end")

    def generate_report(self):
        context = {key: entry.get() for key, entry in self.patient_entries.items()}
        context.update({key: entry.get() for key, entry in self.test_entries.items()})

        try:
            template_name = "SUG PP,FBS,NA,K,U,C,UA.OT,PT.BILI ,CHOL, TGL NEW.docx"
            full_template_path = os.path.join(REPORT_TEMPLATES_FOLDER, template_name)

            doc = DocxTemplate(full_template_path)
            doc.render(context)
            
            target_dir = os.path.join(OUTPUT_FOLDER, "SUG PP,FBS,NA+,K+,U,C,UA.OT,PT.BILI ,CHOL, TGL  NEW")
            os.makedirs(target_dir, exist_ok=True)
            
            patient_id = context.get("patient_id", "NoID").strip().replace(' ', '_') or "NoID"
            base_fname = f"SUG PP,FBS,NA+,K+,U,C,UA.OT,PT.BILI ,CHOL, TGL  NEW_{patient_id}.docx"
            out_path = os.path.join(target_dir, base_fname)
            
            counter = 1
            while os.path.exists(out_path):
                new_fname = f"SUG PP,FBS,NA+,K+,U,C,UA.OT,PT.BILI ,CHOL, TGL  NEW_{patient_id}({counter}).docx"
                out_path = os.path.join(target_dir, new_fname)
                counter += 1
            
            doc.save(out_path)
            
            messagebox.showinfo("Success", f"Report saved:\n{out_path}", parent=self)
            self.clear_all_fields()
            
            os.startfile(out_path)
            self.controller.iconify()

        except FileNotFoundError:
            messagebox.showerror(
                "Template Not Found",
                f"The template file '{template_name}' was not found.\n\nPlease make sure it is inside the '{REPORT_TEMPLATES_FOLDER}' folder.",
                parent=self
            )
        except Exception as e:
            messagebox.showerror("Error", f"Failed to generate report:\n{e}", parent=self)
    def populate_patient_data(self, data):
        """Receives patient data and fills the form's entry fields."""
        # Clear any old data first, except for date fields
        for key, entry in self.patient_entries.items():
            if "date" not in key:
                entry.delete(0, "end")
        
        # Insert new data using keys from the Excel file
        # The .get() method prevents errors if a key is missing
        if 'name' in self.patient_entries:
            self.patient_entries['name'].insert(0, data.get('Patient Name', ''))
        if 'age' in self.patient_entries:
            # Format age to remove .0 if it exists
            age_val = str(data.get('Age', '')).split('.')[0]
            self.patient_entries['age'].insert(0, age_val)
        if 'sex' in self.patient_entries:
            self.patient_entries['sex'].insert(0, data.get('Gender', ''))
        if 'doctor' in self.patient_entries:
            self.patient_entries['doctor'].insert(0, data.get('Ref By', ''))
        if 'patient_id' in self.patient_entries:
            self.patient_entries['patient_id'].insert(0, data.get('Bill Number', ''))

    def clear_patient_fields(self):
        """Clears only the patient-related entry fields."""
        for key, entry in self.patient_entries.items():
            # Keep date fields, clear everything else
            if "date" not in key:
                entry.delete(0, "end")

    def _autocomplete_date(self, event):
        """Autocompletes a date field when focus leaves the widget."""
        widget = event.widget
        day_str = widget.get().strip()

        # Check if the input is a 1 or 2 digit number (a day)
        if day_str.isdigit() and len(day_str) <= 2:
            try:
                day = int(day_str)
                if 1 <= day <= 31:
                    # Get the current month and year
                    now = datetime.now()
                    # Format with leading zeros for day and month
                    full_date = f"{day:02d}.{now.month:02d}.{now.year}"
                    
                    # Update the entry field with the full date
                    widget.delete(0, "end")
                    widget.insert(0, full_date)
            except ValueError:
                # Ignore if it's not a valid number
                pass

class UreaCreatinineFormFrame(ctk.CTkFrame):
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller
        
        self.patient_entries = {}
        self.test_entries = {}

        # --- Top Bar with Back Button ---
        top_bar = ctk.CTkFrame(self, fg_color="transparent")
        top_bar.pack(fill="x", pady=(6, 12), padx=6)
        
        ctk.CTkButton(top_bar, text=" ⬅️ Back to List", width=120,
                      command=lambda: controller.show_frame(ReportSelectionFrame)).pack(side="left", padx=6, pady=6)
        ctk.CTkLabel(top_bar, text="Urea & Creatinine Report Entry", font=ctk.CTkFont(size=18, weight="bold")).pack(side="left", padx=8)

        # --- Main Content Area ---
        content = ctk.CTkFrame(self)
        content.pack(expand=True, fill="both", padx=6, pady=6)
        left = ctk.CTkFrame(content)
        left.pack(side="left", expand=True, fill="both", padx=(0,8), pady=8)
        right = ctk.CTkFrame(content)
        right.pack(side="right", expand=True, fill="both", padx=(8,0), pady=8)

        # --- Patient Details Fields ---
        self.add_patient_field(left, 0, "Patient Name", "name")
        self.add_patient_field(left, 1, "Age", "age")
        self.add_patient_field(left, 2, "Sex", "sex")
        self.add_patient_field(left, 3, "Ref. by Dr.", "doctor")
        self.add_patient_field(left, 4, "Patient ID", "patient_id")
        self.add_patient_field(left, 5, "Collection Date", "collection_date")
        self.add_patient_field(left, 6, "Reporting Date", "report_date")
        
        # --- Test Details Fields ---
        self.add_test_field(right, 0, "SERUM UREA", "urea")
        self.add_test_field(right, 1, "SERUM CREATININE", "creatinine")
        
        ctk.CTkButton(self, text="Save Report", height=40, command=self.generate_report).pack(pady=12, padx=6)

    def add_patient_field(self, parent, row, label, key):
        ctk.CTkLabel(parent, text=label, anchor="w").grid(row=row, column=0, padx=5, pady=4, sticky="w")
        e = ctk.CTkEntry(parent, width=220)
        e.grid(row=row, column=1, padx=5, pady=4)
        self.patient_entries[key] = e
        # This line activates the autocomplete function for date fields
        if "date" in key.lower():
            e.bind("<FocusOut>", self._autocomplete_date)

    def add_test_field(self, parent, row, label, key):
        ctk.CTkLabel(parent, text=label, anchor="w").grid(row=row, column=0, padx=5, pady=4, sticky="w")
        e = ctk.CTkEntry(parent, width=220)
        e.grid(row=row, column=1, padx=5, pady=4)
        self.test_entries[key] = e

    def clear_all_fields(self):
        for e in list(self.patient_entries.values()) + list(self.test_entries.values()):
            e.delete(0, "end")

    def generate_report(self):
        context = {key: entry.get() for key, entry in self.patient_entries.items()}
        context.update({key: entry.get() for key, entry in self.test_entries.items()})

        try:
            template_name = "UREA CREATININE   NEW.docx"
            full_template_path = os.path.join(REPORT_TEMPLATES_FOLDER, template_name)

            doc = DocxTemplate(full_template_path)
            doc.render(context)
            
            target_dir = os.path.join(OUTPUT_FOLDER, "UreaCreatinine")
            os.makedirs(target_dir, exist_ok=True)
            
            patient_id = context.get("patient_id", "NoID").strip().replace(' ', '_') or "NoID"
            base_fname = f"Urea_Creatinine_{patient_id}.docx"
            out_path = os.path.join(target_dir, base_fname)
            
            counter = 1
            while os.path.exists(out_path):
                new_fname = f"Urea_Creatinine_{patient_id}({counter}).docx"
                out_path = os.path.join(target_dir, new_fname)
                counter += 1
            
            doc.save(out_path)
            
            messagebox.showinfo("Success", f"Report saved:\n{out_path}", parent=self)
            self.clear_all_fields()
            
            os.startfile(out_path)
            self.controller.iconify()

        except FileNotFoundError:
            messagebox.showerror(
                "Template Not Found",
                f"The template file '{template_name}' was not found.\n\nPlease make sure it is inside the '{REPORT_TEMPLATES_FOLDER}' folder.",
                parent=self
            )
        except Exception as e:
            messagebox.showerror("Error", f"Failed to generate report:\n{e}", parent=self)
    
    def populate_patient_data(self, data):
        """Receives patient data and fills the form's entry fields."""
        # Clear any old data first, except for date fields
        for key, entry in self.patient_entries.items():
            if "date" not in key:
                entry.delete(0, "end")
        
        # Insert new data using keys from the Excel file
        # The .get() method prevents errors if a key is missing
        if 'name' in self.patient_entries:
            self.patient_entries['name'].insert(0, data.get('Patient Name', ''))
        if 'age' in self.patient_entries:
            # Format age to remove .0 if it exists
            age_val = str(data.get('Age', '')).split('.')[0]
            self.patient_entries['age'].insert(0, age_val)
        if 'sex' in self.patient_entries:
            self.patient_entries['sex'].insert(0, data.get('Gender', ''))
        if 'doctor' in self.patient_entries:
            self.patient_entries['doctor'].insert(0, data.get('Ref By', ''))
        if 'patient_id' in self.patient_entries:
            self.patient_entries['patient_id'].insert(0, data.get('Bill Number', ''))

    def clear_patient_fields(self):
        """Clears only the patient-related entry fields."""
        for key, entry in self.patient_entries.items():
            # Keep date fields, clear everything else
            if "date" not in key:
                entry.delete(0, "end")

    def _autocomplete_date(self, event):
        """Autocompletes a date field when focus leaves the widget."""
        widget = event.widget
        day_str = widget.get().strip()

        # Check if the input is a 1 or 2 digit number (a day)
        if day_str.isdigit() and len(day_str) <= 2:
            try:
                day = int(day_str)
                if 1 <= day <= 31:
                    # Get the current month and year
                    now = datetime.now()
                    # Format with leading zeros for day and month
                    full_date = f"{day:02d}.{now.month:02d}.{now.year}"
                    
                    # Update the entry field with the full date
                    widget.delete(0, "end")
                    widget.insert(0, full_date)
            except ValueError:
                # Ignore if it's not a valid number
                pass

class BSGTFormFrame(ctk.CTkFrame):
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller
        
        self.patient_entries = {}
        self.test_entries = {}

        # --- Top Bar with Back Button ---
        top_bar = ctk.CTkFrame(self, fg_color="transparent")
        top_bar.pack(fill="x", pady=(6, 12), padx=6)
        
        ctk.CTkButton(top_bar, text=" ⬅️ Back to List", width=120,
                      command=lambda: controller.show_frame(ReportSelectionFrame)).pack(side="left", padx=6, pady=6)
        ctk.CTkLabel(top_bar, text="Bilirubin, SGOT, SGPT Report Entry", font=ctk.CTkFont(size=18, weight="bold")).pack(side="left", padx=8)

        # --- Main Content Area ---
        content = ctk.CTkFrame(self)
        content.pack(expand=True, fill="both", padx=6, pady=6)
        left = ctk.CTkFrame(content)
        left.pack(side="left", expand=True, fill="both", padx=(0,8), pady=8)
        right = ctk.CTkFrame(content)
        right.pack(side="right", expand=True, fill="both", padx=(8,0), pady=8)

        # --- Patient Details Fields ---
        self.add_patient_field(left, 0, "Patient Name", "name")
        self.add_patient_field(left, 1, "Age", "age")
        self.add_patient_field(left, 2, "Sex", "sex")
        self.add_patient_field(left, 3, "Ref. by Dr.", "doctor")
        self.add_patient_field(left, 4, "Patient ID", "patient_id")
        self.add_patient_field(left, 5, "Collection Date", "collection_date")
        self.add_patient_field(left, 6, "Reporting Date", "report_date")
        
        # --- Test Details Fields ---
        self.add_test_field(right, 0, "TOTAL BILIRUBIN", "total_bili")
        self.add_test_field(right, 1, "DIRECT BILIRUBIN", "direct_bili")
        self.add_test_field(right, 2, "SGPT", "sgpt")
        self.add_test_field(right, 3, "SGOT", "sgot")
        
        ctk.CTkButton(self, text="Save Report", height=40, command=self.generate_report).pack(pady=12, padx=6)

    def add_patient_field(self, parent, row, label, key):
        ctk.CTkLabel(parent, text=label, anchor="w").grid(row=row, column=0, padx=5, pady=4, sticky="w")
        e = ctk.CTkEntry(parent, width=220)
        e.grid(row=row, column=1, padx=5, pady=4)
        self.patient_entries[key] = e
        # This line activates the autocomplete function for date fields
        if "date" in key.lower():
            e.bind("<FocusOut>", self._autocomplete_date)

    def add_test_field(self, parent, row, label, key):
        ctk.CTkLabel(parent, text=label, anchor="w").grid(row=row, column=0, padx=5, pady=4, sticky="w")
        e = ctk.CTkEntry(parent, width=220)
        e.grid(row=row, column=1, padx=5, pady=4)
        self.test_entries[key] = e

    def clear_all_fields(self):
        for e in list(self.patient_entries.values()) + list(self.test_entries.values()):
            e.delete(0, "end")

    def generate_report(self):
        def get_float(key):
            try:
                return float(self.test_entries[key].get())
            except (ValueError, KeyError):
                return 0.0

        context = {key: entry.get() for key, entry in self.patient_entries.items()}
        context.update({key: entry.get() for key, entry in self.test_entries.items()})
        
        # --- Perform Calculation ---
        total_bili = get_float("total_bili")
        direct_bili = get_float("direct_bili")
        indirect_bili = total_bili - direct_bili
        context['indirect_bili'] = f"{indirect_bili:.2f}"

        try:
            template_name = "BILIRUBIN, SGOT, SGPT  NEW.docx"
            full_template_path = os.path.join(REPORT_TEMPLATES_FOLDER, template_name)

            doc = DocxTemplate(full_template_path)
            doc.render(context)
            
            target_dir = os.path.join(OUTPUT_FOLDER, "BILIRUBIN, SGOT, SGPT  NEW")
            os.makedirs(target_dir, exist_ok=True)
            
            patient_id = context.get("patient_id", "NoID").strip().replace(' ', '_') or "NoID"
            base_fname = f"BILIRUBIN, SGOT, SGPT  NEW_{patient_id}.docx"
            out_path = os.path.join(target_dir, base_fname)
            
            counter = 1
            while os.path.exists(out_path):
                new_fname = f"BILIRUBIN, SGOT, SGPT  NEW_{patient_id}({counter}).docx"
                out_path = os.path.join(target_dir, new_fname)
                counter += 1
            
            doc.save(out_path)
            
            messagebox.showinfo("Success", f"Report saved:\n{out_path}", parent=self)
            self.clear_all_fields()
            
            os.startfile(out_path)
            self.controller.iconify()

        except FileNotFoundError:
            messagebox.showerror(
                "Template Not Found",
                f"The template file '{template_name}' was not found.\n\nPlease make sure it is inside the '{REPORT_TEMPLATES_FOLDER}' folder.",
                parent=self
            )
        except Exception as e:
            messagebox.showerror("Error", f"Failed to generate report:\n{e}", parent=self)

    def populate_patient_data(self, data):
        """Receives patient data and fills the form's entry fields."""
        # Clear any old data first, except for date fields
        for key, entry in self.patient_entries.items():
            if "date" not in key:
                entry.delete(0, "end")
        
        # Insert new data using keys from the Excel file
        # The .get() method prevents errors if a key is missing
        if 'name' in self.patient_entries:
            self.patient_entries['name'].insert(0, data.get('Patient Name', ''))
        if 'age' in self.patient_entries:
            # Format age to remove .0 if it exists
            age_val = str(data.get('Age', '')).split('.')[0]
            self.patient_entries['age'].insert(0, age_val)
        if 'sex' in self.patient_entries:
            self.patient_entries['sex'].insert(0, data.get('Gender', ''))
        if 'doctor' in self.patient_entries:
            self.patient_entries['doctor'].insert(0, data.get('Ref By', ''))
        if 'patient_id' in self.patient_entries:
            self.patient_entries['patient_id'].insert(0, data.get('Bill Number', ''))

    def clear_patient_fields(self):
        """Clears only the patient-related entry fields."""
        for key, entry in self.patient_entries.items():
            # Keep date fields, clear everything else
            if "date" not in key:
                entry.delete(0, "end")

    def _autocomplete_date(self, event):
        """Autocompletes a date field when focus leaves the widget."""
        widget = event.widget
        day_str = widget.get().strip()

        # Check if the input is a 1 or 2 digit number (a day)
        if day_str.isdigit() and len(day_str) <= 2:
            try:
                day = int(day_str)
                if 1 <= day <= 31:
                    # Get the current month and year
                    now = datetime.now()
                    # Format with leading zeros for day and month
                    full_date = f"{day:02d}.{now.month:02d}.{now.year}"
                    
                    # Update the entry field with the full date
                    widget.delete(0, "end")
                    widget.insert(0, full_date)
            except ValueError:
                # Ignore if it's not a valid number
                pass

class SugerFBSFormFrame(ctk.CTkFrame):
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller
        
        self.patient_entries = {}
        self.test_entries = {}

        # --- Top Bar with Back Button ---
        top_bar = ctk.CTkFrame(self, fg_color="transparent")
        top_bar.pack(fill="x", pady=(6, 12), padx=6)
        
        ctk.CTkButton(top_bar, text=" ⬅️ Back to List", width=120,
                      command=lambda: controller.show_frame(ReportSelectionFrame)).pack(side="left", padx=6, pady=6)
        ctk.CTkLabel(top_bar, text="Suger FBS Report Entry", font=ctk.CTkFont(size=18, weight="bold")).pack(side="left", padx=8)

        # --- Main Content Area ---
        content = ctk.CTkFrame(self)
        content.pack(expand=True, fill="both", padx=6, pady=6)
        left = ctk.CTkFrame(content)
        left.pack(side="left", expand=True, fill="both", padx=(0,8), pady=8)
        right = ctk.CTkFrame(content)
        right.pack(side="right", expand=True, fill="both", padx=(8,0), pady=8)

        # --- Patient Details Fields ---
        self.add_patient_field(left, 0, "Patient Name", "name")
        self.add_patient_field(left, 1, "Age", "age")
        self.add_patient_field(left, 2, "Sex", "sex")
        self.add_patient_field(left, 3, "Ref. by Dr.", "doctor")
        self.add_patient_field(left, 4, "Patient ID", "patient_id")
        self.add_patient_field(left, 5, "Collection Date", "collection_date")
        self.add_patient_field(left, 6, "Reporting Date", "report_date")
        
        # --- Test Details Fields ---
        self.add_test_field(right, 0, "GLUCOSE FASTING (PLASMA)", "glucose_f")
        
        ctk.CTkButton(self, text="Save Report", height=40, command=self.generate_report).pack(pady=12, padx=6)

    def add_patient_field(self, parent, row, label, key):
        ctk.CTkLabel(parent, text=label, anchor="w").grid(row=row, column=0, padx=5, pady=4, sticky="w")
        e = ctk.CTkEntry(parent, width=220)
        e.grid(row=row, column=1, padx=5, pady=4)
        self.patient_entries[key] = e
        # This line activates the autocomplete function for date fields
        if "date" in key.lower():
            e.bind("<FocusOut>", self._autocomplete_date)

    def add_test_field(self, parent, row, label, key):
        ctk.CTkLabel(parent, text=label, anchor="w").grid(row=row, column=0, padx=5, pady=4, sticky="w")
        e = ctk.CTkEntry(parent, width=220)
        e.grid(row=row, column=1, padx=5, pady=4)
        self.test_entries[key] = e

    def clear_all_fields(self):
        for e in list(self.patient_entries.values()) + list(self.test_entries.values()):
            e.delete(0, "end")

    def generate_report(self):
        context = {key: entry.get() for key, entry in self.patient_entries.items()}
        context.update({key: entry.get() for key, entry in self.test_entries.items()})

        try:
            template_name = "SUGAR FBS  NEW.docx"
            full_template_path = os.path.join(REPORT_TEMPLATES_FOLDER, template_name)

            doc = DocxTemplate(full_template_path)
            doc.render(context)
            
            target_dir = os.path.join(OUTPUT_FOLDER, "SugerFBS")
            os.makedirs(target_dir, exist_ok=True)
            
            patient_id = context.get("patient_id", "NoID").strip().replace(' ', '_') or "NoID"
            base_fname = f"Suger_FBS_{patient_id}.docx"
            out_path = os.path.join(target_dir, base_fname)
            
            counter = 1
            while os.path.exists(out_path):
                new_fname = f"Suger_FBS_{patient_id}({counter}).docx"
                out_path = os.path.join(target_dir, new_fname)
                counter += 1
            
            doc.save(out_path)
            
            messagebox.showinfo("Success", f"Report saved:\n{out_path}", parent=self)
            self.clear_all_fields()
            
            os.startfile(out_path)
            self.controller.iconify()

        except FileNotFoundError:
            messagebox.showerror(
                "Template Not Found",
                f"The template file '{template_name}' was not found.\n\nPlease make sure it is inside the '{REPORT_TEMPLATES_FOLDER}' folder.",
                parent=self
            )
        except Exception as e:
            messagebox.showerror("Error", f"Failed to generate report:\n{e}", parent=self)

    def populate_patient_data(self, data):
        """Receives patient data and fills the form's entry fields."""
        # Clear any old data first, except for date fields
        for key, entry in self.patient_entries.items():
            if "date" not in key:
                entry.delete(0, "end")
        
        # Insert new data using keys from the Excel file
        # The .get() method prevents errors if a key is missing
        if 'name' in self.patient_entries:
            self.patient_entries['name'].insert(0, data.get('Patient Name', ''))
        if 'age' in self.patient_entries:
            # Format age to remove .0 if it exists
            age_val = str(data.get('Age', '')).split('.')[0]
            self.patient_entries['age'].insert(0, age_val)
        if 'sex' in self.patient_entries:
            self.patient_entries['sex'].insert(0, data.get('Gender', ''))
        if 'doctor' in self.patient_entries:
            self.patient_entries['doctor'].insert(0, data.get('Ref By', ''))
        if 'patient_id' in self.patient_entries:
            self.patient_entries['patient_id'].insert(0, data.get('Bill Number', ''))

    def clear_patient_fields(self):
        """Clears only the patient-related entry fields."""
        for key, entry in self.patient_entries.items():
            # Keep date fields, clear everything else
            if "date" not in key:
                entry.delete(0, "end")

    def _autocomplete_date(self, event):
        """Autocompletes a date field when focus leaves the widget."""
        widget = event.widget
        day_str = widget.get().strip()

        # Check if the input is a 1 or 2 digit number (a day)
        if day_str.isdigit() and len(day_str) <= 2:
            try:
                day = int(day_str)
                if 1 <= day <= 31:
                    # Get the current month and year
                    now = datetime.now()
                    # Format with leading zeros for day and month
                    full_date = f"{day:02d}.{now.month:02d}.{now.year}"
                    
                    # Update the entry field with the full date
                    widget.delete(0, "end")
                    widget.insert(0, full_date)
            except ValueError:
                # Ignore if it's not a valid number
                pass

class SugerFBS_PPBSFormFrame(ctk.CTkFrame):
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller
        
        self.patient_entries = {}
        self.test_entries = {}

        # --- Top Bar with Back Button ---
        top_bar = ctk.CTkFrame(self, fg_color="transparent")
        top_bar.pack(fill="x", pady=(6, 12), padx=6)
        
        ctk.CTkButton(top_bar, text=" ⬅️ Back to List", width=120,
                      command=lambda: controller.show_frame(ReportSelectionFrame)).pack(side="left", padx=6, pady=6)
        ctk.CTkLabel(top_bar, text="Suger FBS & PPBS Report Entry", font=ctk.CTkFont(size=18, weight="bold")).pack(side="left", padx=8)

        # --- Main Content Area ---
        content = ctk.CTkFrame(self)
        content.pack(expand=True, fill="both", padx=6, pady=6)
        left = ctk.CTkFrame(content)
        left.pack(side="left", expand=True, fill="both", padx=(0,8), pady=8)
        right = ctk.CTkFrame(content)
        right.pack(side="right", expand=True, fill="both", padx=(8,0), pady=8)

        # --- Patient Details Fields ---
        self.add_patient_field(left, 0, "Patient Name", "name")
        self.add_patient_field(left, 1, "Age", "age")
        self.add_patient_field(left, 2, "Sex", "sex")
        self.add_patient_field(left, 3, "Ref. by Dr.", "doctor")
        self.add_patient_field(left, 4, "Patient ID", "patient_id")
        self.add_patient_field(left, 5, "Collection Date", "collection_date")
        self.add_patient_field(left, 6, "Reporting Date", "report_date")
        
        # --- Test Details Fields ---
        self.add_test_field(right, 0, "GLUCOSE FASTING (PLASMA)", "glucose_f")
        self.add_test_field(right, 1, "GLUCOSE POSTPRANDIAL (PLASMA)", "glucose_pp")
        
        ctk.CTkButton(self, text="Save Report", height=40, command=self.generate_report).pack(pady=12, padx=6)

    def add_patient_field(self, parent, row, label, key):
        ctk.CTkLabel(parent, text=label, anchor="w").grid(row=row, column=0, padx=5, pady=4, sticky="w")
        e = ctk.CTkEntry(parent, width=220)
        e.grid(row=row, column=1, padx=5, pady=4)
        self.patient_entries[key] = e
        # This line activates the autocomplete function for date fields
        if "date" in key.lower():
            e.bind("<FocusOut>", self._autocomplete_date)

    def add_test_field(self, parent, row, label, key):
        ctk.CTkLabel(parent, text=label, anchor="w").grid(row=row, column=0, padx=5, pady=4, sticky="w")
        e = ctk.CTkEntry(parent, width=220)
        e.grid(row=row, column=1, padx=5, pady=4)
        self.test_entries[key] = e

    def clear_all_fields(self):
        for e in list(self.patient_entries.values()) + list(self.test_entries.values()):
            e.delete(0, "end")

    def generate_report(self):
        context = {key: entry.get() for key, entry in self.patient_entries.items()}
        context.update({key: entry.get() for key, entry in self.test_entries.items()})

        try:
            template_name = "SUGAR FBS, PPBS NEW.docx"
            full_template_path = os.path.join(REPORT_TEMPLATES_FOLDER, template_name)

            doc = DocxTemplate(full_template_path)
            doc.render(context)
            
            target_dir = os.path.join(OUTPUT_FOLDER, "SugerFBS_PPBS")
            os.makedirs(target_dir, exist_ok=True)
            
            patient_id = context.get("patient_id", "NoID").strip().replace(' ', '_') or "NoID"
            base_fname = f"Suger_FBS_PPBS_{patient_id}.docx"
            out_path = os.path.join(target_dir, base_fname)
            
            counter = 1
            while os.path.exists(out_path):
                new_fname = f"Suger_FBS_PPBS_{patient_id}({counter}).docx"
                out_path = os.path.join(target_dir, new_fname)
                counter += 1
            
            doc.save(out_path)
            
            messagebox.showinfo("Success", f"Report saved:\n{out_path}", parent=self)
            self.clear_all_fields()
            
            os.startfile(out_path)
            self.controller.iconify()

        except FileNotFoundError:
            messagebox.showerror(
                "Template Not Found",
                f"The template file '{template_name}' was not found.\n\nPlease make sure it is inside the '{REPORT_TEMPLATES_FOLDER}' folder.",
                parent=self
            )
        except Exception as e:
            messagebox.showerror("Error", f"Failed to generate report:\n{e}", parent=self)
    def populate_patient_data(self, data):
        """Receives patient data and fills the form's entry fields."""
        # Clear any old data first, except for date fields
        for key, entry in self.patient_entries.items():
            if "date" not in key:
                entry.delete(0, "end")
        
        # Insert new data using keys from the Excel file
        # The .get() method prevents errors if a key is missing
        if 'name' in self.patient_entries:
            self.patient_entries['name'].insert(0, data.get('Patient Name', ''))
        if 'age' in self.patient_entries:
            # Format age to remove .0 if it exists
            age_val = str(data.get('Age', '')).split('.')[0]
            self.patient_entries['age'].insert(0, age_val)
        if 'sex' in self.patient_entries:
            self.patient_entries['sex'].insert(0, data.get('Gender', ''))
        if 'doctor' in self.patient_entries:
            self.patient_entries['doctor'].insert(0, data.get('Ref By', ''))
        if 'patient_id' in self.patient_entries:
            self.patient_entries['patient_id'].insert(0, data.get('Bill Number', ''))

    def clear_patient_fields(self):
        """Clears only the patient-related entry fields."""
        for key, entry in self.patient_entries.items():
            # Keep date fields, clear everything else
            if "date" not in key:
                entry.delete(0, "end")

    def _autocomplete_date(self, event):
        """Autocompletes a date field when focus leaves the widget."""
        widget = event.widget
        day_str = widget.get().strip()

        # Check if the input is a 1 or 2 digit number (a day)
        if day_str.isdigit() and len(day_str) <= 2:
            try:
                day = int(day_str)
                if 1 <= day <= 31:
                    # Get the current month and year
                    now = datetime.now()
                    # Format with leading zeros for day and month
                    full_date = f"{day:02d}.{now.month:02d}.{now.year}"
                    
                    # Update the entry field with the full date
                    widget.delete(0, "end")
                    widget.insert(0, full_date)
            except ValueError:
                # Ignore if it's not a valid number
                pass

class SodiumPotassiumFormFrame(ctk.CTkFrame):
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller
        
        self.patient_entries = {}
        self.test_entries = {}

        # --- Top Bar with Back Button ---
        top_bar = ctk.CTkFrame(self, fg_color="transparent")
        top_bar.pack(fill="x", pady=(6, 12), padx=6)
        
        ctk.CTkButton(top_bar, text=" ⬅️ Back to List", width=120,
                      command=lambda: controller.show_frame(ReportSelectionFrame)).pack(side="left", padx=6, pady=6)
        ctk.CTkLabel(top_bar, text="Sodium & Potassium Report Entry", font=ctk.CTkFont(size=18, weight="bold")).pack(side="left", padx=8)

        # --- Main Content Area ---
        content = ctk.CTkFrame(self)
        content.pack(expand=True, fill="both", padx=6, pady=6)
        left = ctk.CTkFrame(content)
        left.pack(side="left", expand=True, fill="both", padx=(0,8), pady=8)
        right = ctk.CTkFrame(content)
        right.pack(side="right", expand=True, fill="both", padx=(8,0), pady=8)

        # --- Patient Details Fields ---
        self.add_patient_field(left, 0, "Patient Name", "name")
        self.add_patient_field(left, 1, "Age", "age")
        self.add_patient_field(left, 2, "Sex", "sex")
        self.add_patient_field(left, 3, "Ref. by Dr.", "doctor")
        self.add_patient_field(left, 4, "Patient ID", "patient_id")
        self.add_patient_field(left, 5, "Collection Date", "collection_date")
        self.add_patient_field(left, 6, "Reporting Date", "report_date")
        
        # --- Test Details Fields ---
        self.add_test_field(right, 0, "SERUM SODIUM", "sodium")
        self.add_test_field(right, 1, "SERUM POTASSIUM", "potassium")
        
        ctk.CTkButton(self, text="Save Report", height=40, command=self.generate_report).pack(pady=12, padx=6)

    def add_patient_field(self, parent, row, label, key):
        ctk.CTkLabel(parent, text=label, anchor="w").grid(row=row, column=0, padx=5, pady=4, sticky="w")
        e = ctk.CTkEntry(parent, width=220)
        e.grid(row=row, column=1, padx=5, pady=4)
        self.patient_entries[key] = e
        # This line activates the autocomplete function for date fields
        if "date" in key.lower():
            e.bind("<FocusOut>", self._autocomplete_date)

    def add_test_field(self, parent, row, label, key):
        ctk.CTkLabel(parent, text=label, anchor="w").grid(row=row, column=0, padx=5, pady=4, sticky="w")
        e = ctk.CTkEntry(parent, width=220)
        e.grid(row=row, column=1, padx=5, pady=4)
        self.test_entries[key] = e

    def clear_all_fields(self):
        for e in list(self.patient_entries.values()) + list(self.test_entries.values()):
            e.delete(0, "end")

    def generate_report(self):
        context = {key: entry.get() for key, entry in self.patient_entries.items()}
        context.update({key: entry.get() for key, entry in self.test_entries.items()})

        try:
            template_name = "SODIUM POTASSIUM NEW.docx"
            full_template_path = os.path.join(REPORT_TEMPLATES_FOLDER, template_name)

            doc = DocxTemplate(full_template_path)
            doc.render(context)
            
            target_dir = os.path.join(OUTPUT_FOLDER, "SodiumPotassium")
            os.makedirs(target_dir, exist_ok=True)
            
            patient_id = context.get("patient_id", "NoID").strip().replace(' ', '_') or "NoID"
            base_fname = f"Sodium_Potassium_{patient_id}.docx"
            out_path = os.path.join(target_dir, base_fname)
            
            counter = 1
            while os.path.exists(out_path):
                new_fname = f"Sodium_Potassium_{patient_id}({counter}).docx"
                out_path = os.path.join(target_dir, new_fname)
                counter += 1
            
            doc.save(out_path)
            
            messagebox.showinfo("Success", f"Report saved:\n{out_path}", parent=self)
            self.clear_all_fields()
            
            os.startfile(out_path)
            self.controller.iconify()

        except FileNotFoundError:
            messagebox.showerror(
                "Template Not Found",
                f"The template file '{template_name}' was not found.\n\nPlease make sure it is inside the '{REPORT_TEMPLATES_FOLDER}' folder.",
                parent=self
            )
        except Exception as e:
            messagebox.showerror("Error", f"Failed to generate report:\n{e}", parent=self)

    def populate_patient_data(self, data):
        """Receives patient data and fills the form's entry fields."""
        # Clear any old data first, except for date fields
        for key, entry in self.patient_entries.items():
            if "date" not in key:
                entry.delete(0, "end")
        
        # Insert new data using keys from the Excel file
        # The .get() method prevents errors if a key is missing
        if 'name' in self.patient_entries:
            self.patient_entries['name'].insert(0, data.get('Patient Name', ''))
        if 'age' in self.patient_entries:
            # Format age to remove .0 if it exists
            age_val = str(data.get('Age', '')).split('.')[0]
            self.patient_entries['age'].insert(0, age_val)
        if 'sex' in self.patient_entries:
            self.patient_entries['sex'].insert(0, data.get('Gender', ''))
        if 'doctor' in self.patient_entries:
            self.patient_entries['doctor'].insert(0, data.get('Ref By', ''))
        if 'patient_id' in self.patient_entries:
            self.patient_entries['patient_id'].insert(0, data.get('Bill Number', ''))

    def clear_patient_fields(self):
        """Clears only the patient-related entry fields."""
        for key, entry in self.patient_entries.items():
            # Keep date fields, clear everything else
            if "date" not in key:
                entry.delete(0, "end")


    def _autocomplete_date(self, event):
        """Autocompletes a date field when focus leaves the widget."""
        widget = event.widget
        day_str = widget.get().strip()

        # Check if the input is a 1 or 2 digit number (a day)
        if day_str.isdigit() and len(day_str) <= 2:
            try:
                day = int(day_str)
                if 1 <= day <= 31:
                    # Get the current month and year
                    now = datetime.now()
                    # Format with leading zeros for day and month
                    full_date = f"{day:02d}.{now.month:02d}.{now.year}"
                    
                    # Update the entry field with the full date
                    widget.delete(0, "end")
                    widget.insert(0, full_date)
            except ValueError:
                # Ignore if it's not a valid number
                pass

class RBS_Urea_CreatinineFormFrame(ctk.CTkFrame):
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller
        
        self.patient_entries = {}
        self.test_entries = {}

        # --- Top Bar with Back Button ---
        top_bar = ctk.CTkFrame(self, fg_color="transparent")
        top_bar.pack(fill="x", pady=(6, 12), padx=6)
        
        ctk.CTkButton(top_bar, text=" ⬅️ Back to List", width=120,
                      command=lambda: controller.show_frame(ReportSelectionFrame)).pack(side="left", padx=6, pady=6)
        ctk.CTkLabel(top_bar, text="RBS & RFT Report Entry", font=ctk.CTkFont(size=18, weight="bold")).pack(side="left", padx=8)

        # --- Main Content Area ---
        content = ctk.CTkFrame(self)
        content.pack(expand=True, fill="both", padx=6, pady=6)
        left = ctk.CTkFrame(content)
        left.pack(side="left", expand=True, fill="both", padx=(0,8), pady=8)
        right = ctk.CTkScrollableFrame(content)
        right.pack(side="right", expand=True, fill="both", padx=(8,0), pady=8)

        # --- Patient Details Fields ---
        self.add_patient_field(left, 0, "Patient Name", "name")
        self.add_patient_field(left, 1, "Age", "age")
        self.add_patient_field(left, 2, "Sex", "sex")
        self.add_patient_field(left, 3, "Ref. by Dr.", "doctor")
        self.add_patient_field(left, 4, "Patient ID", "patient_id")
        self.add_patient_field(left, 5, "Collection Date", "collection_date")
        self.add_patient_field(left, 6, "Reporting Date", "report_date")
        
        # --- Test Details Fields ---
        self.add_test_field(right, 0, "SUGAR (R)", "sugar_r")
        self.add_test_field(right, 1, "UREA", "urea")
        self.add_test_field(right, 2, "CREATININE", "creatinine")
        self.add_test_field(right, 3, "URIC ACID", "uric_acid")
        self.add_test_field(right, 4, "SODIUM (Na+)", "sodium")
        self.add_test_field(right, 5, "POTASSIUM (K+)", "potassium")
        self.add_test_field(right, 6, "CALCIUM (Ca++)", "calcium")
        
        ctk.CTkButton(self, text="Save Report", height=40, command=self.generate_report).pack(pady=12, padx=6)

    def add_patient_field(self, parent, row, label, key):
        ctk.CTkLabel(parent, text=label, anchor="w").grid(row=row, column=0, padx=5, pady=4, sticky="w")
        e = ctk.CTkEntry(parent, width=220)
        e.grid(row=row, column=1, padx=5, pady=4)
        self.patient_entries[key] = e
        # This line activates the autocomplete function for date fields
        if "date" in key.lower():
            e.bind("<FocusOut>", self._autocomplete_date)

    def add_test_field(self, parent, row, label, key):
        ctk.CTkLabel(parent, text=label, anchor="w").grid(row=row, column=0, padx=5, pady=4, sticky="w")
        e = ctk.CTkEntry(parent, width=220)
        e.grid(row=row, column=1, padx=5, pady=4)
        self.test_entries[key] = e

    def clear_all_fields(self):
        for e in list(self.patient_entries.values()) + list(self.test_entries.values()):
            e.delete(0, "end")

    def generate_report(self):
        context = {key: entry.get() for key, entry in self.patient_entries.items()}
        context.update({key: entry.get() for key, entry in self.test_entries.items()})

        try:
            template_name = "RBS, UREA, CREA URIC NEW.docx"
            full_template_path = os.path.join(REPORT_TEMPLATES_FOLDER, template_name)

            doc = DocxTemplate(full_template_path)
            doc.render(context)
            
            target_dir = os.path.join(OUTPUT_FOLDER, "RBS, UREA, CREA URIC NEW")
            os.makedirs(target_dir, exist_ok=True)
            
            patient_id = context.get("patient_id", "NoID").strip().replace(' ', '_') or "NoID"
            base_fname = f"RBS, UREA, CREA URIC NEW_{patient_id}.docx"
            out_path = os.path.join(target_dir, base_fname)
            
            counter = 1
            while os.path.exists(out_path):
                new_fname = f"RBS, UREA, CREA URIC NEW_{patient_id}({counter}).docx"
                out_path = os.path.join(target_dir, new_fname)
                counter += 1
            
            doc.save(out_path)
            
            messagebox.showinfo("Success", f"Report saved:\n{out_path}", parent=self)
            self.clear_all_fields()
            
            os.startfile(out_path)
            self.controller.iconify()

        except FileNotFoundError:
            messagebox.showerror(
                "Template Not Found",
                f"The template file '{template_name}' was not found.\n\nPlease make sure it is inside the '{REPORT_TEMPLATES_FOLDER}' folder.",
                parent=self
            )
        except Exception as e:
            messagebox.showerror("Error", f"Failed to generate report:\n{e}", parent=self)
    def populate_patient_data(self, data):
        """Receives patient data and fills the form's entry fields."""
        # Clear any old data first, except for date fields
        for key, entry in self.patient_entries.items():
            if "date" not in key:
                entry.delete(0, "end")
        
        # Insert new data using keys from the Excel file
        # The .get() method prevents errors if a key is missing
        if 'name' in self.patient_entries:
            self.patient_entries['name'].insert(0, data.get('Patient Name', ''))
        if 'age' in self.patient_entries:
            # Format age to remove .0 if it exists
            age_val = str(data.get('Age', '')).split('.')[0]
            self.patient_entries['age'].insert(0, age_val)
        if 'sex' in self.patient_entries:
            self.patient_entries['sex'].insert(0, data.get('Gender', ''))
        if 'doctor' in self.patient_entries:
            self.patient_entries['doctor'].insert(0, data.get('Ref By', ''))
        if 'patient_id' in self.patient_entries:
            self.patient_entries['patient_id'].insert(0, data.get('Bill Number', ''))
    
    def clear_patient_fields(self):
        """Clears only the patient-related entry fields."""
        for key, entry in self.patient_entries.items():
            # Keep date fields, clear everything else
            if "date" not in key:
                entry.delete(0, "end")

    def _autocomplete_date(self, event):
        """Autocompletes a date field when focus leaves the widget."""
        widget = event.widget
        day_str = widget.get().strip()

        # Check if the input is a 1 or 2 digit number (a day)
        if day_str.isdigit() and len(day_str) <= 2:
            try:
                day = int(day_str)
                if 1 <= day <= 31:
                    # Get the current month and year
                    now = datetime.now()
                    # Format with leading zeros for day and month
                    full_date = f"{day:02d}.{now.month:02d}.{now.year}"
                    
                    # Update the entry field with the full date
                    widget.delete(0, "end")
                    widget.insert(0, full_date)
            except ValueError:
                # Ignore if it's not a valid number
                pass

class BillingApp(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("GREENPATH Diagnostic")
        self.geometry("1200x800")
        self.minsize(1100, 700)
        # Open in maximized mode
        self.after(100, lambda: self.state("zoomed"))
        
        # Configure grid layout
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)

        # Initialize variables
        self.ref_by_var = ctk.StringVar()
        self.address_var = ctk.StringVar()
        self.contact_var = ctk.StringVar()
        self.client_name_var = ctk.StringVar()
        self.age_var = ctk.StringVar()
        self.gender_var = ctk.StringVar()
        self.advanced_var = ctk.StringVar()
        self.test_name_var = ctk.StringVar()
        self.rate_var = ctk.StringVar()
        self.agent_var = ctk.StringVar()
        self.discount_var = ctk.StringVar()
        # --- NEW --- Add traces to automatically call the update function
        self.advanced_var.trace_add("write", self.update_totals_display)
        self.discount_var.trace_add("write", self.update_totals_display)

        self.tests = []
        self.test_data = self.load_test_data()
        self.doctors_data = self.load_doctors_data() #.....NEWWWWW

        self.agent_var = ctk.StringVar()
        self.agents_data = self.load_agents_data()  # Load agents list
        
        # Create UI components
        self.create_sidebar()
        self.create_main_content()
        self.bind_events()

        self.bind("<Button-1>", lambda e: self.clear_suggestions())

        # Initialize counters and folders
        os.makedirs(BILL_FOLDER, exist_ok=True)
        self._fg_color = ctk.ThemeManager.theme["CTkFrame"]["fg_color"]
        self._text_color = ctk.ThemeManager.theme["CTkLabel"]["text_color"]

        self.settings_visible = False
        self.rotation_angle = 0
        self.create_settings_icon()

       
    def show_dashboard(self):
        """Hides manage page and shows dashboard"""
        self.manage_tests_page.grid_forget()
        self.manage_doctors_page.grid_forget() # <<< ADD THIS
        self.manage_agents_page.grid_forget()
        self.history_page.grid_forget()
        self.dashboard_frame.grid(row=0, column=0, sticky="nsew", padx=20, pady=20)

    def show_manage_tests_page(self):
        """Hides dashboard and shows manage page"""
        self.dashboard_frame.grid_forget()
        self.manage_tests_page.grid(row=0, column=0, sticky="nsew")
        self.manage_tests_page.refresh_list() # Ensure list is fresh

    def show_manage_doctors_page(self):
        """Hides dashboard and shows manage doctors page"""
        self.dashboard_frame.grid_forget()
        # Also ensure other manage pages are hidden if they exist
        if hasattr(self, 'manage_tests_page'): self.manage_tests_page.grid_forget()
        
        self.manage_doctors_page.grid(row=0, column=0, sticky="nsew")
        self.manage_doctors_page.refresh_list()

    def show_manage_agents_page(self):
        """Hides dashboard and shows manage agents page"""
        self.dashboard_frame.grid_forget()
        # Ensure other manage pages are hidden
        if hasattr(self, 'manage_tests_page'): self.manage_tests_page.grid_forget()
        if hasattr(self, 'manage_doctors_page'): self.manage_doctors_page.grid_forget()
        
        self.manage_agents_page.grid(row=0, column=0, sticky="nsew")
        self.manage_agents_page.refresh_list()

    def show_history_page(self):
        """Hides dashboard and shows history page"""
        self.dashboard_frame.grid_forget()
        if hasattr(self, 'manage_tests_page'): self.manage_tests_page.grid_forget()
        if hasattr(self, 'manage_doctors_page'): self.manage_doctors_page.grid_forget()
        if hasattr(self, 'manage_agents_page'): self.manage_agents_page.grid_forget()
        
        self.history_page.grid(row=0, column=0, sticky="nsew")
        self.history_page.load_history_data() # Refresh data on show

    def create_sidebar(self):
        self.sidebar_frame = ctk.CTkFrame(self, width=140, corner_radius=0)
        self.sidebar_frame.grid(row=0, column=0, rowspan=4, sticky="nsew")

        # Load and process the circular logo
        self.logo_image = ctk.CTkImage(self.create_circular_logo(LOGO_PATH, size=(80, 80)), size=(80, 80))

        # Display the logo
        self.logo_label = ctk.CTkLabel(self.sidebar_frame, image=self.logo_image, text="")
        self.logo_label.pack(pady=20)

        ctk.CTkLabel(self.sidebar_frame, text="GREENPATH\nDiagnostic",
                    font=ctk.CTkFont("Arial", size=20, weight="bold")).pack(pady=10)

        
        # Buttons
        ctk.CTkButton(self.sidebar_frame, text="Manage Tests", 
                      command=self.show_manage_tests_page).pack(pady=10, padx=20, fill="x")
        ctk.CTkButton(self.sidebar_frame, text="Manage Doctors", 
                      command=self.show_manage_doctors_page).pack(pady=10, padx=20, fill="x")
        ctk.CTkButton(self.sidebar_frame, text="Manage Agents", 
                      command=self.show_manage_agents_page).pack(pady=10, padx=20, fill="x")
        ctk.CTkButton(self.sidebar_frame, text="History", 
                      command=self.show_history_page).pack(pady=10, padx=20, fill="x")
        ctk.CTkButton(self.sidebar_frame, text="Report", command=self.show_report_options).pack(pady=10, padx=20, fill="x")

    def create_circular_logo(self, image_path, size=(80, 80)):
        """
        Creates a perfectly round logo.
        If the image file is not found, it creates a transparent placeholder.
        """
        try:
            # Attempt to open the image file from the provided path
            img = Image.open(image_path).convert("RGBA")
        except FileNotFoundError:
            # If the file doesn't exist, create a blank, transparent image instead
            print(f"Warning: Logo file not found at '{image_path}'. Creating a placeholder.")
            img = Image.new("RGBA", size, (0, 0, 0, 0)) # (R, G, B, Alpha) - fully transparent

        # Note: Using Image.Resampling.LANCZOS is the modern syntax for Image.LANCZOS
        img = img.resize(size, Image.Resampling.LANCZOS)

        # Create a circular mask
        mask = Image.new("L", size, 0)
        draw = ImageDraw.Draw(mask)
        draw.ellipse((0, 0, size[0], size[1]), fill=255)

        # Apply the circular mask to the image (or the placeholder)
        circular_img = Image.new("RGBA", size, (0, 0, 0, 0))
        circular_img.paste(img, (0, 0), mask)

        return circular_img
    

    

    
        
        

    def create_main_content(self):
        """Creates the main container and the dashboard view."""
        
        # 1. Create the Main Container (Holds both Dashboard and Manage Pages)
        # This frame sits in the main window grid
        self.main_container = ctk.CTkFrame(self, corner_radius=0, fg_color="transparent")
        self.main_container.grid(row=0, column=1, sticky="nsew")
        self.main_container.grid_rowconfigure(0, weight=1)
        self.main_container.grid_columnconfigure(0, weight=1)

        # 2. Create the Dashboard Frame (This holds your existing billing UI)
        self.dashboard_frame = ctk.CTkFrame(self.main_container, corner_radius=10)
        self.dashboard_frame.grid(row=0, column=0, sticky="nsew", padx=20, pady=20)
        
        # Configure dashboard grid
        self.dashboard_frame.grid_columnconfigure(0, weight=1)
        self.dashboard_frame.grid_rowconfigure(2, weight=1)

        # --- ATTACH EXISTING DASHBOARD COMPONENTS TO self.dashboard_frame ---
        
        # Pass the dashboard_frame as the parent to these helper methods
        self.create_patient_details(self.dashboard_frame) 
        self.create_test_entry(self.dashboard_frame)      

        # List and Totals Container (Now inside dashboard_frame)
        list_and_totals_container = ctk.CTkFrame(self.dashboard_frame, fg_color="transparent")
        list_and_totals_container.grid(row=2, column=0, columnspan=2, pady=10, padx=10, sticky="nsew")
        
        list_and_totals_container.grid_columnconfigure(0, weight=1)
        list_and_totals_container.grid_columnconfigure(1, weight=0)
        list_and_totals_container.grid_rowconfigure(0, weight=1)

        self.test_list_frame = ctk.CTkScrollableFrame(list_and_totals_container)
        self.test_list_frame.grid(row=0, column=0, pady=0, padx=(0, 10), sticky="nsew")

        self.create_totals_calculator(list_and_totals_container)

        # Print Button (Now inside dashboard_frame)
        self.print_invoice_button = ctk.CTkButton(
            self.dashboard_frame, text="Print Invoice", command=self.print_invoice,
            fg_color="#2E8B57", hover_color="#245c3d",
            font=ctk.CTkFont(size=14, weight="bold"), width=150, height=40
        )
        self.print_invoice_button.grid(row=3, column=0, columnspan=2, pady=10, padx=10, sticky="e")
        
        self.bind("<F5>", lambda e: self.print_invoice())
        self.print_invoice_button.bind("<FocusIn>", lambda e: self.print_invoice_button.configure(fg_color="#FFD700", text_color="black"))
        self.print_invoice_button.bind("<FocusOut>", lambda e: self.print_invoice_button.configure(fg_color="#2E8B57", text_color="white"))

        # 3. Create the Manage Tests Frame (Hidden by default)
        # We create it now but do NOT grid it yet. It sits in memory ready to be shown.
        self.manage_tests_page = ManageTestsFrame(self.main_container, self)
        self.manage_doctors_page = ManageDoctorsFrame(self.main_container, self)
        self.manage_agents_page = ManageAgentsFrame(self.main_container, self)
        self.history_page = HistoryFrame(self.main_container, self)


        # --- NEW --- Function to create the calculator UI
    def create_totals_calculator(self, parent_frame):
        """Creates a simplified totals calculator UI that only shows the gross total."""
        # The calculator is now created inside the parent_frame passed to it.
        totals_frame = ctk.CTkFrame(parent_frame)
        # Position the calculator on the right side (column 1) and stick to the top.
        totals_frame.grid(row=0, column=1, sticky="n", padx=(10, 0))
    
        # Center the content within the totals_frame
        totals_frame.grid_columnconfigure(0, weight=1)

        # --- MODIFIED: Only display the Gross Total ---
        ctk.CTkLabel(totals_frame, text="Total Amount", font=ctk.CTkFont(size=16, weight="bold")).grid(row=0, column=0, padx=10, pady=(10, 5))
    
        self.total_amount_label = ctk.CTkLabel(totals_frame, text="₹0.00", font=ctk.CTkFont(size=20))
        self.total_amount_label.grid(row=1, column=0, padx=10, pady=(5, 10))

        # --- Note: The labels below are created but not displayed. ---
        # This prevents the update_totals_display function from crashing,
        # as it still tries to configure these labels.
        self.discount_display_label = ctk.CTkLabel(totals_frame)
        self.advance_display_label = ctk.CTkLabel(totals_frame)
        self.due_amount_label = ctk.CTkLabel(totals_frame)

        # The update function is still called to update the visible total_amount_label
        self.update_totals_display()
    # --- NEW --- Function to perform all calculations and update the UI
    def update_totals_display(self, *args):
        try:
            total_amount = sum(test[1] for test in self.tests)
            discount_str = self.discount_var.get()
            discount_amount = float(discount_str) if discount_str else 0.0
            advance_str = self.advanced_var.get()
            advanced_payment = float(advance_str) if advance_str else 0.0

            net_total = total_amount - discount_amount
            due_amount = net_total - advanced_payment

            self.total_amount_label.configure(text=f"₹{total_amount:.2f}")
            self.discount_display_label.configure(text=f"₹{discount_amount:.2f}")
            self.advance_display_label.configure(text=f"₹{advanced_payment:.2f}")
            self.due_amount_label.configure(text=f"₹{due_amount:.2f}")

        except (ValueError, TypeError):
            self.due_amount_label.configure(text="Invalid Input")


    

    def create_patient_details(self, parent_frame):
        patient_frame = ctk.CTkFrame(parent_frame)
        patient_frame.grid(row=0, column=0, columnspan=2, pady=10, padx=10, sticky="nsew")
    
        # Create a list to store all entry widgets if needed later
        self.entry_widgets = {}
    
        # Fields now with (label, variable, name) structure
        fields = [
            ("Patient Name:", self.client_name_var, "client_name"),
            ("Age:", self.age_var, "age"),
            ("Gender:", self.gender_var, "gender"),
            ("Ref By Doctor:", self.ref_by_var, "ref_by"),
            ("Agent:", self.agent_var, "agent"),  # New Agent field
            ("Address:", self.address_var, "address"),
            ("Contact No:", self.contact_var, "contact")
        ]
    
        for idx, (label, var, name) in enumerate(fields):
            ctk.CTkLabel(patient_frame, text=label).grid(row=idx, column=0, padx=5, pady=5, sticky="e")
            entry = ctk.CTkEntry(patient_frame, textvariable=var, width=300)
            entry.grid(row=idx, column=1, padx=5, pady=5)
            self.entry_widgets[name] = entry  # Store reference to each entry
        
            # Add suggestion binding to Ref By Doctor field
            if name == "ref_by":
                entry.bind("<KeyRelease>", self.show_doctor_suggestions)
                self.ref_by_entry = entry  # Store direct reference for easy access

            elif name == "agent":  # Add this for agent suggestions
                entry.bind("<KeyRelease>", self.show_agent_suggestions)
                self.agent_entry = entry
        
        # After creating all entries in patient details
        gender_entry = self.entry_widgets["gender"]
        gender_entry.bind("<KeyPress>", self.handle_gender_shortcut)

    def handle_gender_shortcut(self, event):
        """Handles M/F keyboard shortcuts for gender field"""
        char = event.char.lower()
    
        if char == 'm':
            self.gender_var.set("Male")
            self.ref_by_entry.focus()
            return 'break'  # Prevent default key handling
        elif char == 'f':
            self.gender_var.set("Female")
            self.ref_by_entry.focus()
            return 'break'
        return None

    def show_agent_suggestions(self, event):
        # Destroy any existing popup first
        if hasattr(self, 'agent_suggestion_popup'):
            self.agent_suggestion_popup.destroy()

        search_term = self.agent_var.get().lower()
        if len(search_term) < 1:  # Only show after 2+ characters typed
            return

        matches = [agent for agent in self.agents_data if search_term in agent.lower()]
        if not matches:
            return

        x, y = self.agent_entry.winfo_rootx(), self.agent_entry.winfo_rooty() + self.agent_entry.winfo_height()
        self.agent_suggestion_popup = ctk.CTkToplevel(self)
        self.agent_suggestion_popup.overrideredirect(True)
        self.agent_suggestion_popup.geometry(f"+{x}+{y}")
        self.agent_suggestion_popup.lift()
        self.agent_suggestion_popup.attributes('-topmost', True)
        HOVER_BG = "#3a7ebf" if ctk.get_appearance_mode() == "Dark" else "#e0e0e0"
        HOVER_FG = "#ffffff" if ctk.get_appearance_mode() == "Dark" else "#000000"
        NORMAL_BG = "#2b2b2b" if ctk.get_appearance_mode() == "Dark" else "#ffffff"
        NORMAL_FG = "#ffffff" if ctk.get_appearance_mode() == "Dark" else "#000000"

        for doc in matches:
            frame = ctk.CTkFrame(self.agent_suggestion_popup, fg_color=NORMAL_BG)
            frame.pack(fill="x")
    
            lbl = ctk.CTkLabel(
                frame, 
                text=doc, 
                anchor="w", 
                cursor="hand2",
                text_color=NORMAL_FG
            )
            lbl.pack(fill="x", padx=5, pady=2)
        
            # Bind the click event to both the label and the frame
            for widget in [lbl, frame]:
                widget.bind("<Button-1>", lambda e, d=doc: self.select_agent_suggestion(d))
        
            # Hover effects
            lbl.bind("<Enter>", lambda e, l=lbl, f=frame: (
                f.configure(fg_color=HOVER_BG),
                l.configure(text_color=HOVER_FG)
            ))
            lbl.bind("<Leave>", lambda e, l=lbl, f=frame: (
                f.configure(fg_color=NORMAL_BG),
                l.configure(text_color=NORMAL_FG)
            ))

        # Bind click outside to close suggestions
        self.bind("<Button-1>", self.check_click_outside_agent_suggestions)
        self.bind("<Tab>", self.check_click_outside_agent_suggestions)

    def select_agent_suggestion(self, agent):
        """Handle selection of a doctor from the suggestions"""
        self.agent_var.set(agent)  # Set the value in the AGENT entry
        self.close_agent_suggestions()  # Close the popup
        self.agent_entry.focus()  # Return focus to the AGENT entry
        self.agent_entry.icursor("end")  # Move cursor to end of text

    def check_click_outside_agent_suggestions(self, event):
        """Check if click was outside the agent suggestion popup"""
        if not hasattr(self, 'agent_suggestion_popup') or not self.agent_suggestion_popup.winfo_exists():
            self.unbind("<Button-1>")
            return

        try:
            popup = self.agent_suggestion_popup
            popup_x = popup.winfo_rootx()
            popup_y = popup.winfo_rooty()
            popup_width = popup.winfo_width()
            popup_height = popup.winfo_height()
    
            # Check if click was outside the popup
            if not (popup_x <= event.x_root <= popup_x + popup_width and
                    popup_y <= event.y_root <= popup_y + popup_height):
                self.close_agent_suggestions()
        except:
            self.close_agent_suggestions()

    def close_agent_suggestions(self):
        """Clean up the agent suggestion popup"""
        try:
            if hasattr(self, 'agent_suggestion_popup') and self.agent_suggestion_popup.winfo_exists():
                self.agent_suggestion_popup.destroy()
        except:
            pass
        finally:
            if hasattr(self, 'agent_suggestion_popup'):
                del self.agent_suggestion_popup
            self.unbind("<Button-1>")  # Remove the global click binding

    def show_doctor_suggestions(self, event):
        # Destroy any existing popup first
        if hasattr(self, 'doctor_suggestion_popup'):
            self.doctor_suggestion_popup.destroy()

        search_term = self.ref_by_var.get().lower()
        if len(search_term) < 1:  # Only show after 2+ characters typed
            return

        matches = [doc for doc in self.doctors_data if search_term in doc.lower()]
        if not matches:
            return

        x, y = self.ref_by_entry.winfo_rootx(), self.ref_by_entry.winfo_rooty() + self.ref_by_entry.winfo_height()
        self.doctor_suggestion_popup = ctk.CTkToplevel(self)
        self.doctor_suggestion_popup.overrideredirect(True)
        self.doctor_suggestion_popup.geometry(f"+{x}+{y}")
        self.doctor_suggestion_popup.lift()

        # Make sure the popup stays above other windows
        self.doctor_suggestion_popup.attributes('-topmost', True)

        HOVER_BG = "#3a7ebf" if ctk.get_appearance_mode() == "Dark" else "#e0e0e0"
        HOVER_FG = "#ffffff" if ctk.get_appearance_mode() == "Dark" else "#000000"
        NORMAL_BG = "#2b2b2b" if ctk.get_appearance_mode() == "Dark" else "#ffffff"
        NORMAL_FG = "#ffffff" if ctk.get_appearance_mode() == "Dark" else "#000000"

        for doc in matches:
            frame = ctk.CTkFrame(self.doctor_suggestion_popup, fg_color=NORMAL_BG)
            frame.pack(fill="x")
    
            lbl = ctk.CTkLabel(
                frame, 
                text=doc, 
                anchor="w", 
                cursor="hand2",
                text_color=NORMAL_FG
            )
            lbl.pack(fill="x", padx=5, pady=2)
        
            # Bind the click event to both the label and the frame
            for widget in [lbl, frame]:
                widget.bind("<Button-1>", lambda e, d=doc: self.select_doctor_suggestion(d))
        
            # Hover effects
            lbl.bind("<Enter>", lambda e, l=lbl, f=frame: (
                f.configure(fg_color=HOVER_BG),
                l.configure(text_color=HOVER_FG)
            ))
            lbl.bind("<Leave>", lambda e, l=lbl, f=frame: (
                f.configure(fg_color=NORMAL_BG),
                l.configure(text_color=NORMAL_FG)
            ))

        # Bind click outside to close suggestions
        self.bind("<Button-1>", self.check_click_outside_doctor_suggestions)
        self.bind("<Tab>", self.check_click_outside_doctor_suggestions)

    def select_doctor_suggestion(self, doctor):
        """Handle selection of a doctor from the suggestions"""
        self.ref_by_var.set(doctor)  # Set the value in the entry
        self.close_doctor_suggestions()  # Close the popup
        self.ref_by_entry.focus()  # Return focus to the entry
        self.ref_by_entry.icursor("end")  # Move cursor to end of text

    def check_click_outside_doctor_suggestions(self, event):
        """Check if click was outside the doctor suggestion popup"""
        if not hasattr(self, 'doctor_suggestion_popup') or not self.doctor_suggestion_popup.winfo_exists():
            self.unbind("<Button-1>")
            return

        try:
            popup = self.doctor_suggestion_popup
            popup_x = popup.winfo_rootx()
            popup_y = popup.winfo_rooty()
            popup_width = popup.winfo_width()
            popup_height = popup.winfo_height()
    
            # Check if click was outside the popup
            if not (popup_x <= event.x_root <= popup_x + popup_width and
                    popup_y <= event.y_root <= popup_y + popup_height):
                self.close_doctor_suggestions()
        except:
            self.close_doctor_suggestions()

    def close_doctor_suggestions(self):
        """Clean up the doctor suggestion popup"""
        try:
            if hasattr(self, 'doctor_suggestion_popup') and self.doctor_suggestion_popup.winfo_exists():
                self.doctor_suggestion_popup.destroy()
        except:
            pass
        finally:
            if hasattr(self, 'doctor_suggestion_popup'):
                del self.doctor_suggestion_popup
            self.unbind("<Button-1>")  # Remove the global click binding

    def create_test_entry(self, parent_frame):
        test_frame = ctk.CTkFrame(parent_frame)
        test_frame.grid(row=1, column=0, columnspan=2, pady=10, padx=10, sticky="nsew")

    # Test Name Entry
        ctk.CTkLabel(test_frame, text="Test Name:").grid(row=0, column=0, padx=5, pady=5)
        self.test_name_entry = ctk.CTkEntry(test_frame, textvariable=self.test_name_var, width=400)
        self.test_name_entry.grid(row=0, column=1, padx=5, pady=5)


        self.test_name_entry.bind("<KeyRelease>", self.show_suggestions)

    # Test Amount Entry
        ctk.CTkLabel(test_frame, text="Amount:").grid(row=1, column=0, padx=5, pady=5)
        self.rate_entry = ctk.CTkEntry(test_frame, textvariable=self.rate_var, width=400)
        self.rate_entry.grid(row=1, column=1, padx=5, pady=5)

        #Enter key to add amount

        self.rate_entry.bind("<Return>", lambda e: self.add_test())

    # Advance Payment Entry
        ctk.CTkLabel(test_frame, text="Advance Payment:").grid(row=2, column=0, padx=5, pady=5)
        self.advanced_entry = ctk.CTkEntry(test_frame, textvariable=self.advanced_var, width=400)
        self.advanced_entry.grid(row=2, column=1, padx=5, pady=5)


    # Discount Amount Entry (New)
        ctk.CTkLabel(test_frame, text="Discount Amount:").grid(row=3, column=0, padx=5, pady=5)
        self.discount_entry = ctk.CTkEntry(test_frame, textvariable=self.discount_var, width=400)
        self.discount_entry.grid(row=3, column=1, padx=5, pady=5)

    # Add Test Button
        self.add_test_button = ctk.CTkButton(
            test_frame,
            text="Add Test",
            command=self.add_test,
            fg_color="#2E8B57",  # Normal color (green)
            hover_color="#245c3d",  # Hover color
            font=ctk.CTkFont(size=14, weight="bold")
        )
        self.add_test_button.grid(row=4, column=1, pady=10)

        # Focus behavior for Add Test button
        self.add_test_button.bind("<FocusIn>", lambda e: self.add_test_button.configure(fg_color="#FFD700", text_color="black"))
        self.add_test_button.bind("<FocusOut>", lambda e: self.add_test_button.configure(fg_color="#2E8B57",text_color="white"))
        


    # Fix: Use focus_set() instead of focus_widget()
        self.advanced_entry.bind("<Tab>", lambda e: self.move_focus(self.add_test_button))
        self.add_test_button.bind("<Tab>", lambda e: self.move_focus(self.print_invoice_button))

    def move_focus(self, widget):
        """Moves focus to the given widget."""
        widget.focus_set()
        # Bind Space and Enter keys to simulate a button click
        widget.bind("<Return>", lambda e: widget.invoke())
        widget.bind("<space>", lambda e: widget.invoke())
        return "break"  # Prevents default tab behavior
        



    def bind_events(self):
        self.test_name_entry.bind("<KeyRelease>", self.show_suggestions)
        self.bind("<F5>", lambda event: self.print_invoice())

    def load_test_data(self):
        """Loads test names and amounts from file."""
        test_dict = {}
        if os.path.exists("test_amount.txt"):
            with open("test_amount.txt", "r", encoding="utf-8") as f:
                for line in f:
                    try:
                        test_name, test_amount = line.strip().split(" - ")
                        test_dict[test_name] = float(test_amount)
                    except ValueError:
                        continue
        return test_dict
    
    def get_next_bill_number(self):
        if not os.path.exists(BILL_COUNTER_FILE):
            with open(BILL_COUNTER_FILE, "w") as f:
                f.write("1")

        with open(BILL_COUNTER_FILE, "r") as f:
            serial_number = int(f.read().strip())

        bill_number = f"GPDL{serial_number:04d}"

        with open(BILL_COUNTER_FILE, "w") as f:
            f.write(str(serial_number + 1))

        return bill_number


    def show_suggestions(self, event):
        # Destroy any existing popup first
        if hasattr(self, 'suggestion_popup'):
            self.suggestion_popup.destroy()
    
        search_term = self.test_name_var.get().lower()
    
        if not search_term:
            return
    
        matches = [test for test in self.test_data.keys() if search_term in test.lower()]
    
        if not matches:  # No matches found
            return
    
        # Create new suggestion popup
        x, y = self.test_name_entry.winfo_rootx(), self.test_name_entry.winfo_rooty() + self.test_name_entry.winfo_height()
        self.suggestion_popup = ctk.CTkToplevel(self)
        self.suggestion_popup.overrideredirect(True)
        self.suggestion_popup.geometry(f"+{x}+{y}")
        self.suggestion_popup.lift()

        # Custom colors for better visibility
        HOVER_BG = "#3a7ebf"  # Blue background on hover
        HOVER_FG = "#ffffff"  # White text on hover
        NORMAL_BG = "#2b2b2b"  # Dark background
        NORMAL_FG = "#ffffff"  # White text
    
        # Create suggestion items
        for test in matches:
            frame = ctk.CTkFrame(self.suggestion_popup, fg_color=NORMAL_BG)
            frame.pack(fill="x")
        
            lbl = ctk.CTkLabel(
                frame, 
                text=test, 
                anchor="w", 
                cursor="hand2",
                text_color=NORMAL_FG
            )
            lbl.pack(fill="x", padx=5, pady=2)
        
            # Bind click event
            lbl.bind("<Button-1>", lambda e, t=test: self.select_suggestion(t))        
            
            
            # Hover effects - more visible highlighting
            lbl.bind("<Enter>", lambda e, l=lbl, f=frame: (
                f.configure(fg_color=HOVER_BG),
                l.configure(text_color=HOVER_FG)
            ))
            lbl.bind("<Leave>", lambda e, l=lbl, f=frame: (
                f.configure(fg_color=NORMAL_BG),
                l.configure(text_color=NORMAL_FG)
            ))

        
 

    def select_suggestion(self, test):
        self.test_name_var.set(test)
        self.rate_var.set(self.test_data[test])
        if hasattr(self, 'suggestion_popup'):
            self.suggestion_popup.destroy()
        self.test_name_entry.focus()

    def clear_suggestions(self):
        if hasattr(self, 'suggestion_popup'):
            self.suggestion_popup.destroy()

    def add_test(self):
        """Add a test with maximum 10 tests limit"""
        if len(self.tests) >= 10:
            messagebox.showwarning("Limit Reached", "Maximum 10 tests per bill allowed")
            return
        
        test_name = self.test_name_var.get()
        rate = self.rate_var.get()
    
        if not test_name or not rate:
            return
        
        try:
            rate = float(rate)
            self.tests.append((test_name, rate))
            self.add_test_to_list(test_name, rate)
            self.test_name_var.set("")
            self.rate_var.set("")
            self.test_name_entry.focus()
            # --- MODIFIED --- Update calculator when a test is added
            self.update_totals_display()
        except ValueError:
            pass

    def add_test_to_list(self, test_name, rate):
        # Create the main frame for the row. The color will be managed by the hover events.
        frame = ctk.CTkFrame(self.test_list_frame)
        frame.pack(fill="x", pady=2)

        # Store labels in variables so we can bind events to them
        label_name = ctk.CTkLabel(frame, text=test_name, bg_color="transparent")
        label_name.pack(side="left", padx=5)

        label_amount = ctk.CTkLabel(frame, text=f"₹{rate:.2f}", bg_color="transparent")
        label_amount.pack(side="right", padx=5)

        # --- NEW: Dynamic hover functions ---
        def on_enter(event):
            """Checks the current theme and applies the correct hover color."""
            # Define your desired hover colors here
            if ctk.get_appearance_mode() == "Dark":
                hover_color = "#2E8B57"  # SeaGreen, good on dark backgrounds
            else: # Light mode
                hover_color = "#98FB98"  # PaleGreen, good on light backgrounds
            frame.configure(fg_color=hover_color)

        def on_leave(event):
            """Resets the frame to the theme's default background color."""
            # Get the current default background color for any CTkFrame
            normal_color = ctk.ThemeManager.theme["CTkFrame"]["fg_color"]
            frame.configure(fg_color=normal_color)

        # Bind events to ALL widgets in the row for a smooth effect
        widgets_to_bind = [frame, label_name, label_amount]
        for widget in widgets_to_bind:
            widget.bind("<Enter>", on_enter)
            widget.bind("<Leave>", on_leave)
            # Re-bind the right-click menu to all widgets
            widget.bind("<Button-3>", lambda e, f=frame: self.show_test_list_context_menu(e, f))

    

    def show_test_list_context_menu(self, event, frame):
        # Create standard Tkinter menu with CustomTkinter styling
        menu = tk.Menu(self, tearoff=0)
        menu.configure(
            bg="#2b2b2b",  # Background color
            fg="#ffffff",  # Text color
            activebackground="#3b3b3b",  # Hover color
            activeforeground="#ffffff"
        )
    
        menu.add_command(
            label="Delete",
            command=lambda f=frame: self.delete_test(f)
        )
    
        try:
            menu.tk_popup(event.x_root, event.y_root)
        finally:
            menu.grab_release()

    def delete_test(self, frame):
        test_name = frame.winfo_children()[0].cget("text")
        amount_str = frame.winfo_children()[1].cget("text").replace("₹", "")
        
        # Find and remove the first matching test
        test_to_remove = None
        for test in self.tests:
            if test[0] == test_name and f"{test[1]:.2f}" == f"{float(amount_str):.2f}":
                test_to_remove = test
                break
        
        if test_to_remove:
            self.tests.remove(test_to_remove)
            frame.destroy()
            # --- MODIFIED --- Update calculator when a test is deleted
            self.update_totals_display()

    
    def show_test_context_menu(self, event, test_name, test_amount):
        """Shows the right-click menu for managing tests."""
        # Get current theme mode
        current_mode = ctk.get_appearance_mode()
    
        # Define colors for light and dark mode
        if current_mode == "Dark":
            bg_color = "#333333"  # Dark gray background
            fg_color = "#ffffff"  # White text
            hover_bg = "#555555"  # Slightly lighter gray for hover
        else:
            bg_color = "#ffffff"  # White background
            fg_color = "#000000"  # Black text
            hover_bg = "#e5e5e5"  # Light gray for hover

        # Create a new menu
        menu = tk.Menu(self, tearoff=0)
        menu.configure(
            bg=bg_color,
            fg=fg_color,
            activebackground=hover_bg,
            activeforeground=fg_color
        )

        # Add menu options
        menu.add_command(label="Edit", command=lambda: self.open_edit_window(test_name, test_amount))
        menu.add_command(label="Delete", command=lambda: self.delete_test_from_file(test_name, self.manage_window))

        try:
            menu.tk_popup(event.x_root, event.y_root)
        finally:
            menu.grab_release()


    
    def set_cell_text(self, cell, text, font_size=14):
        """Set text in a table cell while preserving formatting with a fixed font size."""
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(text)
        run.font.size = Pt(font_size)  # Set font size to 14 pt


    def save_invoice(self):
        """Generates and saves the invoice in DOCX format."""
        if not self.client_name_var.get() or not self.age_var.get() or not self.gender_var.get():
            messagebox.showerror("Input Error", "Please enter Name, Age, and Gender.")
            return

        if not self.tests:
            messagebox.showerror("Input Error", "Please add at least one test.")
            return

        # Load template and set bill number
        document = Document(TEMPLATE_FILE)
        bill_number = self.get_next_bill_number()
        date_today = datetime.today().strftime("%d-%m-%Y")
        total_amount = sum(t[1] for t in self.tests)
        advanced_payment = float(self.advanced_var.get()) if self.advanced_var.get() else 0.0
        due_amount = total_amount - advanced_payment

        # Fill patient details in the first table
        tables = document.tables
        patient_table = tables[0]
        self.set_cell_text(patient_table.cell(0, 1), bill_number)
        self.set_cell_text(patient_table.cell(0, 3), date_today)
        self.set_cell_text(patient_table.cell(1, 1), self.client_name_var.get())
        self.set_cell_text(patient_table.cell(1, 3), self.age_var.get())
        self.set_cell_text(patient_table.cell(2, 1), self.ref_by_var.get())
        self.set_cell_text(patient_table.cell(2, 3), self.gender_var.get())
        self.set_cell_text(patient_table.cell(3, 1), self.address_var.get())
        self.set_cell_text(patient_table.cell(3, 3), self.contact_var.get())

        # Fill test details in the second table
        test_table = tables[1]
        for _ in range(len(test_table.rows) - 1):
            test_table._element.remove(test_table.rows[1]._element)

        # Add exactly 10 rows (filled or blank)
        for i in range(10):
            row_cells = test_table.add_row().cells
            if i < len(self.tests):
                # Fill with actual test data
                test_name, amount = self.tests[i]
                self.set_cell_text(row_cells[0], test_name)
                self.set_cell_text(row_cells[1], f"₹{amount:.2f}")
            else:
                # Add blank rows to maintain formatting
                self.set_cell_text(row_cells[0], "")
                self.set_cell_text(row_cells[1], "")
                

        # Calculate amounts with discount
        total_amount = sum(t[1] for t in self.tests)
        discount = float(self.discount_var.get()) if self.discount_var.get() else 0.0
        advanced_payment = float(self.advanced_var.get()) if self.advanced_var.get() else 0.0
        due_amount = (total_amount - discount) - advanced_payment

        # Fill total, advanced, and due in the third table
        total_table = tables[2]
        self.set_cell_text(total_table.cell(0, 1), f"₹{total_amount:.2f}")
        self.set_cell_text(total_table.cell(1, 1), f"₹{discount:.2f}")  # Show discount
        self.set_cell_text(total_table.cell(2, 1), f"₹{advanced_payment:.2f}")
        self.set_cell_text(total_table.cell(3, 1), f"₹{due_amount:.2f}")

        # Save the document
        self.invoice_filename = os.path.join(BILL_FOLDER, f"{bill_number}.docx")
        os.makedirs(BILL_FOLDER, exist_ok=True)
        document.save(self.invoice_filename)

        return self.invoice_filename




    def print_invoice(self):
        """Creates a 2-page PDF with identical copies of the invoice"""
        """Converts the DOCX invoice to PDF and opens it for printing."""
        docx_file = self.save_invoice()
        if not docx_file:
            return
        
        # Save new doctor and agent if they don't exist
        self.save_new_doctor_if_not_exists(self.ref_by_var.get())
        self.save_new_agent_if_not_exists(self.agent_var.get())

        # Convert DOCX to PDF
        pdf_file = docx_file.replace(".docx", ".pdf")
        self.convert_docx_to_pdf(docx_file, pdf_file)
        # First convert to single PDF
        single_pdf = docx_file.replace(".docx", "_single.pdf")
        self.convert_docx_to_pdf(docx_file, single_pdf)
    
        # Merge two copies
        from PyPDF2 import PdfMerger
        merger = PdfMerger()
        merger.append(single_pdf)
        merger.append(single_pdf)
    
        two_page_pdf = single_pdf.replace("_single.pdf", "_2page.pdf")
        merger.write(two_page_pdf)
        merger.close()
    
        webbrowser.open(two_page_pdf)
    
        # Clean up temporary files
        os.remove(single_pdf)


        

        # Save billing details to Excel
        total_amount = sum(t[1] for t in self.tests)
        discount_amount = float(self.discount_var.get()) if self.discount_var.get() else 0.0
        advanced_payment = float(self.advanced_var.get()) if self.advanced_var.get() else 0.0

        net_total_after_discount = total_amount - discount_amount
        due_amount = net_total_after_discount - advanced_payment
        bill_number = os.path.basename(docx_file).replace(".docx", "")

        self.save_to_excel(bill_number, total_amount, advanced_payment, due_amount)

        # Clear all patient details after printing
        self.clear_patient_details()

    


    def convert_docx_to_pdf(self, docx_path, pdf_path):
        """Converts a DOCX file to PDF using Microsoft Word."""
        word = comtypes.client.CreateObject("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(os.path.abspath(docx_path))
        doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)  # 17 is PDF format
        doc.Close()
        word.Quit()


    def save_to_excel(self, bill_number, total_amount, advanced_payment, due_amount):
        discount = float(self.discount_var.get()) if self.discount_var.get() else 0.0
    
        bill_data = {
            "Bill Number": bill_number,
            "Date": datetime.today().strftime("%d-%m-%Y"),
            "Patient Name": self.client_name_var.get(),
            "Age": self.age_var.get(),
            "Gender": self.gender_var.get(),
            "Ref By": self.ref_by_var.get(),
            "Agent": self.agent_var.get(),  # Include agent in Excel
            "Address": self.address_var.get(),
            "Contact": self.contact_var.get(),
            "Total Amount": total_amount,
            "Discount": discount,  # Include discount in Excel
            "Advanced Payment": advanced_payment,
            "Due Amount": due_amount,
            "Due Payment Date": "",  # new column
            "Tests": ", ".join([f"{test[0]} ({test[1]})" for test in self.tests])
        }

        # Convert to DataFrame
        new_data = pd.DataFrame([bill_data])

        try:
            if os.path.exists(EXCEL_FILE):
                df = pd.read_excel(EXCEL_FILE)
                df = pd.concat([df, new_data], ignore_index=True)
            else:   
                df = new_data

            # Save to Excel
            df.to_excel(EXCEL_FILE, index=False, engine="openpyxl")

            # Load workbook to modify columns
            wb = openpyxl.load_workbook(EXCEL_FILE)
            ws = wb.active

            # Auto-adjust column width based on the longest text in each column
            for col in ws.columns:
                max_length = 0
                col_letter = col[0].column_letter  # Get column letter (A, B, C, etc.)
                for cell in col:
                    try:
                        if cell.value:
                            max_length = max(max_length, len(str(cell.value)))
                    except:
                        pass
                adjusted_width = max_length + 2  # Extra padding for readability
                ws.column_dimensions[col_letter].width = adjusted_width

            # Save the modified Excel file
            wb.save(EXCEL_FILE)
            print(f"✅ Bill {bill_number} saved in Excel with adjusted column width.")

        except Exception as e:
            print(f"❌ Error saving to Excel: {e}")


    def update_test_list(self):
        """Clears the test list in the UI after printing the invoice."""
        for widget in self.test_list_frame.winfo_children():
            widget.destroy()  # Remove all test entries from the UI



    def clear_patient_details(self):
        """Clears all form fields after printing the invoice"""
        # Clear existing fields
        self.client_name_var.set("")
        self.age_var.set("")
        self.gender_var.set("")
        self.ref_by_var.set("")
        self.agent_var.set("")  # Clear agent field
        self.address_var.set("")
        self.contact_var.set("")
    
        # Clear test-related fields
        self.test_name_var.set("")
        self.rate_var.set("")
        self.advanced_var.set("")
        self.discount_var.set("")  # Clear discount field
    
        # Clear tests list
        self.tests.clear()
        self.update_test_list()
    
        # Set focus back to patient name field for new entry
        self.entry_widgets["client_name"].focus_set()
        self.update_totals_display()

    
    def show_report_options(self):
        """Opens the main report module window."""
        if hasattr(self, 'report_module') and self.report_module.winfo_exists():
            self.report_module.lift()
            return
        self.report_module = ReportModuleWindow(self)



    def change_appearance_mode(self, mode):
        ctk.set_appearance_mode(mode)


    def create_settings_icon(self):
        # Settings icon frame
        self.settings_frame = ctk.CTkFrame(self.sidebar_frame, fg_color="transparent")
        self.settings_frame.pack(side="bottom", fill="x", pady=20)
        
        # Load icon images using CTkImage
        self.settings_img = CTkImage(
            Image.open("settings_icon.png"),
            size=(25, 25)
        )
        
        self.moon_img = CTkImage(
            Image.open("moon_icon.png"),
            size=(20, 20)
        )
        
        self.sun_img = CTkImage(
            Image.open("sun_icon.png"),
            size=(20, 20)
        )

        # Settings button
        self.settings_btn = ctk.CTkButton(
            self.settings_frame,
            image=self.settings_img,
            text="",
            width=30,
            height=30,
            fg_color="transparent",
            hover_color=("#e0e0e0", "#2b2b2b"),
            command=self.toggle_settings
        )
        self.settings_btn.pack(side="left", padx=10)

    def show_settings_menu(self):
        """Shows the settings menu with theme-aware text colors."""

        if hasattr(self, 'menu_window') and self.menu_window.winfo_exists():
            self.menu_window.destroy()

        self.menu_window = ctk.CTkToplevel(self)
        self.menu_window.attributes("-topmost", True)
        self.menu_window.overrideredirect(True)

        # --- NEW: Define colors based on the current theme ---
        if ctk.get_appearance_mode() == "Dark":
            text_color = "white"
            hover_color = "#1f6aa5"  # A nice blue
        else: # Light mode
            text_color = "black"
            hover_color = "#1f6aa5"  # Same blue for consistency

        popup_width = 180
        popup_height = 120
        self.menu_window.geometry(f"{popup_width}x{popup_height}")

        btn_x = self.settings_btn.winfo_rootx()
        btn_y = self.settings_btn.winfo_rooty()
        button_width = self.settings_btn.winfo_width()
        button_height = self.settings_btn.winfo_height()

        x = btn_x + button_width + 5
        y = btn_y + (button_height // 2) - (popup_height // 2) - 50
        
        # (Positioning logic remains the same)
        screen_width = self.winfo_screenwidth()
        if x + popup_width > screen_width: x = btn_x - popup_width - 5
        if y < 5: y = 5
        self.menu_window.geometry(f"+{int(x)}+{int(y)}")

        # === Appearance ===
        appearance_frame = ctk.CTkFrame(self.menu_window, fg_color="transparent")
        appearance_frame.pack(fill="x", pady=2, padx=5)
        
        # Apply the dynamic text color
        appearance_label = ctk.CTkLabel(appearance_frame, text="Appearance", text_color=text_color)
        appearance_label.pack(side="left", padx=5)

        self.theme_icon = ctk.CTkLabel(
            appearance_frame, text="",
            image=self.sun_img if ctk.get_appearance_mode() == "Dark" else self.moon_img,
            width=20
        )
        self.theme_icon.pack(side="right", padx=5)
        self.theme_icon.bind("<Button-1>", lambda e: self.safe_toggle_theme())

        # === Settings ===
        settings_frame = ctk.CTkFrame(self.menu_window, fg_color="transparent")
        settings_frame.pack(fill="x", pady=2, padx=5)

        # Apply dynamic colors and update hover bindings
        settings_label = ctk.CTkLabel(settings_frame, text="Settings", cursor="hand2", text_color=text_color)
        settings_label.pack(side="left", padx=5)
        settings_label.bind("<Enter>", lambda e, lbl=settings_label: lbl.configure(text_color=hover_color))
        settings_label.bind("<Leave>", lambda e, lbl=settings_label: lbl.configure(text_color=text_color))

        def open_help(event):
            # (This function's content remains unchanged)
            self.menu_window.destroy()
            help_window = ctk.CTkToplevel(self)
            help_window.title("How to Use")
            help_window.geometry("500x400")
            help_window.attributes("-topmost", True)
            help_window.resizable(False, False)
            instructions = (
                "📋 How to Use GreenPath Billing Software:\n\n"
                "1. Enter patient details: Name, Age, Gender, etc.\n"
                "2. Start typing a test name — suggestions will appear.\n"
                "3. Click a test to autofill its amount.\n"
                "4. Add more tests as needed.\n"
                "5. Click 'Print Invoice' to generate a PDF.\n"
                "6. Invoices are saved inside the 'bill' folder.\n"
                "7. Use the ⚙️ Settings Panel to edit tests and theme.\n\n"
                "📧 Still need help? Click 'Report Us' to contact support."
            )
            help_label = ctk.CTkLabel(
                help_window, text=instructions, wraplength=460, justify="left",
                font=ctk.CTkFont(size=14), anchor="nw"
            )
            help_label.pack(padx=20, pady=20, fill="both", expand=True)
        settings_label.bind("<Button-1>", open_help)

        # === Report Us ===
        report_frame = ctk.CTkFrame(self.menu_window, fg_color="transparent")
        report_frame.pack(fill="x", pady=2, padx=5)

        # Apply dynamic colors and update hover bindings
        report_label = ctk.CTkLabel(report_frame, text="Report Us", cursor="hand2", text_color=text_color)
        report_label.pack(side="left", padx=5)
        report_label.bind("<Enter>", lambda e, lbl=report_label: lbl.configure(text_color=hover_color))
        report_label.bind("<Leave>", lambda e, lbl=report_label: lbl.configure(text_color=text_color))

        def open_mail(event):
            # (This function's content remains unchanged)
            self.menu_window.destroy()
            import webbrowser
            subject = "Complaint about GreenPath Billing Software"
            body = "What is the issue are you facing?"
            email = "sk.samimjul.islam.it.2023@tint.edu.in"
            url = f"https://mail.google.com/mail/?view=cm&fs=1&to={email}&su={subject}&body={body}"
            webbrowser.open(url)
        report_label.bind("<Button-1>", open_mail)

        self.bind_all("<Button-1>", self.safe_check_click_location)
    

    def safe_toggle_theme(self):
        """Toggles appearance mode and rebuilds the settings menu to apply new colors."""
        # Toggle appearance mode
        current_mode = ctk.get_appearance_mode()
        new_mode = "Light" if current_mode == "Dark" else "Dark"
        ctk.set_appearance_mode(new_mode)
        
        # --- NEW: Rebuild the menu to apply new colors ---
        # Hide the old menu first
        if hasattr(self, 'menu_window') and self.menu_window.winfo_exists():
            # Unbind the global click listener before destroying
            self.unbind_all("<Button-1>")
            self.menu_window.destroy()
            del self.menu_window

        # Show a new menu with the updated theme colors
        self.show_settings_menu()

    def safe_check_click_location(self, event):
        if not hasattr(self, 'menu_window') or not self.menu_window.winfo_exists():
            return

        # Get menu window geometry
        menu_x = self.menu_window.winfo_rootx()
        menu_y = self.menu_window.winfo_rooty()
        menu_width = self.menu_window.winfo_width()
        menu_height = self.menu_window.winfo_height()

        # Check if click is inside menu
        click_in_menu = (
            menu_x <= event.x_root <= menu_x + menu_width and
            menu_y <= event.y_root <= menu_y + menu_height
        )

        # Get settings button geometry
        btn_x = self.settings_btn.winfo_rootx()
        btn_y = self.settings_btn.winfo_rooty()
        btn_width = self.settings_btn.winfo_width()
        btn_height = self.settings_btn.winfo_height()

        # Check if click is inside settings button
        click_in_btn = (
            btn_x <= event.x_root <= btn_x + btn_width and
            btn_y <= event.y_root <= btn_y + btn_height
        )

        # Close menu only if clicking outside both areas
        if not click_in_menu and not click_in_btn:
            self.hide_settings_menu()

        
    

    def hide_settings_menu(self):
        if hasattr(self, 'menu_window'):
            try:
                self.unbind_all("<Button-1>")  # Remove global click binding
                if self.menu_window.winfo_exists():
                    self.menu_window.destroy()
            except Exception as e:
                print(f"Error closing menu: {e}")
            finally:
                del self.menu_window
        if hasattr(self, 'theme_icon'):
            del self.theme_icon

    def toggle_settings(self):
        # Rotate icon safely
        self.rotation_angle = (self.rotation_angle + 180) % 360
        rotated_img = CTkImage(
            Image.open("settings_icon.png").rotate(self.rotation_angle),
            size=(25, 25)
        )
        self.settings_btn.configure(image=rotated_img)
    
        # Toggle menu visibility
        if hasattr(self, 'menu_window') and self.menu_window.winfo_exists():
            self.hide_settings_menu()
        else:
            self.show_settings_menu()


    def load_doctors_data(self):
        """Loads doctor names from file."""
        doctors = []
        if os.path.exists(DOCTORS_FILE):
            with open(DOCTORS_FILE, "r", encoding="utf-8") as f:
                for line in f:
                    doctor = line.strip()
                    if doctor:
                        doctors.append(doctor)
        return doctors

    def save_doctors_data(self):
        """Saves doctors list to file."""
        with open(DOCTORS_FILE, "w", encoding="utf-8") as f:
            for doctor in self.doctors_data:
                f.write(f"{doctor}\n")

    
    
    def show_doctor_context_menu(self, event, doctor_name):
        """Shows the right-click menu for managing doctors."""
        current_mode = ctk.get_appearance_mode()
        bg_color = "#333333" if current_mode == "Dark" else "#ffffff"
        fg_color = "#ffffff" if current_mode == "Dark" else "#000000"
        hover_bg = "#555555" if current_mode == "Dark" else "#e5e5e5"

        menu = tk.Menu(self, tearoff=0)
        menu.configure(
            bg=bg_color,
            fg=fg_color,
            activebackground=hover_bg,
            activeforeground=fg_color
        )

        menu.add_command(label="Edit", command=lambda: self.open_edit_doctor_window(doctor_name))
        menu.add_command(label="Delete", command=lambda: self.delete_doctor(doctor_name))

        try:
            menu.tk_popup(event.x_root, event.y_root)
        finally:
            menu.grab_release()



        


    def load_agents_data(self):
        """Loads agent names from file."""
        agents = []
        if os.path.exists(AGENTS_FILE):
            with open(AGENTS_FILE, "r", encoding="utf-8") as f:
                for line in f:
                    agent = line.strip()
                    if agent:
                        agents.append(agent)
        return agents

    def save_agents_data(self):
        """Saves agents list to file."""
        with open(AGENTS_FILE, "w", encoding="utf-8") as f:
            for agent in self.agents_data:
                f.write(f"{agent}\n")

    

    def save_new_doctor_if_not_exists(self, doctor_name):
        """Save a new doctor to the file if it doesn't exist"""
        if not doctor_name or doctor_name in self.doctors_data:
            return
    
        self.doctors_data.append(doctor_name)
        with open(DOCTORS_FILE, "a", encoding="utf-8") as f:
            f.write(f"{doctor_name}\n")

    def save_new_agent_if_not_exists(self, agent_name):
        """Save a new agent to the file if it doesn't exist"""
        if not agent_name or agent_name in self.agents_data:
            return
    
        self.agents_data.append(agent_name)
        with open(AGENTS_FILE, "a", encoding="utf-8") as f:
            f.write(f"{agent_name}\n")


    
    def edit_bill(self, bill_number):
        """Open edit window for the selected bill"""
        # Load the data from Excel
        df = pd.read_excel(EXCEL_FILE)
        bill_data = df[df['Bill Number'] == bill_number].iloc[0].to_dict()

        # Replace NaN values with appropriate defaults
        for key, value in bill_data.items():
            if pd.isna(value):
                if key in ["Discount", "Advanced Payment", "Due Amount"]:
                    bill_data[key] = 0  # Set numeric fields to 0
                else:
                    bill_data[key] = ""  # Set text fields to blank

        # Create edit window
        self.edit_bill_window = ctk.CTkToplevel(self)
        self.edit_bill_window.title(f"Edit Bill - {bill_number}")
        # Set window state to maximized
        self.edit_bill_window.state("zoomed")

        self.edit_bill_window.attributes('-topmost', True)

        # Main frame
        main_frame = ctk.CTkScrollableFrame(self.edit_bill_window)
        main_frame.pack(fill="both", expand=True, padx=10, pady=10)

        # Patient details section
        ctk.CTkLabel(main_frame, text="Patient Details", font=ctk.CTkFont(weight="bold")).pack(pady=5, anchor="w")
    
        # Create form fields for patient details
        fields = [
            ("Patient Name:", "Patient Name", bill_data.get('Patient Name', '')),
            ("Age:", "Age", bill_data.get('Age', '')),
            ("Gender:", "Gender", bill_data.get('Gender', '')),
            ("Ref By Doctor:", "Ref By", bill_data.get('Ref By', '')),
            ("Agent:", "Agent", bill_data.get('Agent', '')),
            ("Address:", "Address", bill_data.get('Address', '')),
            ("Contact:", "Contact", bill_data.get('Contact', '')),
            ("Total Amount:", "Total Amount", bill_data.get('Total Amount', '')),
            ("Discount:", "Discount", bill_data.get('Discount', '0')),
            ("Advanced Payment:", "Advanced Payment", bill_data.get('Advanced Payment', '0')),
            ("Due Amount:", "Due Amount", bill_data.get('Due Amount', '0')),  # Add this line
            ("Due Payment Date:", "Due Payment Date", bill_data.get('Due Payment Date', '')),  # Add this line
        ]

        self.edit_vars = {}
        for idx, (label, key, value) in enumerate(fields):
            frame = ctk.CTkFrame(main_frame)
            frame.pack(fill="x", pady=2)
        
            ctk.CTkLabel(frame, text=label).pack(side="left", padx=5)
        
            # Ensure numeric fields are properly formatted
            if key == "Contact":
            # Remove any trailing '.0' if it’s there (from float conversion)
                value = str(value).rstrip(".0") if str(value).endswith(".0") else str(value)
                var = ctk.StringVar(value=value)
            elif key in ["Discount", "Advanced Payment", "Due Amount", "Total Amount"]:
                value = float(value) if value else 0.0
                var = ctk.StringVar(value=f"{value:.2f}" if value != 0 else "0.00")
            else:
                var = ctk.StringVar(value=str(value) if not pd.isna(value) else "")

            entry = ctk.CTkEntry(frame, textvariable=var)
            entry.pack(side="left", fill="x", expand=True, padx=5)
    
            self.edit_vars[key] = var

        # Tests section
        ctk.CTkLabel(main_frame, text="Tests", font=ctk.CTkFont(weight="bold")).pack(pady=(10,5), anchor="w")
    
        # Parse tests from the string
        tests_str = bill_data.get('Tests', '')
        self.edit_tests = []
    
        # Test list frame
        test_list_frame = ctk.CTkScrollableFrame(main_frame, height=150)
        test_list_frame.pack(fill="x", pady=5)
    
        # Add existing tests
        if tests_str:
            test_entries = tests_str.split(", ")
            for entry in test_entries:
                try:
                    name_part, sep, amount_part = entry.rpartition(" (")
                    if sep and amount_part.endswith(")"):
                        test_name = name_part.strip()
                        test_amount = float(amount_part[:-1])  # remove trailing ')'
                        self.edit_tests.append((test_name, test_amount))
                        self.add_test_to_edit_list(test_list_frame, test_name, test_amount)
                    else:
                        print(f"Skipping malformed test entry: '{entry}'")
                except Exception as e:
                    print(f"Error parsing test entry '{entry}': {e}")



        # Add test section
        add_test_frame = ctk.CTkFrame(main_frame)
        add_test_frame.pack(fill="x", pady=5)
    
        self.new_test_var = ctk.StringVar()
        self.new_amount_var = ctk.StringVar()
    
        ctk.CTkLabel(add_test_frame, text="Test Name:").pack(side="left", padx=5)
        test_entry = ctk.CTkEntry(add_test_frame, textvariable=self.new_test_var)
        test_entry.pack(side="left", fill="x", expand=True, padx=5)
    
        ctk.CTkLabel(add_test_frame, text="Amount:").pack(side="left", padx=5)
        amount_entry = ctk.CTkEntry(add_test_frame, textvariable=self.new_amount_var)
        amount_entry.pack(side="left", padx=5)
    
        add_btn = ctk.CTkButton(
            add_test_frame, 
            text="Add Test", 
            command=lambda: self.add_test_to_edit(test_list_frame),
            state="normal"
        )
        add_btn.pack(side="left", padx=5)

        # Buttons frame
        btn_frame = ctk.CTkFrame(main_frame)
        btn_frame.pack(fill="x", pady=10)
    
        ctk.CTkButton(
            btn_frame, 
            text="Save & Print", 
            command=lambda: self.save_and_print_edited_bill(bill_number),
            fg_color="#2E8B57"
        ).pack(side="left", padx=5)
    
        ctk.CTkButton(
            btn_frame, 
            text="Cancel", 
            command=self.edit_bill_window.destroy
        ).pack(side="right", padx=5)

        # Bind suggestion functionality to the test entry
        test_entry.bind("<KeyRelease>", 
            lambda e: self.show_edittest_suggestions(e, test_entry, self.new_test_var, self.new_amount_var))

        
        
        
    def show_edittest_suggestions(self, event, entry_widget, test_var, rate_var):
        """Generic suggestion display for any test entry"""
        # Destroy existing popup
        if hasattr(self, 'edit_suggestion_popup'):
            self.edit_suggestion_popup.destroy()

        search_term = test_var.get().lower()
        if len(search_term) < 1:  # Only show after 1+ characters
            return

        matches = [test for test in self.test_data.keys() if search_term in test.lower()]
        if not matches:
            return

        # Get screen position of the entry widget
        x = entry_widget.winfo_rootx()
        y = entry_widget.winfo_rooty() + entry_widget.winfo_height()

        # Create suggestion popup
        self.edit_suggestion_popup = ctk.CTkToplevel(self)
        self.edit_suggestion_popup.overrideredirect(True)
        self.edit_suggestion_popup.geometry(f"+{x}+{y}")
        self.edit_suggestion_popup.lift()
        self.edit_suggestion_popup.attributes('-topmost', True)

        
        self.edit_suggestion_popup.bind_all("<Button-1>", self.check_click_outside_edit_suggestions)



        # Custom colors for better visibility
        HOVER_BG = "#3a7ebf"  # Blue background on hover
        HOVER_FG = "#ffffff"  # White text on hover
        NORMAL_BG = "#2b2b2b"  # Dark background
        NORMAL_FG = "#ffffff"  # White text

        # Create suggestion items
        for test in matches:
            frame = ctk.CTkFrame(self.edit_suggestion_popup, fg_color=NORMAL_BG)
            frame.pack(fill="x")
        
            lbl = ctk.CTkLabel(frame, text=test, anchor="w", cursor="hand2", text_color=NORMAL_FG)
            lbl.pack(fill="x", padx=5, pady=2)
        
            # Bind click to set both test and amount variables
            lbl.bind("<Button-1>", 
                   lambda e, t=test: self.select_edittest_suggestion(t, test_var, rate_var))
            
            # Hover effects - more visible highlighting
            lbl.bind("<Enter>", lambda e, l=lbl, f=frame: (
                f.configure(fg_color=HOVER_BG),
                l.configure(text_color=HOVER_FG)
            ))
            lbl.bind("<Leave>", lambda e, l=lbl, f=frame: (
                f.configure(fg_color=NORMAL_BG),
                l.configure(text_color=NORMAL_FG)
            ))

    
    def check_click_outside_edit_suggestions(self, event):
        """Check if click was outside suggestion popup"""
        if not hasattr(self, 'edit_suggestion_popup') or not self.edit_suggestion_popup.winfo_exists():
            self.edit_suggestion_popup.unbind_all("<Button-1>")
            return

        try:
            popup = self.edit_suggestion_popup
            in_popup = (popup.winfo_rootx() <= event.x_root <= popup.winfo_rootx() + popup.winfo_width() and
                        popup.winfo_rooty() <= event.y_root <= popup.winfo_rooty() + popup.winfo_height())

            if not in_popup:
                self.close_edit_suggestions()
        except:
            self.close_edit_suggestions()



    def close_edit_suggestions(self):
        """Clean up suggestion popup"""
        try:
            if hasattr(self, 'edit_suggestion_popup') and self.edit_suggestion_popup.winfo_exists():
                self.edit_suggestion_popup.unbind_all("<Button-1>")
                self.edit_suggestion_popup.destroy()
        except:
            pass
        finally:
            if hasattr(self, 'edit_suggestion_popup'):
                del self.edit_suggestion_popup


    def select_edittest_suggestion(self, test, test_var, rate_var):
        """Handle suggestion selection for any entry"""
        test_var.set(test)
        rate_var.set(self.test_data.get(test, 0))
    
        if hasattr(self, 'edit_suggestion_popup'):
            self.edit_suggestion_popup.destroy()
    
        # Move focus to amount field
        if rate_var.get() == 0:
            self.focus_set()  # Or specific widget as needed

    

    def add_test_to_edit_list(self, parent_frame, test_name, test_amount):
        """Add a test to the edit list"""
        normal_color = ("#E5E5E5", "#2b2b2b")   # (Light Mode, Dark Mode)
        hover_color = ("#CCCCCC", "#3a3a3a")    # (Light Mode, Dark Mode)
        

        frame = ctk.CTkFrame(parent_frame, fg_color=normal_color)
        frame.pack(fill="x", pady=2)

        label_name = ctk.CTkLabel(frame, text=test_name)
        label_name.pack(side="left", padx=5)

        label_amount = ctk.CTkLabel(frame, text=f"₹{test_amount:.2f}")
        label_amount.pack(side="right", padx=5)

        # Hover effects
        def on_enter(e):
            frame.configure(fg_color=hover_color)

        def on_leave(e):
            frame.configure(fg_color=normal_color)

        frame.bind("<Enter>", on_enter)
        frame.bind("<Leave>", on_leave)

        # Right-click menu
        def show_right_click_menu(event, frame, test_name, test_amount):
            # Create right-click menu
            menu = tk.Menu(self.edit_bill_window, tearoff=False)
            menu.add_command(label="Delete", command=lambda: self.delete_test_from_edit(frame, test_name, test_amount))

            # Show menu at mouse pointer position
            menu.post(event.x_root, event.y_root)

        # Bind right-click to show delete option
        frame.bind("<Button-3>", lambda e, f=frame, t_name=test_name, t_amount=test_amount: show_right_click_menu(e, f, t_name, t_amount))


    def add_test_to_edit(self, parent_frame):
        """Add a new test to the edit list with 10-test maximum"""
        if len(self.edit_tests) >= 10:
            self.show_top_warning("Maximum 10 tests reached!")
            return
        
        test_name = self.new_test_var.get()
        amount = self.new_amount_var.get()
    
        if not test_name or not amount:
            return
    
        try:
            amount = float(amount)
            self.edit_tests.append((test_name, amount))
            self.add_test_to_edit_list(parent_frame, test_name, amount)
            self.new_test_var.set("")
            self.new_amount_var.set("")
        
            
        except ValueError:
            messagebox.showerror("Error", "Please enter a valid amount")

    def show_top_warning(self, message):
        """Show a temporary warning popup at top of screen"""
        x = self.winfo_x() + (self.winfo_width() // 2) - 150  # Center horizontally
        y = self.winfo_y() + 50  # 50px from top of main window
    
        warning = ctk.CTkToplevel(self)
        warning.title("Warning")
        warning.geometry(f"300x80+{x}+{y}")
        warning.attributes('-topmost', True)
        warning.overrideredirect(True)  # Remove window decorations
    
        # Warning frame with attention color
        frame = ctk.CTkFrame(warning, fg_color="#FFF3CD", corner_radius=5)
        frame.pack(fill="both", expand=True, padx=5, pady=5)
    
        ctk.CTkLabel(frame, text="⚠️ " + message, 
                    text_color="#856404",  # Dark yellow text
                    font=ctk.CTkFont(weight="bold")).pack(pady=10)
    
        # Auto-close after 2 seconds
        warning.after(2000, warning.destroy)


    def delete_test_from_edit(self, frame, test_name, test_amount):
        """Delete a test from the edit list."""
        # Remove from tests list
        self.edit_tests = [(t, a) for t, a in self.edit_tests if not (t == test_name and a == test_amount)]

        # Destroy the frame
        frame.destroy()

        # Re-enable add button if under limit
        if len(self.edit_tests) < 10:
            for widget in self.edit_bill_window.winfo_children():
                if isinstance(widget, ctk.CTkFrame):
                    for subwidget in widget.winfo_children():
                        if isinstance(subwidget, ctk.CTkButton) and "Add Test" in subwidget.cget("text"):
                            subwidget.configure(state="normal")

    def save_and_print_edited_bill(self, original_bill_number):
        """Save the edited bill and print it"""
        # Get all edited values
        patient_data = {key: var.get() for key, var in self.edit_vars.items()}

        # Calculate new totals
        try:
            total_amount = sum(t[1] for t in self.edit_tests)
            discount = float(patient_data.get('Discount', 0))
            advanced = float(patient_data.get('Advanced Payment', 0))
            due_amount = (total_amount - discount) - advanced
           
            # Update the totals in the display
            self.edit_vars['Total Amount'].set(f"{total_amount:.2f}")
            if 'Due Amount' in self.edit_vars:
                self.edit_vars['Due Amount'].set(f"{due_amount:.2f}")
        except ValueError:
            messagebox.showerror("Error", "Invalid amount values")
            return
       
        # Generate the new bill document
        document = Document(TEMPLATE_FILE)
        date_today = datetime.today().strftime("%d-%m-%Y")
       
        # Fill patient details in the first table
        tables = document.tables
        patient_table = tables[0]
        self.set_cell_text(patient_table.cell(0, 1), original_bill_number)
        self.set_cell_text(patient_table.cell(0, 3), date_today)
        self.set_cell_text(patient_table.cell(1, 1), patient_data['Patient Name'])
        self.set_cell_text(patient_table.cell(1, 3), patient_data['Age'])
        self.set_cell_text(patient_table.cell(2, 1), patient_data['Ref By'])
        self.set_cell_text(patient_table.cell(2, 3), patient_data['Gender'])
        self.set_cell_text(patient_table.cell(3, 1), patient_data['Address'])
        self.set_cell_text(patient_table.cell(3, 3), patient_data['Contact'])
       
        # Fill test details in the second table
        test_table = tables[1]
        for _ in range(len(test_table.rows) - 1):
            test_table._element.remove(test_table.rows[1]._element)
       
        for i in range(10):
            row_cells = test_table.add_row().cells
            if i < len(self.edit_tests):
                test_name, amount = self.edit_tests[i]
                self.set_cell_text(row_cells[0], test_name)
                self.set_cell_text(row_cells[1], f"₹{amount:.2f}")
            else:
                self.set_cell_text(row_cells[0], "")
                self.set_cell_text(row_cells[1], "")
       
        # Fill amounts in the third table
        total_table = tables[2]
        self.set_cell_text(total_table.cell(0, 1), f"₹{total_amount:.2f}")
        self.set_cell_text(total_table.cell(1, 1), f"₹{discount:.2f}")
        self.set_cell_text(total_table.cell(2, 1), f"₹{advanced:.2f}")
        self.set_cell_text(total_table.cell(3, 1), f"₹{due_amount:.2f}")
       
        # --- FIX: Check for Permission Error before saving ---
        docx_file = os.path.join(BILL_FOLDER, f"{original_bill_number}.docx")
        try:
            document.save(docx_file)
        except PermissionError:
            messagebox.showerror("File Open Error", 
                               f"The file '{original_bill_number}.docx' is currently open.\n\nPlease close Microsoft Word and try again.")
            return
        except Exception as e:
            messagebox.showerror("Error", f"Could not save file: {e}")
            return

        # Convert to PDF
        pdf_file = docx_file.replace(".docx", ".pdf")
        try:
            self.convert_docx_to_pdf(docx_file, pdf_file)
            single_pdf = docx_file.replace(".docx", "_single.pdf")
            self.convert_docx_to_pdf(docx_file, single_pdf)
       
            # Merge two copies
            from PyPDF2 import PdfMerger
            merger = PdfMerger()
            merger.append(single_pdf)
            merger.append(single_pdf)
       
            two_page_pdf = single_pdf.replace("_single.pdf", "_2page.pdf")
            merger.write(two_page_pdf)
            merger.close()
       
            webbrowser.open(two_page_pdf)
       
            # Clean up temporary files
            if os.path.exists(single_pdf):
                os.remove(single_pdf)
                
        except Exception as e:
            messagebox.showerror("PDF Error", f"Could not generate PDF: {e}")
            return
       
        # Update Excel data
        self.update_excel_data(original_bill_number, patient_data, total_amount, discount, advanced, due_amount)
       
        # Close edit window
        self.edit_bill_window.destroy()
       
        # Refresh history view
        if hasattr(self, 'history_window') and self.history_window.winfo_exists():
            self.load_history_data()

    def update_excel_data(self, bill_number, patient_data, total_amount, discount, advanced, due):
        """Update the Excel file with edited data"""
        try:
            df = pd.read_excel(EXCEL_FILE)
        
            # Find the row to update
            mask = df['Bill Number'] == bill_number
            if not mask.any():
                return
        
            # Update the row
            df.loc[mask, 'Date'] = datetime.today().strftime("%d-%m-%Y")
            df.loc[mask, 'Patient Name'] = patient_data['Patient Name']
            df.loc[mask, 'Age'] = patient_data['Age']
            df.loc[mask, 'Gender'] = patient_data['Gender']
            df.loc[mask, 'Ref By'] = patient_data['Ref By']
            df.loc[mask, 'Agent'] = patient_data['Agent']
            df.loc[mask, 'Address'] = patient_data['Address']
          
            df.loc[mask, 'Contact'] = patient_data['Contact']
            df.loc[mask, 'Total Amount'] = total_amount
            df.loc[mask, 'Discount'] = discount
            df.loc[mask, 'Advanced Payment'] = advanced
            df.loc[mask, 'Due Amount'] = due
            df.loc[mask, 'Due Payment Date'] = patient_data.get('Due Payment Date', '')
            df.loc[mask, 'Tests'] = ", ".join([f"{t[0]} ({t[1]})" for t in self.edit_tests])
        
            # Save back to Excel
            df.to_excel(EXCEL_FILE, index=False, engine="openpyxl")
        
        except Exception as e:
            messagebox.showerror("Error", f"Failed to update Excel: {str(e)}")

    def pay_due(self, bill_number, due_amount):
        """Handle payment of due amount"""
        # Create payment window
        pay_window = ctk.CTkToplevel(self)
        pay_window.title(f"Pay Due - {bill_number}")
        pay_window.geometry("400x200")
        pay_window.attributes('-topmost', True)
    
        # Payment amount
        ctk.CTkLabel(pay_window, text=f"Due Amount: ₹{due_amount:.2f}").pack(pady=10)

        ctk.CTkLabel(pay_window, text="Amount Paid:").pack()
        self.paid_amount_var = ctk.StringVar(value=str(due_amount))
        paid_entry = ctk.CTkEntry(pay_window, textvariable=self.paid_amount_var)
        paid_entry.pack(pady=5)
    
        # Buttons
        btn_frame = ctk.CTkFrame(pay_window)
        btn_frame.pack(pady=10)
    
        ctk.CTkButton(
            btn_frame, 
            text="Confirm Payment", 
            command=lambda: self.confirm_payment(bill_number, due_amount, pay_window),
            fg_color="#2E8B57"
        ).pack(side="left", padx=5)
    
        ctk.CTkButton(
            btn_frame, 
            text="Cancel", 
            command=pay_window.destroy
        ).pack(side="right", padx=5)

    def confirm_payment(self, bill_number, due_amount, pay_window):
        """Confirm the payment and update records"""
        try:
            paid_amount = float(self.paid_amount_var.get())
       
            if paid_amount <= 0:
                messagebox.showerror("Error", "Payment amount must be positive")
                return
           
            if paid_amount > due_amount:
                messagebox.showerror("Error", "Payment amount cannot be more than due amount")
                return
           
            # Update Excel data
            df = pd.read_excel(EXCEL_FILE)
            mask = df['Bill Number'] == bill_number
       
            if not mask.any():
                messagebox.showerror("Error", "Bill not found")
                return
           
            # Update advanced payment and due amount
            current_advanced = float(df.loc[mask, 'Advanced Payment'].values[0])
            new_advanced = current_advanced + paid_amount
            new_due = float(df.loc[mask, 'Total Amount'].values[0]) - float(df.loc[mask, 'Discount'].values[0]) - new_advanced
       
            df.loc[mask, 'Advanced Payment'] = new_advanced
            df.loc[mask, 'Due Amount'] = max(0, new_due)  # Ensure due doesn't go negative
            df.loc[mask, 'Date'] = datetime.today().strftime("%d-%m-%Y")  # Update date
            df.loc[mask, 'Due Payment Date'] = datetime.today().strftime("%d-%m-%Y")  # Add payment date
       
            # Save back to Excel
            df.to_excel(EXCEL_FILE, index=False, engine="openpyxl")
       
            # Close payment window
            pay_window.destroy()
       
            # --- FIX: Auto-Refresh the History Page ---
            # Instead of looking for a popup window, we update the frame directly.
            if hasattr(self, 'history_page'):
                self.history_page.load_history_data()
           
            messagebox.showinfo("Success", f"Payment of ₹{paid_amount:.2f} recorded on {datetime.today().strftime('%d-%m-%Y')}")
            
        except ValueError:
            messagebox.showerror("Error", "Invalid payment amount")

    def download_excel_copy(self):
        """Save a copy of the Excel file to user-specified location"""
        if not os.path.exists(EXCEL_FILE):
            messagebox.showerror("Error", "No Excel file exists to download", parent=self)
            return

        # FIX: Use 'self' as parent instead of the deleted 'self.history_window'
        file_path = filedialog.asksaveasfilename(
            parent=self, 
            defaultextension=".xlsx",
            filetypes=[("Excel Files", "*.xlsx"), ("All Files", "*.*")],
            title="Save Excel Copy As",
            initialfile=f"bills_backup_{datetime.today().strftime('%Y%m%d')}.xlsx"
        )

        if not file_path:  # User cancelled
            return

        try:
            # Read the original Excel file
            df = pd.read_excel(EXCEL_FILE)
       
            # Save to the new location
            df.to_excel(file_path, index=False, engine="openpyxl")
       
            # Adjust column widths in the new file (Optional styling)
            try:
                wb = openpyxl.load_workbook(file_path)
                ws = wb.active
                for col in ws.columns:
                    max_length = 0
                    col_letter = col[0].column_letter
                    for cell in col:
                        try:
                            if cell.value:
                                max_length = max(max_length, len(str(cell.value)))
                        except:
                            pass
                    adjusted_width = max_length + 2
                    ws.column_dimensions[col_letter].width = adjusted_width
                wb.save(file_path)
            except Exception as e:
                print(f"Warning: Could not adjust column widths - {e}")

            # FIX: Use standard messagebox which is safer and simpler
            messagebox.showinfo("Success", f"Excel file saved successfully to:\n{file_path}", parent=self)
       
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save Excel file:\n{str(e)}", parent=self)


    def center_window(self, window):
        """Center a window on the screen"""
        window.update_idletasks()
        width = window.winfo_width()
        height = window.winfo_height()
        x = (window.winfo_screenwidth() // 2) - (width // 2)
        y = (window.winfo_screenheight() // 2) - (height // 2)
        window.geometry(f'+{x}+{y}')
            



if __name__ == "__main__":
    app = BillingApp()
    app.mainloop()